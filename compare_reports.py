#!/usr/bin/env python3
"""Compare v2 vs v1 Word reports.

Usage:
    python3 compare_reports.py --client "ClientName" --oferta 1
    python3 compare_reports.py --client "ClientName"  # All offers
"""

import sys
import argparse
from pathlib import Path
from docx import Document
from typing import Dict, Tuple

ROOT = Path(__file__).parent
OUTPUT_BASE = ROOT / "output_AO"


def extract_report_text(docx_path: Path) -> str:
    """Extract all text from a DOCX file.

    Args:
        docx_path: Path to .docx file

    Returns:
        Full text content (paragraphs + tables)
    """
    if not docx_path.exists():
        return ""

    doc = Document(str(docx_path))
    lines = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)

    return "\n".join(lines)


def extract_metrics(text: str) -> Dict[str, str]:
    """Extract summary metrics from report text.

    Args:
        text: Full report text

    Returns:
        Dict with metrics (key=metric_name, value=metric_value)
    """
    metrics = {}

    # Parse lines for "Metric: Value" or "Metric | Value" patterns
    for line in text.split("\n"):
        line = line.strip()

        # Skip empty lines
        if not line:
            continue

        # Parse "Key: Value" or "Key | Value"
        if ":" in line:
            parts = line.split(":", 1)
        elif "|" in line:
            parts = line.split("|", 1)
        else:
            continue

        if len(parts) == 2:
            key = parts[0].strip()
            val = parts[1].strip()
            if key and val:
                # Normalize key
                key_lower = key.lower()
                if any(x in key_lower for x in ["matched", "ref", "oferta", "count", "group", "article"]):
                    metrics[key] = val

    return metrics


def compare_reports(v1_path: Path, v2_path: Path) -> Tuple[bool, Dict]:
    """Compare two reports and return results.

    Args:
        v1_path: Path to v1 report
        v2_path: Path to v2 report

    Returns:
        (is_identical: bool, differences: dict)
    """
    # Extract text
    v1_text = extract_report_text(v1_path)
    v2_text = extract_report_text(v2_path)

    if not v1_text:
        return False, {"error": f"v1 report not found: {v1_path}"}
    if not v2_text:
        return False, {"error": f"v2 report not found: {v2_path}"}

    # Extract metrics
    v1_metrics = extract_metrics(v1_text)
    v2_metrics = extract_metrics(v2_text)

    # Compare
    differences = {}

    # Check all metrics from both reports
    all_keys = set(v1_metrics.keys()) | set(v2_metrics.keys())

    for key in sorted(all_keys):
        v1_val = v1_metrics.get(key, "N/A")
        v2_val = v2_metrics.get(key, "N/A")

        if v1_val != v2_val:
            differences[key] = {"v1": v1_val, "v2": v2_val}

    # Report is identical if no differences in metrics
    is_identical = len(differences) == 0

    return is_identical, differences


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Compare v2 vs v1 Word reports",
    )
    parser.add_argument("--client", required=True, help="Client name")
    parser.add_argument("--oferta", type=int, default=None, help="Offer number (1, 2, ...)")

    args = parser.parse_args()

    client_name = args.client
    client_dir = OUTPUT_BASE / client_name

    if not client_dir.exists():
        print(f"Error: Client directory not found: {client_dir}")
        sys.exit(1)

    # Determine offers to compare
    if args.oferta is not None:
        offers = [args.oferta]
    else:
        # Find all offers
        v1_reports = list(client_dir.glob("Raport_Oferta_*.docx"))
        v1_reports = [
            p for p in v1_reports
            if "_v2" not in p.name  # Exclude v2 reports
        ]
        offers = []
        for p in v1_reports:
            # Extract offer number from "Raport_Oferta_N.docx"
            try:
                num = int(p.name.split("_")[2].split(".")[0])
                offers.append(num)
            except (ValueError, IndexError):
                pass
        offers = sorted(set(offers))

    if not offers:
        print(f"No offers found for {client_name}")
        sys.exit(1)

    # Compare each offer
    total_identical = 0
    total_mismatches = 0
    all_differences = {}

    for oferta_num in offers:
        v1_path = client_dir / f"Raport_Oferta_{oferta_num}.docx"
        v2_path = client_dir / f"Raport_Oferta_{oferta_num}_v2.docx"

        is_identical, differences = compare_reports(v1_path, v2_path)

        if is_identical:
            print(f"✓ Oferta {oferta_num}: IDENTICAL")
            total_identical += 1
        else:
            print(f"✗ Oferta {oferta_num}: MISMATCH")
            total_mismatches += 1
            all_differences[oferta_num] = differences

            # Print differences
            if "error" in differences:
                print(f"  Error: {differences['error']}")
            else:
                for key, diffs in differences.items():
                    print(f"  {key}:")
                    print(f"    v1: {diffs['v1']}")
                    print(f"    v2: {diffs['v2']}")

    # Summary
    print()
    print(f"Summary: {total_identical} identical, {total_mismatches} mismatches")

    # Exit code
    if total_mismatches == 0:
        print("All reports match!")
        sys.exit(0)
    else:
        print("Some reports differ.")
        sys.exit(1)


if __name__ == "__main__":
    main()
