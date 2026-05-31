#!/usr/bin/env python3
"""
Generate Word reports from v2 extraction JSON (oferta_N_v2.json).
Usage: python3 generate_v2_reports.py "Drum Tatarani"
"""
import json
import sys
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
OUTPUT_BASE = ROOT / "output_AO"

def load_json(path):
    with open(path) as f:
        return json.load(f)

def generate_v2_report(client_name, oferta_num):
    client_dir = OUTPUT_BASE / client_name
    v2_file = client_dir / f"oferta_{oferta_num}_v2.json"

    if not v2_file.exists():
        print(f"Missing: {v2_file}")
        return

    data = load_json(v2_file)

    # Create Word doc
    doc = Document()
    doc.add_heading(f"Raport Extraction v2 - Oferta {oferta_num}", 0)
    doc.add_paragraph(f"Client: {client_name}")
    doc.add_paragraph(f"Method: Table-Native Extraction (v2)")
    doc.add_paragraph(f"Date: {date.today()}")

    # Metadata
    doc.add_heading("Metadata", level=1)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Light Grid Accent 1"

    meta_table.rows[0].cells[0].text = "Template ID"
    meta_table.rows[0].cells[1].text = data.get("template_id", "UNKNOWN")

    meta_table.rows[1].cells[0].text = "Extraction Version"
    meta_table.rows[1].cells[1].text = data.get("extraction_version", "2.0")

    meta_table.rows[2].cells[0].text = "Total Groups"
    total_groups = len(data.get("grupos", []))
    meta_table.rows[2].cells[1].text = str(total_groups)

    meta_table.rows[3].cells[0].text = "Total Articles"
    total_articles = sum(len(g.get("articole", [])) for g in data.get("grupos", []))
    meta_table.rows[3].cells[1].text = str(total_articles)

    # Articles by group
    doc.add_heading("Articles by Group", level=1)

    for grupo in data.get("grupos", []):
        deviz_cod = grupo.get("deviz_cod", "UNKNOWN")
        deviz_den = grupo.get("deviz_den", "")
        articole = grupo.get("articole", [])

        heading = f"{deviz_cod} - {deviz_den[:60]}" if deviz_den else deviz_cod
        doc.add_heading(heading, level=2)

        # Group summary
        summary = doc.add_paragraph()
        summary.add_run(f"{len(articole)} articles\n")

        sources = {}
        for art in articole:
            src = art.get("extraction_source", "UNKNOWN")
            sources[src] = sources.get(src, 0) + 1

        summary.add_run("Sources: ")
        summary.add_run(", ".join([f"{src}={count}" for src, count in sorted(sources.items())]))

        # Articles table
        if articole:
            art_table = doc.add_table(rows=len(articole) + 1, cols=5)
            art_table.style = "Light Grid Accent 1"

            # Header
            hdr = art_table.rows[0].cells
            hdr[0].text = "NR"
            hdr[1].text = "Descriere"
            hdr[2].text = "UM"
            hdr[3].text = "Cant"
            hdr[4].text = "Source"

            # Data rows
            for i, art in enumerate(articole):
                row = art_table.rows[i + 1].cells
                row[0].text = art.get("nr") or art.get("cod") or ""
                row[1].text = (art.get("descriere") or "")[:80]
                row[2].text = art.get("um") or ""
                row[3].text = str(art.get("cant") or art.get("cantitate") or "")
                row[4].text = art.get("extraction_source", "UNKNOWN")

    # Save
    output_path = client_dir / f"Raport_Oferta_{oferta_num}_v2.docx"
    doc.save(output_path)
    print(f"✓ Saved: {output_path.name}")

def generate_v2_report_from_holistic(holistic_data, output_path, client_name, oferta_num):
    """Generate v2 report from holistic matching data."""
    doc = Document()
    doc.add_heading(f"Raport Extraction v2 - Oferta {oferta_num}", 0)
    doc.add_paragraph(f"Client: {client_name}")
    doc.add_paragraph(f"Method: v2 Extraction (Table-Native) + v1 Matching")
    doc.add_paragraph(f"Date: {date.today()}")

    # Summary
    doc.add_heading("Summary", level=1)
    matched = len([g for g in holistic_data.get("matched_groups", []) if g.get("articole")])
    ref_only = len(holistic_data.get("ref_only_groups", []))
    oferta_only = len(holistic_data.get("oferta_only_groups", []))

    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Light Grid Accent 1"
    summary_table.rows[0].cells[0].text = "Matched Groups"
    summary_table.rows[0].cells[1].text = str(matched)
    summary_table.rows[1].cells[0].text = "Reference-Only Groups"
    summary_table.rows[1].cells[1].text = str(ref_only)
    summary_table.rows[2].cells[0].text = "Oferta-Only Groups"
    summary_table.rows[2].cells[1].text = str(oferta_only)
    summary_table.rows[3].cells[0].text = "Total Groups"
    summary_table.rows[3].cells[1].text = str(matched + ref_only + oferta_only)

    # Matched groups
    doc.add_heading("Matched Groups", level=1)
    for grupo in holistic_data.get("matched_groups", [])[:10]:  # First 10
        deviz = grupo.get("deviz_cod", "UNKNOWN")
        articole_count = len(grupo.get("articole", []))
        doc.add_paragraph(f"{deviz}: {articole_count} articles")

    # Save
    doc.save(output_path)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 generate_v2_reports.py 'ClientName'")
        sys.exit(1)

    client = sys.argv[1]
    for oferta in [1, 2]:
        generate_v2_report(client, oferta)
