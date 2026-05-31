"""
Generate Word reports showing hierarchy corrections per group.

Entry point: generate_hierarchy_report(client_name, oferta_num)
Output: output_AO/{ClientName}/Raport_Hierarchy_{oferta_num}_v2.docx
"""

import json
import os
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# Colors
BLACK = RGBColor(0x00, 0x00, 0x00)
GREEN_FILL = "92D050"  # Green for PASS
RED_FILL = "FF6B6B"    # Red for FAIL
GRAY_FILL = "D9D9D9"   # Gray for header


def _set_cell_shading(cell, fill_hex: str):
    """Set background color of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}" w:val="clear"/>')
    tcPr.append(shd)


def _style_cell(cell, size_pt: float, bold: bool = False,
                color: RGBColor = None, center: bool = False):
    """Apply font size, bold, color, and alignment to all runs in a cell."""
    for para in cell.paragraphs:
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(size_pt)
            run.font.color.rgb = color or BLACK
            if bold:
                run.font.bold = True


def _load_oferta_v2(client_name: str, oferta_num: int) -> Optional[dict]:
    """Load oferta_{oferta_num}_v2.json file."""
    base_dir = Path("/Users/gabrielchitu/analiza-oferte-local/output_AO")
    client_dir = base_dir / client_name
    filepath = client_dir / f"oferta_{oferta_num}_v2.json"

    if not filepath.exists():
        print(f"Warning: File not found: {filepath}")
        return None

    with open(filepath, 'r', encoding='utf-8') as f:
        return json.load(f)


def _extract_hierarchy_stats(grupos: list) -> dict:
    """
    Extract hierarchy correction statistics from grupos.

    Returns:
    {
        "total_detected": int,
        "total_corrected": int,
        "groups": [
            {
                "deviz_cod": str,
                "deviz_den": str,
                "detected_count": int,
                "corrected_count": int,
                "validation": "PASS" | "FAIL",
                "message": str
            }
        ]
    }
    """
    total_detected = 0
    total_corrected = 0
    groups = []

    for grupo in grupos:
        deviz_cod = grupo.get("deviz_cod", "UNKNOWN")
        deviz_den = grupo.get("deviz_den", "")
        articole = grupo.get("articole", [])

        # Count detected hierarchy issues (components without explicit markers)
        detected = 0
        corrected = 0

        for art in articole:
            is_component = art.get("is_component", False)
            has_marker = art.get("explicit_component_marker", False)
            parent_code = art.get("parent_code")

            if is_component:
                detected += 1
                if has_marker or parent_code:
                    corrected += 1

        # Determine validation status
        validation = "PASS" if (detected == 0 or corrected == detected) else "FAIL"
        message = f"Detected: {detected}, Corrected: {corrected}"

        groups.append({
            "deviz_cod": deviz_cod,
            "deviz_den": deviz_den,
            "detected_count": detected,
            "corrected_count": corrected,
            "validation": validation,
            "message": message
        })

        total_detected += detected
        total_corrected += corrected

    return {
        "total_detected": total_detected,
        "total_corrected": total_corrected,
        "groups": groups
    }


def _create_summary_section(doc: Document, client_name: str, oferta_num: int,
                            stats: dict):
    """Add summary section to document."""
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(f"Raport Corectii Ierarhie - Oferta {oferta_num}")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = BLACK
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph(f"Client: {client_name}")
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(f"Data raportului: {date.today().strftime('%d.%m.%Y')}")
    date_para_run = date_para.runs[0]
    date_para_run.font.size = Pt(10)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Blank line

    # Summary statistics
    summary = doc.add_paragraph()
    summary_run = summary.add_run(
        f"Total subcomponente detectate: {stats['total_detected']} | "
        f"Corectate: {stats['total_corrected']}"
    )
    summary_run.font.size = Pt(11)
    summary_run.font.bold = True
    summary.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # Blank line


def _create_groups_table(doc: Document, groups: list):
    """Create table with per-group statistics."""
    # Table header
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    headers = ["Deviz Cod", "Deviz Denumire", "Detectate", "Corectate", "Status"]

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        _set_cell_shading(cell, GRAY_FILL)
        _style_cell(cell, 10, bold=True, center=True)

    # Data rows
    for group in groups:
        row = table.add_row()
        cells = row.cells

        # Deviz Cod
        cells[0].text = group["deviz_cod"]
        _style_cell(cells[0], 10, center=False)

        # Deviz Den
        cells[1].text = group["deviz_den"][:50] if group["deviz_den"] else ""
        _style_cell(cells[1], 10, center=False)

        # Detected
        cells[2].text = str(group["detected_count"])
        _style_cell(cells[2], 10, center=True)

        # Corrected
        cells[3].text = str(group["corrected_count"])
        _style_cell(cells[3], 10, center=True)

        # Status
        status = group["validation"]
        cells[4].text = status
        status_color = GREEN_FILL if status == "PASS" else RED_FILL
        _set_cell_shading(cells[4], status_color)
        _style_cell(cells[4], 10, bold=True, center=True)

    # Set column widths
    table.autofit = False
    table.allow_autofit = False
    widths = [Inches(1.2), Inches(2.5), Inches(1.0), Inches(1.0), Inches(0.8)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def generate_hierarchy_report(client_name: str, oferta_num: int) -> bool:
    """
    Generate hierarchy correction report for given client and oferta.

    Args:
        client_name: Name of client (e.g., "Drum Tatarani")
        oferta_num: Oferta number (1-4)

    Returns:
        True if successful, False otherwise
    """
    # Load data
    data = _load_oferta_v2(client_name, oferta_num)
    if not data:
        return False

    # Extract statistics
    stats = _extract_hierarchy_stats(data.get("grupos", []))

    # Create document
    doc = Document()
    doc.sections[0].page_height = Cm(29.7)  # A4
    doc.sections[0].page_width = Cm(21)

    # Add content
    _create_summary_section(doc, client_name, oferta_num, stats)

    if stats["groups"]:
        doc.add_heading("Statistici per Grup", level=2)
        _create_groups_table(doc, stats["groups"])
    else:
        doc.add_paragraph("Nu au fost găsite grupuri cu informații de ierarhie.")

    # Add footer with validation summary
    doc.add_paragraph()
    footer = doc.add_paragraph()
    passed = sum(1 for g in stats["groups"] if g["validation"] == "PASS")
    failed = sum(1 for g in stats["groups"] if g["validation"] == "FAIL")
    footer_text = f"Rezumat: {passed} grupuri PASS, {failed} grupuri FAIL"
    footer_run = footer.add_run(footer_text)
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True

    # Save document
    base_dir = Path("/Users/gabrielchitu/analiza-oferte-local/output_AO")
    client_dir = base_dir / client_name
    client_dir.mkdir(parents=True, exist_ok=True)

    output_path = client_dir / f"Raport_Hierarchy_{oferta_num}_v2.docx"
    doc.save(str(output_path))

    print(f"Report saved: {output_path}")
    return True


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("Usage: python3 generate_hierarchy_report.py <client_name> <oferta_num>")
        print("Example: python3 generate_hierarchy_report.py 'Drum Tatarani' 1")
        sys.exit(1)

    client_name = sys.argv[1]
    try:
        oferta_num = int(sys.argv[2])
    except ValueError:
        print("Error: oferta_num must be an integer")
        sys.exit(1)

    success = generate_hierarchy_report(client_name, oferta_num)
    sys.exit(0 if success else 1)
