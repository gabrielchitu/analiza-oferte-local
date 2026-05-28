#!/usr/bin/env python3
"""
di_to_docx.py — Convert Azure DI JSON files to human-readable DOCX.

Usage:
    python3 di_to_docx.py --client "BR BLOC A"

Output:
    output_AO/<client>/DI_Referinta.docx
    output_AO/<client>/DI_Oferta_1.docx
    ...
"""
import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from shared.client_config import ClientConfig

INPUT_BASE = Path("input_AO")
OUTPUT_BASE = Path("output_AO")

GRAY_HEADER = "D9D9D9"
GRAY_SUBHEADER = "F2F2F2"


def _set_cell_shading(cell, fill_hex: str) -> None:
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _cells_by_row(cells: list) -> dict:
    rows: dict = {}
    for c in cells:
        r, col = c["row_index"], c["column_index"]
        rows.setdefault(r, {})[col] = c.get("content", "")
    return rows


def _render_metadata(doc: Document, table: dict) -> None:
    rows = _cells_by_row(table["cells"])
    for r_idx in sorted(rows):
        row = rows[r_idx]
        key = row.get(0, "").strip()
        val = row.get(1, "").strip()
        if not key:
            continue
        p = doc.add_paragraph()
        run_k = p.add_run(key + "  ")
        run_k.bold = True
        p.add_run(val)


def _render_f3_table(doc: Document, table: dict) -> None:
    row_count = table["row_count"]
    col_count = table["column_count"]
    if row_count == 0 or col_count < 2:
        print(f"  [WARN] Tabel ignorat: rows={row_count} cols={col_count}")
        return

    cells_by_row = _cells_by_row(table["cells"])
    tbl = doc.add_table(rows=row_count, cols=col_count)
    tbl.style = "Table Grid"

    for r_idx in range(row_count):
        row_data = cells_by_row.get(r_idx, {})
        doc_row = tbl.rows[r_idx]
        for c_idx in range(col_count):
            content = row_data.get(c_idx, "")
            cell = doc_row.cells[c_idx]
            cell.text = content
            runs = cell.paragraphs[0].runs
            if r_idx == 0:
                _set_cell_shading(cell, GRAY_HEADER)
                if runs:
                    runs[0].bold = True
                    runs[0].font.size = Pt(9)
            elif r_idx == 1:
                _set_cell_shading(cell, GRAY_SUBHEADER)
                if runs:
                    runs[0].bold = True
                    runs[0].font.size = Pt(9)
            else:
                if runs:
                    runs[0].font.size = Pt(9)


def _render_pages(doc: Document, pages: list) -> None:
    for page in pages:
        p = doc.add_paragraph(f"--- Pagina {page['page_number']} ---")
        p.style = "Caption"
        for line in page.get("lines", []):
            doc.add_paragraph(line.get("content", ""))


def convert_di_to_docx(di_path: Path, out_path: Path) -> None:
    try:
        data = json.loads(di_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, UnicodeDecodeError) as e:
        print(f"[ERROR] JSON invalid: {di_path}: {e}")
        sys.exit(1)
    doc = Document()

    doc.add_heading(di_path.stem, level=0)

    pages = data.get("pages", [])
    tables = data.get("tables", [])

    if pages:
        _render_pages(doc, pages)

    doc.add_page_break()
    doc.add_heading("Tabele", level=1)

    if tables:
        _render_metadata(doc, tables[0])
        for i, table in enumerate(tables[1:], start=1):
            doc.add_heading(f"Tabel {i}", level=2)
            _render_f3_table(doc, table)

    doc.save(out_path)
    print(f"  [OK] {out_path.name}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert DI JSON to readable DOCX")
    parser.add_argument("--client", required=True, help="Client name (must match input_AO folder)")
    args = parser.parse_args()

    cfg = ClientConfig.from_folder(args.client, INPUT_BASE, OUTPUT_BASE)
    if not cfg.input_dir.exists():
        print(f"[ERROR] Client '{args.client}' not found in {INPUT_BASE}/")
        sys.exit(1)

    cfg.ensure_output_dirs()

    if cfg.reference_file.exists():
        print(f"[DI] Convertesc {cfg.reference_file.name}...")
        convert_di_to_docx(cfg.reference_file, cfg.output_dir / "DI_Referinta.docx")
    else:
        print(f"[WARN] {cfg.reference_file} lipsește, skip")

    for di_path in cfg.offer_files:
        n = di_path.stem.split("_")[-1]
        print(f"[DI] Convertesc {di_path.name}...")
        convert_di_to_docx(di_path, cfg.output_dir / f"DI_Oferta_{n}.docx")


if __name__ == "__main__":
    main()
