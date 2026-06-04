#!/usr/bin/env python3
"""
Generate Word reports from v2 matching (holistic JSON).
Usage: python3 generate_matching_report.py "ClientName" [oferta_num]
"""
import json
import sys
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
OUTPUT_BASE = ROOT / "output_AO"

def load_json(path):
    """Load JSON file."""
    with open(path) as f:
        return json.load(f)

def generate_matching_report(client_name, oferta_num):
    """Generate Word report from holistic matching JSON."""
    client_dir = OUTPUT_BASE / client_name

    # Check both v2 and v1 (fallback) formats
    holistic_file_v2 = client_dir / f"holistic_oferta_{oferta_num}_v2.json"
    holistic_file_v1 = client_dir / f"holistic_oferta_{oferta_num}.json"

    if holistic_file_v2.exists():
        holistic_file = holistic_file_v2
        version = "v2"
    elif holistic_file_v1.exists():
        holistic_file = holistic_file_v1
        version = "v1"
    else:
        print(f"Missing holistic JSON for {client_name} oferta {oferta_num}")
        return

    holistic = load_json(holistic_file)

    # Create Word doc
    doc = Document()
    doc.add_heading(f"Raport Matching - Oferta {oferta_num}", 0)
    doc.add_paragraph(f"Client: {client_name}")
    doc.add_paragraph(f"Method: Set-Based Matching ({version})")
    doc.add_paragraph(f"Date: {date.today()}")

    # Extract summary stats
    matched_groups = holistic.get('matched_groups', [])
    ref_only_groups = holistic.get('ref_only_groups', [])
    oferta_only_groups = holistic.get('oferta_only_groups', [])

    # Count matched articles
    matched_articles = sum(
        len(g.get('articole', [])) or len(g.get('ref_articles', []))
        for g in matched_groups
    )

    # Count ref-only articles
    ref_only_articles = sum(
        len(g.get('ref_articles', [])) or len(g.get('articole', []))
        for g in ref_only_groups
    )

    # Count oferta-only articles
    oferta_only_articles = sum(
        len(g.get('oferta_articles', [])) or len(g.get('articole', []))
        for g in oferta_only_groups
    )

    # Summary section
    doc.add_heading("Summary", level=1)
    stats = holistic.get('stats', {})

    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = "Light Grid Accent 1"

    summary_table.rows[0].cells[0].text = "Matched Groups"
    summary_table.rows[0].cells[1].text = str(len(matched_groups))

    summary_table.rows[1].cells[0].text = "Reference-Only Groups"
    summary_table.rows[1].cells[1].text = str(len(ref_only_groups))

    summary_table.rows[2].cells[0].text = "Oferta-Only Groups"
    summary_table.rows[2].cells[1].text = str(len(oferta_only_groups))

    summary_table.rows[3].cells[0].text = "Matched Articles"
    summary_table.rows[3].cells[1].text = str(matched_articles)

    summary_table.rows[4].cells[0].text = "Reference-Only Articles"
    summary_table.rows[4].cells[1].text = str(ref_only_articles)

    summary_table.rows[5].cells[0].text = "Oferta-Only Articles"
    summary_table.rows[5].cells[1].text = str(oferta_only_articles)

    total_articles = matched_articles + ref_only_articles + oferta_only_articles
    summary_table.rows[6].cells[0].text = "Total Articles"
    summary_table.rows[6].cells[1].text = str(total_articles)

    # Matched groups section
    doc.add_heading("Matched Groups", level=1)
    doc.add_paragraph(f"Total: {len(matched_groups)} groups")

    for i, group in enumerate(matched_groups[:10]):  # First 10
        deviz_cod = group.get('ref_deviz_cod') or group.get('deviz_cod', 'UNKNOWN')
        deviz_den = group.get('deviz_denumire', '')

        ref_articles = group.get('ref_articles', [])
        oferta_articles = group.get('oferta_articles', [])
        matched = group.get('articole', [])

        heading = f"{deviz_cod} - {deviz_den[:60]}" if deviz_den else deviz_cod
        doc.add_heading(heading, level=2)

        summary = doc.add_paragraph()
        summary.add_run(f"Matched: {len(matched)} articles\n")
        summary.add_run(f"Ref-Only: {len(ref_articles)} articles\n")
        summary.add_run(f"Oferta-Only: {len(oferta_articles)} articles")

    # Reference-only groups section
    doc.add_heading("Reference-Only Groups", level=1)
    doc.add_paragraph(f"Total: {len(ref_only_groups)} groups")

    if ref_only_groups:
        for i, group in enumerate(ref_only_groups[:5]):  # First 5
            deviz_cod = group.get('ref_deviz_cod') or group.get('deviz_cod', 'UNKNOWN')
            deviz_den = group.get('deviz_denumire', '')
            ref_articles = group.get('ref_articles', [])

            heading = f"{deviz_cod} - {deviz_den[:60]}" if deviz_den else deviz_cod
            doc.add_heading(heading, level=2)
            doc.add_paragraph(f"{len(ref_articles)} reference articles")

    # Oferta-only groups section
    doc.add_heading("Oferta-Only Groups", level=1)
    doc.add_paragraph(f"Total: {len(oferta_only_groups)} groups")

    if oferta_only_groups:
        for i, group in enumerate(oferta_only_groups[:5]):  # First 5
            deviz_cod = group.get('oferta_deviz_cod') or group.get('deviz_cod', 'UNKNOWN')
            deviz_den = group.get('deviz_denumire', '')
            oferta_articles = group.get('oferta_articles', [])

            heading = f"{deviz_cod} - {deviz_den[:60]}" if deviz_den else deviz_cod
            doc.add_heading(heading, level=2)
            doc.add_paragraph(f"{len(oferta_articles)} oferta articles")

    # Save
    output_path = client_dir / f"Raport_Matching_Oferta_{oferta_num}.docx"
    doc.save(output_path)
    print(f"✓ Saved: {output_path.name}")
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 generate_matching_report.py 'ClientName' [oferta_num]")
        sys.exit(1)

    client = sys.argv[1]
    oferta = int(sys.argv[2]) if len(sys.argv) > 2 else 1

    generate_matching_report(client, oferta)
