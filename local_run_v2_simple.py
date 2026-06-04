#!/usr/bin/env python3
"""
v2 extraction pipeline (no matching, extraction only).
Usage: python3 local_run_v2_simple.py --client "Drum Tatarani"
"""
import argparse
import json
import logging
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from shared.client_config import ClientConfig
from shared.extraction_v2 import ExtractionOrchestrator
from generate_v2_reports import generate_v2_report

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

ROOT = Path(__file__).parent
OUTPUT_BASE = ROOT / "output_AO"

def run_v2_simple(client_config):
    """Extract (v2 only) + simple reports"""
    client_name = client_config.name
    client_dir = OUTPUT_BASE / client_name
    client_dir.mkdir(parents=True, exist_ok=True)

    logger.info(f"\n{'='*60}")
    logger.info(f"v2 EXTRACTION PIPELINE: {client_name}")
    logger.info(f"{'='*60}\n")

    extractor = ExtractionOrchestrator()

    # Reference
    logger.info("Extracting reference (v2)")
    with open(client_config.reference_file) as f:
        ref_di = json.load(f)
    ref_extracted = extractor.extract_from_di(ref_di, client_name)
    ref_extracted["di_file"] = "di_referinta.json"

    ref_file = client_dir / "referinta_v2_full.json"
    with open(ref_file, "w") as f:
        json.dump(ref_extracted, f, indent=2, ensure_ascii=False)
    
    ref_articles = sum(len(g.get("articole", [])) for g in ref_extracted.get("grupos", []))
    logger.info(f"✓ Reference: {ref_articles} articles\n")

    # Offers
    offer_files = sorted(client_config.input_dir.glob("di_oferta_*.json"))
    for oferta_file in offer_files:
        match = oferta_file.name.replace("di_oferta_", "").replace(".json", "")
        try:
            oferta_num = int(match)
        except:
            continue

        logger.info(f"Extracting oferta {oferta_num}")
        with open(oferta_file) as f:
            oferta_di = json.load(f)

        oferta_extracted = extractor.extract_from_di(oferta_di, client_name)
        oferta_extracted["di_file"] = f"di_oferta_{oferta_num}.json"

        oferta_out = client_dir / f"oferta_{oferta_num}_v2_full.json"
        with open(oferta_out, "w") as f:
            json.dump(oferta_extracted, f, indent=2, ensure_ascii=False)

        oferta_articles = sum(len(g.get("articole", [])) for g in oferta_extracted.get("grupos", []))
        logger.info(f"✓ Oferta {oferta_num}: {oferta_articles} articles")

        # Generate report
        report_file = client_dir / f"Raport_Oferta_{oferta_num}_v2_full.docx"
        try:
            generate_v2_report_from_holistic(oferta_extracted, str(report_file), client_name, oferta_num)
            logger.info(f"  ✓ Report: {report_file.name}\n")
        except Exception as e:
            logger.warning(f"  ✗ Report failed: {e}\n")

    logger.info(f"{'='*60}")
    logger.info(f"✓ v2 Extraction Pipeline Complete")
    logger.info(f"{'='*60}\n")

def generate_v2_report_from_holistic(extracted_data, output_path, client_name, oferta_num):
    """Generate v2 report from extraction JSON."""
    from datetime import date
    from docx import Document

    doc = Document()
    doc.add_heading(f"Raport Extraction v2 - Oferta {oferta_num}", 0)
    doc.add_paragraph(f"Client: {client_name}")
    doc.add_paragraph(f"Method: v2 Table-Native Extraction (no matching)")
    doc.add_paragraph(f"Date: {date.today()}")

    # Summary
    doc.add_heading("Summary", level=1)
    total_articles = sum(len(g.get("articole", [])) for g in extracted_data.get("grupos", []))
    total_groups = len(extracted_data.get("grupos", []))

    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.style = "Light Grid Accent 1"
    summary_table.rows[0].cells[0].text = "Total Groups"
    summary_table.rows[0].cells[1].text = str(total_groups)
    summary_table.rows[1].cells[0].text = "Total Articles"
    summary_table.rows[1].cells[1].text = str(total_articles)
    summary_table.rows[2].cells[0].text = "Template"
    summary_table.rows[2].cells[1].text = extracted_data.get("template_id", "UNKNOWN")

    # Groups
    doc.add_heading("Groups", level=1)
    for grupo in extracted_data.get("grupos", [])[:20]:
        deviz = grupo.get("deviz_cod", "UNKNOWN")
        den = grupo.get("deviz_den", "")[:50]
        count = len(grupo.get("articole", []))
        doc.add_paragraph(f"{deviz} - {den}: {count} articles")

    doc.save(output_path)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--client", help="Client name")
    args = parser.parse_args()

    INPUT_BASE = ROOT / "input_AO"

    if args.client:
        config = ClientConfig.from_folder(args.client, INPUT_BASE, OUTPUT_BASE)
    else:
        clients = ClientConfig.detect_clients(INPUT_BASE)
        if not clients:
            logger.error("No clients found")
            sys.exit(1)
        for i, c in enumerate(clients, 1):
            print(f"{i}. {c}")
        choice = int(input(f"Select (1-{len(clients)}): ")) - 1
        config = ClientConfig.from_folder(clients[choice], INPUT_BASE, OUTPUT_BASE)

    run_v2_simple(config)
