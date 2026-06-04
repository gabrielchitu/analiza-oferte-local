"""Tests for report comparison tool.

Validates report extraction and comparison logic.
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt

from compare_reports import extract_report_text, extract_metrics, compare_reports


# Fixtures

@pytest.fixture
def temp_docx(tmp_path):
    """Create temporary DOCX with sample content."""
    def _create_docx(filename: str, content_text: str):
        path = tmp_path / filename
        doc = Document()
        doc.add_heading("Test Report", 0)
        doc.add_paragraph(content_text)

        # Add summary table
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Count"
        table.rows[1].cells[0].text = "Matched Articles"
        table.rows[1].cells[1].text = "10"
        table.rows[2].cells[0].text = "Ref-Only Articles"
        table.rows[2].cells[1].text = "5"

        doc.save(str(path))
        return path

    return _create_docx


# Tests

def test_extract_report_text(temp_docx):
    """Test extracting text from DOCX."""
    docx_path = temp_docx("test.docx", "Test content here")

    text = extract_report_text(docx_path)

    assert "Test Report" in text
    assert "Test content here" in text


def test_extract_report_text_includes_tables(temp_docx):
    """Test extracting text includes table content."""
    docx_path = temp_docx("test_table.docx", "Main content")

    text = extract_report_text(docx_path)

    # Should include table headers and content
    assert "Metric" in text
    assert "Count" in text
    assert "Matched Articles" in text
    assert "10" in text


def test_extract_report_text_missing_file():
    """Test extracting text from missing file returns empty string."""
    result = extract_report_text(Path("/nonexistent/file.docx"))
    assert result == ""


def test_extract_metrics(temp_docx):
    """Test extracting metrics from report text."""
    docx_path = temp_docx("metrics.docx", "Some report")

    text = extract_report_text(docx_path)
    metrics = extract_metrics(text)

    assert "Matched Articles" in metrics
    assert metrics["Matched Articles"] == "10"
    assert "Ref-Only Articles" in metrics
    assert metrics["Ref-Only Articles"] == "5"


def test_extract_metrics_empty_text():
    """Test extracting metrics from empty text."""
    metrics = extract_metrics("")
    assert metrics == {}


def test_compare_reports_identical(temp_docx):
    """Test comparing identical reports."""
    docx_path = temp_docx("report1.docx", "Identical content")

    # Create second identical docx
    path1 = temp_docx("report_v1.docx", "Identical content")
    path2 = temp_docx("report_v2.docx", "Identical content")

    is_identical, differences = compare_reports(path1, path2)

    assert is_identical
    assert len(differences) == 0


def test_compare_reports_different(tmp_path):
    """Test comparing reports with different metrics."""
    # Create v1 report
    doc1 = Document()
    doc1.add_heading("Report V1", 0)
    table1 = doc1.add_table(rows=2, cols=2)
    table1.rows[0].cells[0].text = "Matched Articles"
    table1.rows[0].cells[1].text = "10"
    table1.rows[1].cells[0].text = "Extra Articles"
    table1.rows[1].cells[1].text = "5"
    v1_path = tmp_path / "report_v1.docx"
    doc1.save(str(v1_path))

    # Create v2 report with different values
    doc2 = Document()
    doc2.add_heading("Report V2", 0)
    table2 = doc2.add_table(rows=2, cols=2)
    table2.rows[0].cells[0].text = "Matched Articles"
    table2.rows[0].cells[1].text = "12"  # Different value
    table2.rows[1].cells[0].text = "Extra Articles"
    table2.rows[1].cells[1].text = "5"
    v2_path = tmp_path / "report_v2.docx"
    doc2.save(str(v2_path))

    is_identical, differences = compare_reports(v1_path, v2_path)

    assert not is_identical
    assert len(differences) > 0
    assert "Matched Articles" in differences


def test_compare_reports_missing_v1(tmp_path):
    """Test comparing when v1 report is missing."""
    # Create v2 only
    doc2 = Document()
    doc2.add_heading("Report V2", 0)
    v2_path = tmp_path / "report_v2.docx"
    doc2.save(str(v2_path))

    missing_path = tmp_path / "missing_v1.docx"

    is_identical, differences = compare_reports(missing_path, v2_path)

    assert not is_identical
    assert "error" in differences


def test_compare_reports_missing_v2(tmp_path):
    """Test comparing when v2 report is missing."""
    # Create v1 only
    doc1 = Document()
    doc1.add_heading("Report V1", 0)
    v1_path = tmp_path / "report_v1.docx"
    doc1.save(str(v1_path))

    missing_path = tmp_path / "missing_v2.docx"

    is_identical, differences = compare_reports(v1_path, missing_path)

    assert not is_identical
    assert "error" in differences
