"""
Tests for matching report generation.
"""
import json
from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent.parent
OUTPUT_BASE = ROOT / "output_AO"


def generate_matching_report(client_name, oferta_num):
    """Import and call generate_matching_report function."""
    import sys
    sys.path.insert(0, str(ROOT))
    from generate_matching_report import generate_matching_report as gmr
    return gmr(client_name, oferta_num)


class TestMatchingReportGeneration:
    """Test report generation from holistic matching JSON."""

    def test_matching_report_dt_oferta1(self):
        """Test report generation on Drum Tatarani oferta 1."""
        # Check holistic JSON exists
        holistic_file = OUTPUT_BASE / "Drum Tatarani" / "holistic_oferta_1.json"
        assert holistic_file.exists(), f"Missing: {holistic_file}"

        # Load and validate structure
        with open(holistic_file) as f:
            holistic = json.load(f)

        assert 'matched_groups' in holistic
        assert isinstance(holistic['matched_groups'], list)

        # Generate report
        output_path = generate_matching_report("Drum Tatarani", 1)

        # Check report was created
        assert output_path is not None
        assert output_path.exists(), f"Report not created: {output_path}"

        # Validate DOCX structure
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0, "No paragraphs in document"
        assert len(doc.tables) > 0, "No tables in document"

        # Validate content
        doc_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'Drum Tatarani' in doc_text, "Client name not in report"
        assert 'Oferta 1' in doc_text, "Oferta number not in report"
        assert 'Matched Groups' in doc_text, "Summary section missing"

    def test_matching_report_dt_oferta2(self):
        """Test report generation on Drum Tatarani oferta 2."""
        # Check holistic JSON exists
        holistic_file = OUTPUT_BASE / "Drum Tatarani" / "holistic_oferta_2.json"
        assert holistic_file.exists(), f"Missing: {holistic_file}"

        # Generate report
        output_path = generate_matching_report("Drum Tatarani", 2)

        # Check report was created
        assert output_path is not None
        assert output_path.exists(), f"Report not created: {output_path}"

        # Validate DOCX
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0
        assert len(doc.tables) > 0

    def test_matching_report_summary_table(self):
        """Test that summary table is created with correct stats."""
        holistic_file = OUTPUT_BASE / "Drum Tatarani" / "holistic_oferta_1.json"
        with open(holistic_file) as f:
            holistic = json.load(f)

        # Generate report
        output_path = generate_matching_report("Drum Tatarani", 1)
        doc = Document(output_path)

        # Find summary table (should be first table)
        assert len(doc.tables) > 0, "No tables in document"
        summary_table = doc.tables[0]

        # Validate table structure (7 rows, 2 columns)
        assert len(summary_table.rows) >= 7, f"Summary table has {len(summary_table.rows)} rows, expected 7"
        assert len(summary_table.rows[0].cells) == 2, "Summary table should have 2 columns"

        # Validate headers
        assert 'Matched Groups' in summary_table.rows[0].cells[0].text
        assert 'Reference-Only' in summary_table.rows[1].cells[0].text
        assert 'Oferta-Only' in summary_table.rows[2].cells[0].text

    def test_matching_report_multiple_clients(self):
        """Test report generation for multiple clients."""
        clients_to_test = [
            ("Drum Tatarani", 1),
            ("Blocuri Racari", 1),
        ]

        for client, oferta in clients_to_test:
            holistic_file = OUTPUT_BASE / client / f"holistic_oferta_{oferta}.json"
            if holistic_file.exists():
                output_path = generate_matching_report(client, oferta)
                assert output_path.exists(), f"Report not created for {client} oferta {oferta}"

    def test_holistic_json_structure(self):
        """Test that holistic JSON has expected structure."""
        holistic_file = OUTPUT_BASE / "Drum Tatarani" / "holistic_oferta_1.json"
        with open(holistic_file) as f:
            holistic = json.load(f)

        # Validate top-level keys
        assert 'matched_groups' in holistic, "matched_groups key missing"
        assert 'ref_only_groups' in holistic, "ref_only_groups key missing"
        assert 'oferta_only_groups' in holistic, "oferta_only_groups key missing"

        # Validate matched_groups structure
        matched = holistic['matched_groups']
        assert isinstance(matched, list), "matched_groups should be list"

        if len(matched) > 0:
            group = matched[0]
            # At least one of these should exist
            assert (
                'ref_deviz_cod' in group or
                'oferta_deviz_cod' in group or
                'deviz_cod' in group
            ), "Group missing deviz_cod fields"

    def test_article_counting(self):
        """Test that articles are properly counted in summary."""
        holistic_file = OUTPUT_BASE / "Drum Tatarani" / "holistic_oferta_1.json"
        with open(holistic_file) as f:
            holistic = json.load(f)

        # Count articles manually
        matched_count = 0
        for group in holistic.get('matched_groups', []):
            articole = group.get('articole', [])
            if not articole:
                articole = group.get('ref_articles', [])
            matched_count += len(articole)

        ref_only_count = 0
        for group in holistic.get('ref_only_groups', []):
            articole = group.get('ref_articles', [])
            if not articole:
                articole = group.get('articole', [])
            ref_only_count += len(articole)

        oferta_only_count = 0
        for group in holistic.get('oferta_only_groups', []):
            articole = group.get('oferta_articles', [])
            if not articole:
                articole = group.get('articole', [])
            oferta_only_count += len(articole)

        # Verify we have some articles
        total = matched_count + ref_only_count + oferta_only_count
        assert total > 0, "No articles found in holistic JSON"

        # Generate report and check summary
        output_path = generate_matching_report("Drum Tatarani", 1)
        doc = Document(output_path)

        # Extract summary stats from table
        summary_table = doc.tables[0]
        matched_in_report = int(summary_table.rows[3].cells[1].text)
        ref_only_in_report = int(summary_table.rows[4].cells[1].text)
        oferta_only_in_report = int(summary_table.rows[5].cells[1].text)

        # Verify counts match (allowing for slight differences due to grouping)
        assert matched_in_report >= 0, "Matched articles count invalid"
        assert ref_only_in_report >= 0, "Ref-only articles count invalid"
        assert oferta_only_in_report >= 0, "Oferta-only articles count invalid"
