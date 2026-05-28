# tests/test_di_to_docx.py
import sys
import pytest
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from di_to_docx import _render_metadata, _render_f3_table, _render_pages


def _make_metadata_table():
    return {
        "row_count": 4, "column_count": 2,
        "cells": [
            {"row_index": 0, "column_index": 0, "kind": "", "content": "Beneficiar:"},
            {"row_index": 0, "column_index": 1, "kind": "", "content": "Primaria Racari"},
            {"row_index": 1, "column_index": 0, "kind": "", "content": ""},
            {"row_index": 1, "column_index": 1, "kind": "", "content": ""},
            {"row_index": 2, "column_index": 0, "kind": "", "content": "Obiectivul:"},
            {"row_index": 2, "column_index": 1, "kind": "", "content": "EFICIENTIZARE ENERGETICA"},
            {"row_index": 3, "column_index": 0, "kind": "", "content": "Obiectul:"},
            {"row_index": 3, "column_index": 1, "kind": "", "content": "1 LUCRARI BLOC A"},
        ],
    }


def _make_f3_table(row_count=4, col_count=6):
    cells = []
    r0 = ["SECTIUNEA TEHNICA", "", "", "SECTIUNEA FINANCIARA", "", ""]
    r1 = ["Nr.", "Capitol de lucrari", "U.M.", "Cantitatea", "Pret unitar", "TOTAL"]
    for c, v in enumerate(r0):
        cells.append({"row_index": 0, "column_index": c, "kind": "", "content": v})
    for c, v in enumerate(r1):
        cells.append({"row_index": 1, "column_index": c, "kind": "columnHeader", "content": v})
    for r in range(2, row_count):
        for c in range(col_count):
            cells.append({"row_index": r, "column_index": c, "kind": "", "content": f"v{r}{c}"})
    return {"row_count": row_count, "column_count": col_count, "cells": cells}


def test_metadata_renders_key_value_paragraphs():
    doc = Document()
    _render_metadata(doc, _make_metadata_table())
    texts = [p.text for p in doc.paragraphs]
    assert any("Beneficiar:" in t and "Primaria Racari" in t for t in texts)
    assert any("Obiectivul:" in t and "EFICIENTIZARE ENERGETICA" in t for t in texts)


def test_metadata_skips_empty_key_rows():
    doc = Document()
    _render_metadata(doc, _make_metadata_table())
    texts = [p.text for p in doc.paragraphs]
    # Rows with empty key (row 1) must not produce output paragraphs
    assert len([t for t in texts if t.strip()]) == 3  # Beneficiar, Obiectivul, Obiectul


def test_f3_table_row_col_count():
    doc = Document()
    _render_f3_table(doc, _make_f3_table(row_count=5, col_count=6))
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 5
    assert len(doc.tables[0].columns) == 6


def test_f3_table_header_content():
    doc = Document()
    _render_f3_table(doc, _make_f3_table())
    tbl = doc.tables[0]
    assert tbl.rows[0].cells[0].text == "SECTIUNEA TEHNICA"
    assert tbl.rows[1].cells[1].text == "Capitol de lucrari"


def test_page_lines_rendered_with_separators():
    doc = Document()
    pages = [
        {"page_number": 1, "lines": [{"content": "Linie A"}, {"content": "Linie B"}]},
        {"page_number": 2, "lines": [{"content": "Linie C"}]},
    ]
    _render_pages(doc, pages)
    texts = [p.text for p in doc.paragraphs]
    assert "--- Pagina 1 ---" in texts
    assert "--- Pagina 2 ---" in texts
    assert "Linie A" in texts
    assert "Linie C" in texts
