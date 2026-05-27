from docx import Document
from shared.report_word import _count_main_articles, _add_group_totals_row, GRAY_FILL


def test_count_main_articles_empty():
    assert _count_main_articles([]) == 0


def test_count_main_articles_all_main():
    articles = [
        {"cod": "A01", "is_component": False, "cantitate": 1.0},
        {"cod": "A02", "cantitate": 2.5},  # missing is_component → not a component
        {"cod": "A03", "is_component": False, "cantitate": 0.5},
    ]
    assert _count_main_articles(articles) == 3


def test_count_main_articles_filters_components():
    articles = [
        {"cod": "A01", "is_component": False, "cantitate": 1.0},
        {"cod": "A01-sub1", "is_component": True, "cantitate": 1.0},
        {"cod": "A01-sub2", "is_component": True, "cantitate": 1.0},
        {"cod": "A02", "is_component": False, "cantitate": 1.0},
    ]
    assert _count_main_articles(articles) == 2


def test_count_main_articles_filters_zero_cantitate():
    articles = [
        {"cod": "A01", "is_component": False, "cantitate": 1.0},
        {"cod": "A02", "is_component": False, "cantitate": 0.0},  # cant=0 → excluded
        {"cod": "A03", "is_component": False},  # no cantitate → excluded
    ]
    assert _count_main_articles(articles) == 1


def test_count_main_articles_all_components():
    articles = [
        {"cod": "X01", "is_component": True, "cantitate": 1.0},
        {"cod": "X02", "is_component": True, "cantitate": 1.0},
    ]
    assert _count_main_articles(articles) == 0


# ── Tests for _add_group_totals_row ────────────────────────────────────────


def _make_table(doc):
    """11-column table with one data row."""
    table = doc.add_table(rows=1, cols=11)
    return table


def _get_shading_fill(cell):
    """Extract fill hex from cell shading XML, or None."""
    tc = cell._tc
    tcPr = tc.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr"
    )
    if tcPr is None:
        return None
    shd = tcPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
    )
    if shd is None:
        return None
    return shd.get(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"
    )


def test_add_group_totals_row_matched_adds_one_row():
    doc = Document()
    table = _make_table(doc)
    initial_count = len(table.rows)
    _add_group_totals_row(table, ref_count=10, oferta_count=8)
    assert len(table.rows) == initial_count + 1
    # All cells must be gray-shaded
    row = table.rows[-1]
    for cell in [row.cells[0], row.cells[2], row.cells[6], row.cells[10]]:
        assert _get_shading_fill(cell).upper() == GRAY_FILL.upper()


def test_add_group_totals_row_matched_text():
    doc = Document()
    table = _make_table(doc)
    _add_group_totals_row(table, ref_count=10, oferta_count=8)
    row = table.rows[-1]
    # After merge, cell 0 = label, cell 2 = ref text, cell 6 = offer text
    label_text = row.cells[0].text
    ref_text = row.cells[2].text
    offer_text = row.cells[6].text
    assert "TOTAL" in label_text
    assert "10" in ref_text
    assert "8" in offer_text


def test_add_group_totals_row_ref_only_no_offer_text():
    doc = Document()
    table = _make_table(doc)
    _add_group_totals_row(table, ref_count=5, oferta_count=None)
    row = table.rows[-1]
    ref_text = row.cells[2].text
    offer_text = row.cells[6].text
    assert "5" in ref_text
    assert offer_text.strip() == ""


def test_add_group_totals_row_oferta_only_no_ref_text():
    doc = Document()
    table = _make_table(doc)
    _add_group_totals_row(table, ref_count=None, oferta_count=7)
    row = table.rows[-1]
    ref_text = row.cells[2].text
    offer_text = row.cells[6].text
    assert ref_text.strip() == ""
    assert "7" in offer_text


# ── Tests for _generate_word_holistic with group totals ────────────────────


from shared.report_word import _generate_word_holistic


def _make_holistic_comp():
    """Minimal comp dict for _generate_word_holistic."""
    return {"ofertant": "Test SRL", "source_file": "test.pdf"}


def _make_matched_group(n_ref=3, n_oferta=2, n_neconformitati=0):
    return {
        "ref_deviz_cod": "REF01",
        "oferta_deviz_cod": "OFF01",
        "ref_header": None,
        "oferta_header": None,
        "deviz_denumire": "Test deviz",
        "ref_articles": [{"cod": f"R{i}", "is_component": False, "cantitate": 1.0} for i in range(n_ref)],
        "oferta_articles": [{"cod": f"O{i}", "is_component": False, "cantitate": 1.0} for i in range(n_oferta)],
        "neconformitati": [],
        "matches": n_oferta,
    }


def _make_ref_only_group(n_articles=4):
    return {
        "ref_deviz_cod": "REF02",
        "ref_header": None,
        "deviz_denumire": "Ref only deviz",
        "articles": [{"cod": f"R{i}", "is_component": False, "cantitate": 1.0} for i in range(n_articles)],
        "neconformitati": [],
    }


def _make_oferta_only_group(n_articles=5):
    return {
        "oferta_deviz_cod": "OFF02",
        "oferta_header": None,
        "deviz_denumire": "Oferta only deviz",
        "articles": [{"cod": f"O{i}", "is_component": False, "cantitate": 1.0} for i in range(n_articles)],
        "neconformitati": [],
    }


def _all_row_texts(table):
    texts = []
    for row in table.rows:
        row_text = " | ".join(c.text for c in row.cells)
        texts.append(row_text)
    return texts


def test_holistic_matched_group_has_totals_row():
    doc = Document()
    raport = {
        "matched_groups": [_make_matched_group(n_ref=3, n_oferta=2)],
        "ref_only_groups": [],
        "oferta_only_groups": [],
        "ungrouped": [],
        "unassigned_articles": [],
    }
    _generate_word_holistic(doc, raport, _make_holistic_comp())
    all_texts = _all_row_texts(doc.tables[0])
    assert any("TOTAL GRUP" in t for t in all_texts), "Missing TOTAL GRUP row for matched group"
    assert any("3" in t and "Referință" in t for t in all_texts), "Missing ref count"
    assert any("2" in t and "Ofertă" in t for t in all_texts), "Missing offer count"


def test_holistic_ref_only_group_has_totals_row_ref_side_only():
    doc = Document()
    raport = {
        "matched_groups": [],
        "ref_only_groups": [_make_ref_only_group(n_articles=4)],
        "oferta_only_groups": [],
        "ungrouped": [],
        "unassigned_articles": [],
    }
    _generate_word_holistic(doc, raport, _make_holistic_comp())
    all_texts = _all_row_texts(doc.tables[0])
    assert any("TOTAL GRUP" in t for t in all_texts)
    assert any("4" in t and "Referință" in t for t in all_texts)
    # Offer side must be empty
    total_rows = [t for t in all_texts if "TOTAL GRUP" in t]
    assert all("Ofertă" not in t for t in total_rows)


def test_holistic_oferta_only_group_has_totals_row_offer_side_only():
    doc = Document()
    raport = {
        "matched_groups": [],
        "ref_only_groups": [],
        "oferta_only_groups": [_make_oferta_only_group(n_articles=5)],
        "ungrouped": [],
        "unassigned_articles": [],
    }
    _generate_word_holistic(doc, raport, _make_holistic_comp())
    all_texts = _all_row_texts(doc.tables[0])
    assert any("TOTAL GRUP" in t for t in all_texts)
    assert any("5" in t and "Ofertă" in t for t in all_texts)
    total_rows = [t for t in all_texts if "TOTAL GRUP" in t]
    assert all("Referință" not in t for t in total_rows)
