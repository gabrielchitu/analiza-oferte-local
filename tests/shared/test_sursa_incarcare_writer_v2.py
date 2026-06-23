import os
import tempfile
import pytest
from docx import Document
from docx.oxml.ns import qn
from shared.sursa_incarcare_writer import write_docx_v2


def _make_deviz(n_main=2, n_sub=1):
    """Minimal deviz for testing."""
    articles = []
    for i in range(1, n_main + 1):
        art = {
            "nr_crt": str(i),
            "cod": f"TST0{i}A",
            "denumire": f"Articol test {i}",
            "um": "mc",
            "cantitate": float(i),
            "pret_unitar": 0.0,
            "total": 0.0,
            "breakdown": {
                "material": {"pret": 10.0 * i, "total": 0.0},
                "manopera": {"pret": 0.0, "total": 0.0},
                "utilaj": {"pret": 0.0, "total": 0.0},
                "transport": {"pret": 0.0, "total": 0.0},
            },
            "sub_items": [],
        }
        if i == 1 and n_sub > 0:
            art["sub_items"] = [{
                "nr_crt": "1.1",
                "cod": "SPEC_SUB",
                "denumire": "Subcomponent test",
                "um": "buc",
                "cantitate": 5.0,
            }]
        articles.append(art)
    return {
        "obiectivul": "OBIECTIV TEST",
        "obiectul": "Ob 001",
        "categoria": "Cat Test",
        "capitole": [{"titlu": "Capitol 1", "articole": articles}],
    }


def _write_and_open(devize, metadata=None):
    out = tempfile.mktemp(suffix=".docx")
    write_docx_v2(devize, out, metadata=metadata)
    return out, Document(out)


def test_write_docx_v2_creates_file():
    out, doc = _write_and_open([_make_deviz()])
    try:
        assert os.path.exists(out)
        assert os.path.getsize(out) > 3000
    finally:
        os.unlink(out)


def test_write_docx_v2_document_header_paragraphs():
    """Primele 4 paragrafe: Lista articole, Client, Ofertant, Generat."""
    meta = {"offer_label": "Oferta 2", "client": "Test Client", "ofertant": "SC TEST SRL", "date": "2026-06-23"}
    out, doc = _write_and_open([_make_deviz()], metadata=meta)
    try:
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        assert paras[0] == "Lista articole — Oferta 2"
        assert "Test Client" in paras[1]
        assert "SC TEST SRL" in paras[2]
        assert "2026-06-23" in paras[3]
    finally:
        os.unlink(out)


def test_write_docx_v2_document_header_bold_title_only():
    """Doar titlul bold; Client/Ofertant/Generat nu sunt bold — identic template."""
    out, doc = _write_and_open([_make_deviz()])
    try:
        header_paras = [p for p in doc.paragraphs if p.text.strip()][:4]
        assert len(header_paras) == 4
        assert any(r.bold for r in header_paras[0].runs), "Titlul trebuie bold"
        for p in header_paras[1:]:
            for run in p.runs:
                assert not run.bold, f"Linia nu trebuie bold: {p.text!r}"
    finally:
        os.unlink(out)


def test_write_docx_v2_group_header_pipe_format():
    """Paragraf grup: 'OBIECTIV TEST | Ob 001 | Cat Test'."""
    out, doc = _write_and_open([_make_deviz()])
    try:
        group_paras = [p for p in doc.paragraphs if p.text.strip()][4:]
        assert group_paras[0].text == "OBIECTIV TEST | Ob 001 | Cat Test"
    finally:
        os.unlink(out)


def test_write_docx_v2_table_12_cols():
    """Fiecare grup are exact un tabel cu 12 coloane (adăugat Total)."""
    out, doc = _write_and_open([_make_deviz()])
    try:
        assert len(doc.tables) == 1
        tblGrid = doc.tables[0]._tbl.find(qn('w:tblGrid'))
        assert tblGrid is not None
        cols = tblGrid.findall(qn('w:gridCol'))
        assert len(cols) == 12
    finally:
        os.unlink(out)


def test_write_docx_v2_col_widths_exact():
    from shared.sursa_incarcare_writer import _COL_WIDTHS_TWIPS_V2
    out, doc = _write_and_open([_make_deviz()])
    try:
        tblGrid = doc.tables[0]._tbl.find(qn('w:tblGrid'))
        widths = [int(gc.get(qn('w:w'))) for gc in tblGrid.findall(qn('w:gridCol'))]
        assert widths == _COL_WIDTHS_TWIPS_V2
    finally:
        os.unlink(out)


def test_write_docx_v2_table_grid_style():
    """Tabel foloseste TableGrid style — borduri vizibile ca in template."""
    out, doc = _write_and_open([_make_deviz()])
    try:
        tbl = doc.tables[0]
        tblPr = tbl._tbl.find(qn('w:tblPr'))
        tblStyle = tblPr.find(qn('w:tblStyle'))
        assert tblStyle is not None, "tblStyle missing"
        assert tblStyle.get(qn('w:val')) == 'TableGrid'
    finally:
        os.unlink(out)


def test_write_docx_v2_table_header_text():
    out, doc = _write_and_open([_make_deviz()])
    try:
        tbl = doc.tables[0]
        assert tbl.rows[0].cells[0].text == "Nr."
        assert tbl.rows[0].cells[2].text == "Cod"
        assert tbl.rows[0].cells[7].text == "Total\n(lei)"   # col 7 = Total
        assert tbl.rows[0].cells[8].text == "Pret unitar (lei/UM)"  # cols 8-11
        assert tbl.rows[1].cells[8].text == "Material"
        assert tbl.rows[1].cells[11].text == "Transport"
    finally:
        os.unlink(out)


def test_write_docx_v2_article_data_rows():
    """Rânduri date: Nr, Nr.crt, Cod, Cod principal, Denumire, UM, Cantitate."""
    out, doc = _write_and_open([_make_deviz(n_main=2, n_sub=0)])
    try:
        tbl = doc.tables[0]
        row2 = tbl.rows[2]  # first data row (after 2 header rows)
        assert row2.cells[0].text == "1"       # Nr local
        assert row2.cells[1].text == "1"       # nr_crt
        assert row2.cells[2].text == "TST01A"  # Cod
        assert row2.cells[3].text == ""         # Cod principal empty for main
        assert row2.cells[4].text == "Articol test 1"
        assert row2.cells[5].text == "mc"
        assert row2.cells[6].text == "1.00"    # cantitate
    finally:
        os.unlink(out)


def test_write_docx_v2_subitem_cod_principal():
    """Sub-item: Cod principal = parent.cod."""
    out, doc = _write_and_open([_make_deviz(n_main=2, n_sub=1)])
    try:
        tbl = doc.tables[0]
        # rows: 0,1=header; 2=main art1; 3=sub_item(1.1); 4=main art2; 5=total
        sub_row = tbl.rows[3]
        assert sub_row.cells[1].text == "1.1"    # nr_crt decimal
        assert sub_row.cells[2].text == "SPEC_SUB"
        assert sub_row.cells[3].text == "TST01A"  # parent.cod
    finally:
        os.unlink(out)


def test_write_docx_v2_breakdown_prices():
    """Col 7=Total, col 8=Material, col 9=Manopera (0→empty)."""
    out, doc = _write_and_open([_make_deviz(n_main=1, n_sub=0)])
    try:
        tbl = doc.tables[0]
        row2 = tbl.rows[2]
        # art 1: cantitate=1, material.pret=10.0 → total=10.0
        assert row2.cells[7].text == "10,00"   # Total
        assert row2.cells[8].text == "10,00"   # Material pret unitar
        assert row2.cells[9].text == ""        # Manopera = 0 → empty
    finally:
        os.unlink(out)


def test_write_docx_v2_total_grup_row():
    """Ultima linie: 'Total grup: N articole principale / M subcomponente'."""
    out, doc = _write_and_open([_make_deviz(n_main=2, n_sub=1)])
    try:
        tbl = doc.tables[0]
        last_row = tbl.rows[-1]
        text = last_row.cells[0].text
        assert "Total grup:" in text
        assert "2 articole principale" in text
        assert "1 subcomponente" in text
    finally:
        os.unlink(out)


def test_write_docx_v2_multiple_devize():
    """Un paragraf + un tabel per deviz."""
    out, doc = _write_and_open([_make_deviz(), _make_deviz(n_main=1, n_sub=0)])
    try:
        assert len(doc.tables) == 2
        group_paras = [p for p in doc.paragraphs if "OBIECTIV TEST" in p.text]
        assert len(group_paras) == 2
    finally:
        os.unlink(out)
