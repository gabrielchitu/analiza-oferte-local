from shared.report_word import _observatie_text, _add_neconf_row, SEMANTIC_MATCH_FILL, SPEC_DIFF_FILL
from docx import Document


# ── Helpers ────────────────────────────────────────────────────────────────

def _get_shading_fill(cell):
    """Extract fill hex from cell shading XML, or None."""
    tc = cell._tc
    tcPr = tc.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr"
    )
    if tcPr is None:
        return None
    shd = tcPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
    )
    if shd is None:
        return None
    return shd.get(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"
    )


def _make_table_with_row(doc):
    """Create an 11-column table; return table."""
    table = doc.add_table(rows=3, cols=11)
    table.style = "Table Grid"
    return table


# ── _observatie_text tests ─────────────────────────────────────────────────

def test_observatie_text_cod_normativ_diferit_with_diferente():
    """NC with diferente list containing camp/ref/oferta entries returns formatted text."""
    nc = {
        "tip": "COD_NORMATIV_DIFERIT",
        "diferente": [
            {"camp": "cod_normativ", "ref": "ES08A4", "oferta": "TCB30A1"}
        ],
        "motiv_llm": "Aceeași lucrare de armare",
    }
    result = _observatie_text(nc)
    assert "cod_normativ" in result
    assert "ES08A4" in result
    assert "TCB30A1" in result


def test_observatie_text_cod_normativ_diferit_empty_diferente():
    """NC with empty diferente returns fallback message."""
    nc = {
        "tip": "COD_NORMATIV_DIFERIT",
        "diferente": [],
    }
    result = _observatie_text(nc)
    assert result == "Cod normativ diferit"


def test_observatie_text_specificatie_diferita_with_nota():
    """NC with nota_specialist returns that exact string."""
    nota = "Specificație diferită: beton C25/30 față de C20/25 din referință"
    nc = {
        "tip": "SPECIFICATIE_DIFERITA",
        "nota_specialist": nota,
    }
    result = _observatie_text(nc)
    assert result == nota


def test_observatie_text_specificatie_diferita_empty_nota():
    """NC with no nota_specialist returns fallback message."""
    nc = {
        "tip": "SPECIFICATIE_DIFERITA",
    }
    result = _observatie_text(nc)
    assert result == "Specificație diferită"


# ── _add_nc_row color and content tests ───────────────────────────────────

def _make_cod_normativ_diferit_nc():
    return {
        "tip": "COD_NORMATIV_DIFERIT",
        "ref_cod": "ES08A4",
        "ref_denumire": "Armare stâlpi",
        "oferta_cod": "TCB30A1",
        "oferta_denumire": "Armare stâlpi conform ofertă",
        "diferente": [{"camp": "cod_normativ", "ref": "ES08A4", "oferta": "TCB30A1"}],
        "motiv_llm": "Aceeași lucrare fizică",
        "deviz_ref": "D01",
        "deviz_denumire": "Structura | Armare",
        "ref_um": "kg",
        "ref_cantitate": 100.0,
        "oferta_um": "kg",
        "oferta_cantitate": 100.0,
    }


def _make_specificatie_diferita_nc():
    return {
        "tip": "SPECIFICATIE_DIFERITA",
        "ref_cod": "BCR30A1",
        "ref_denumire": "Beton C30/37",
        "oferta_cod": "BCR30A1",
        "oferta_denumire": "Beton C25/30",
        "nota_specialist": "Rezistența betonului diferă: C30/37 vs C25/30",
        "deviz_ref": "D01",
        "deviz_denumire": "Structura | Beton",
        "ref_um": "mc",
        "ref_cantitate": 50.0,
        "oferta_um": "mc",
        "oferta_cantitate": 50.0,
    }


def test_add_neconf_row_cod_normativ_diferit_fill_color():
    """COD_NORMATIV_DIFERIT rows must use SEMANTIC_MATCH_FILL background."""
    doc = Document()
    table = _make_table_with_row(doc)
    nc = _make_cod_normativ_diferit_nc()
    _add_neconf_row(table, 1, nc, {})
    row = table.rows[-1]
    fill = _get_shading_fill(row.cells[0])
    assert fill is not None
    assert fill.upper() == SEMANTIC_MATCH_FILL.upper()


def test_add_neconf_row_specificatie_diferita_fill_color():
    """SPECIFICATIE_DIFERITA rows must use SPEC_DIFF_FILL background."""
    doc = Document()
    table = _make_table_with_row(doc)
    nc = _make_specificatie_diferita_nc()
    _add_neconf_row(table, 1, nc, {})
    row = table.rows[-1]
    fill = _get_shading_fill(row.cells[0])
    assert fill is not None
    assert fill.upper() == SPEC_DIFF_FILL.upper()


def test_add_neconf_row_cod_normativ_diferit_col2_text():
    """Col 2 (ref side) must contain REF: prefix and ref_cod."""
    doc = Document()
    table = _make_table_with_row(doc)
    nc = _make_cod_normativ_diferit_nc()
    _add_neconf_row(table, 1, nc, {})
    row = table.rows[-1]
    col2_text = row.cells[2].text
    assert "REF:" in col2_text
    assert "ES08A4" in col2_text


def test_add_neconf_row_cod_normativ_diferit_col3_text():
    """Col 3 (oferta side) must contain OFF: prefix and oferta_cod."""
    doc = Document()
    table = _make_table_with_row(doc)
    nc = _make_cod_normativ_diferit_nc()
    _add_neconf_row(table, 1, nc, {})
    row = table.rows[-1]
    col3_text = row.cells[3].text
    assert "OFF:" in col3_text
    assert "TCB30A1" in col3_text


def test_add_neconf_row_specificatie_diferita_col2_text():
    """Col 2 for SPECIFICATIE_DIFERITA must contain REF: and ref_denumire."""
    doc = Document()
    table = _make_table_with_row(doc)
    nc = _make_specificatie_diferita_nc()
    _add_neconf_row(table, 1, nc, {})
    row = table.rows[-1]
    col2_text = row.cells[2].text
    assert "REF:" in col2_text
    assert "Beton C30/37" in col2_text


def test_add_neconf_row_specificatie_diferita_col3_text():
    """Col 3 for SPECIFICATIE_DIFERITA must contain OFF: and oferta_denumire."""
    doc = Document()
    table = _make_table_with_row(doc)
    nc = _make_specificatie_diferita_nc()
    _add_neconf_row(table, 1, nc, {})
    row = table.rows[-1]
    col3_text = row.cells[3].text
    assert "OFF:" in col3_text
    assert "Beton C25/30" in col3_text
