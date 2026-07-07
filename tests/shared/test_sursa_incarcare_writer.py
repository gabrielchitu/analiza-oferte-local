import subprocess
import pytest
from pathlib import Path
from docx import Document
from openpyxl import load_workbook
from shared.sursa_incarcare_writer import make_acronym, write_docx, write_xlsx, write_pdf


def test_make_acronym_standard():
    assert make_acronym("CONSTRUIRE UNITATE DE CAZARE - TARGOVISTE") == "CUCT"


def test_make_acronym_with_numeric_prefix():
    assert make_acronym("0232 000000232 DRUMURI TATARANI") == "DT"


def test_make_acronym_short():
    assert make_acronym("DRUMURI TATARANI") == "DT"


def test_make_acronym_max_6():
    result = make_acronym("CONSTRUIRE REABILITARE MODERNIZARE EXTINDERE CONSOLIDARE RENOVARE")
    assert len(result) <= 6


def _make_sample_deviz(status="OK"):
    return {
        'status': status,
        'deviz_key': 'test',
        'obiectivul': 'CONSTRUIRE UNITATE DE CAZARE - TARGOVISTE',
        'obiectul': '3 ARHITECTURA',
        'categoria': '3.1 ARHITECTURA',
        'total_deviz': 24220.05,
        'capitole': [
            {
                'titlu': 'INFRASTRUCTURA',
                'total_capitol': 24220.05,
                'articole': [
                    {
                        'nr_crt': '1',
                        'cod': 'CF38A*',
                        'denumire': 'Tencuiala pe baza de ciment',
                        'um': 'MP',
                        'cantitate': 225.0,
                        'pret_unitar': 33.22,
                        'total': 7473.71,
                        'suspect': False,
                        'breakdown': {
                            'material': {'pret': 13.22, 'total': 2973.71},
                            'manopera': {'pret': 20.00, 'total': 4500.00},
                            'utilaj': {'pret': 0.00, 'total': 0.00},
                            'transport': {'pret': 0.00, 'total': 0.00},
                            'control_ok': True,
                        },
                        'sub_items': [],
                    },
                    {
                        'nr_crt': '2',
                        'cod': 'RPCE27A+',
                        'denumire': 'Mastic bituminos',
                        'um': 'MP',
                        'cantitate': 307.0,
                        'pret_unitar': 54.55,
                        'total': 16746.33,
                        'suspect': False,
                        'breakdown': None,
                        'sub_items': [],
                    },
                ],
            }
        ],
    }


def test_write_docx_creates_file(tmp_path):
    deviz = _make_sample_deviz()
    out = tmp_path / "test.docx"
    write_docx([deviz], out)
    assert out.exists()
    assert out.stat().st_size > 0


def test_write_docx_has_rows(tmp_path):
    deviz = _make_sample_deviz()
    out = tmp_path / "test.docx"
    write_docx([deviz], out)
    doc = Document(str(out))
    assert len(doc.tables) >= 1
    tbl = doc.tables[0]
    # header(2) + capitol(1) + art1(1) + breakdown(4) + art2(1) + total_cap(1) + total_deviz(1) = 11
    assert len(tbl.rows) >= 8


def test_write_docx_no_breakdown_rows_when_none(tmp_path):
    deviz = _make_sample_deviz()
    deviz['capitole'][0]['articole'][0]['breakdown'] = None
    out = tmp_path / "test.docx"
    write_docx([deviz], out)
    doc = Document(str(out))
    tbl = doc.tables[0]
    # Without breakdown: header(2) + capitol(1) + art1(1) + art2(1) + total_cap(1) + total_deviz(1) = 7
    assert len(tbl.rows) == 7


def test_write_docx_red_flag_when_red(tmp_path):
    deviz = _make_sample_deviz(status="RED")
    out = tmp_path / "test.docx"
    write_docx([deviz], out)
    doc = Document(str(out))
    all_text = ' '.join(
        cell.text for tbl in doc.tables for row in tbl.rows for cell in row.cells
    )
    assert 'NECONFIRMAT' in all_text


def test_write_xlsx_creates_file(tmp_path):
    deviz = _make_sample_deviz()
    out = tmp_path / "test.xlsx"
    write_xlsx([deviz], out)
    assert out.exists()
    assert out.stat().st_size > 0


def test_write_xlsx_sheet_name(tmp_path):
    deviz = _make_sample_deviz()
    out = tmp_path / "test.xlsx"
    write_xlsx([deviz], out)
    wb = load_workbook(str(out))
    assert '3.1 ARHITECTURA' in wb.sheetnames[0]


def test_write_xlsx_has_rows(tmp_path):
    deviz = _make_sample_deviz()
    out = tmp_path / "test.xlsx"
    write_xlsx([deviz], out)
    wb = load_workbook(str(out))
    ws = wb.worksheets[0]  # first deviz sheet, not the Recapitulatie sheet
    assert ws.max_row >= 8


def test_write_xlsx_red_flag(tmp_path):
    deviz = _make_sample_deviz(status="RED")
    out = tmp_path / "test.xlsx"
    write_xlsx([deviz], out)
    wb = load_workbook(str(out))
    ws = wb.active
    all_text = ' '.join(
        str(ws.cell(r, c).value or '')
        for r in range(1, ws.max_row + 1)
        for c in range(1, 7)
    )
    assert 'NECONFIRMAT' in all_text
    assert 'necesară' in all_text


def test_write_pdf_skips_gracefully_if_no_libreoffice(tmp_path, monkeypatch):
    deviz = _make_sample_deviz()
    docx_path = tmp_path / "test.docx"
    write_docx([deviz], docx_path)

    def fake_run(*args, **kwargs):
        raise FileNotFoundError("soffice not found")
    monkeypatch.setattr(subprocess, "run", fake_run)

    result = write_pdf(docx_path, tmp_path)
    assert result is False
