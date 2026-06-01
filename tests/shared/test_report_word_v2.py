"""Tests for Word report generation from v2 holistic JSON.

Validates DOCX report generation, structure, and content accuracy.
"""

import pytest
from pathlib import Path
from docx import Document
from typing import Dict

from shared.report_word_v2 import generate_holistic_report


# Fixtures

@pytest.fixture
def mock_holistic_json() -> Dict:
    """Create mock holistic JSON structure."""
    return {
        "client": "TestClient",
        "di_file": "oferta_1.json",
        "extraction_version": "2.0",
        "matched_groups": [
            {
                "deviz_cod": "0001",
                "deviz_den": "GROUP 1 | Category A | SubCat A",
                "articole": [
                    {
                        "cod": "CA01A",
                        "um": "BUC",
                        "cant": 5,
                        "denumire": "Component 1",
                    },
                    {
                        "cod": "CA02B",
                        "um": "M",
                        "cant": 10,
                        "denumire": "Component 2",
                    },
                ],
            },
            {
                "deviz_cod": "0002",
                "deviz_den": "GROUP 2 | Category B | SubCat B",
                "articole": [
                    {
                        "cod": "CB01A",
                        "um": "KG",
                        "cant": 25,
                        "denumire": "Material A",
                    },
                ],
            },
        ],
        "ref_only_groups": [
            {
                "deviz_cod": "0003",
                "deviz_den": "GROUP 3 | Category C | SubCat C",
                "articole": [
                    {
                        "cod": "CC01A",
                        "um": "TON",
                        "cant": 50,
                        "denumire": "Material Not in Offer",
                    },
                ],
            },
        ],
        "oferta_only_groups": [
            {
                "deviz_cod": "0004",
                "deviz_den": "GROUP 4 | Category D | SubCat D",
                "articole": [
                    {
                        "cod": "CD01A",
                        "um": "BUC",
                        "cant": 3,
                        "denumire": "Extra Component",
                    },
                ],
            },
        ],
        "stats": {
            "matched_articles_count": 3,
            "ref_only_articles_count": 1,
            "oferta_only_articles_count": 1,
            "matched_groups_count": 2,
            "ref_only_groups_count": 1,
            "oferta_only_groups_count": 1,
        },
    }


# Tests

def test_generate_holistic_report_creates_file(tmp_path, mock_holistic_json):
    """Test report generation creates output file."""
    output_path = tmp_path / "test_report.docx"

    generate_holistic_report(
        holistic=mock_holistic_json,
        output_path=output_path,
        client_name="TestClient",
    )

    assert output_path.exists()
    assert output_path.stat().st_size > 0


def test_generate_holistic_report_valid_docx(tmp_path, mock_holistic_json):
    """Test report file is valid DOCX (can be opened by python-docx)."""
    output_path = tmp_path / "test_report.docx"

    generate_holistic_report(
        holistic=mock_holistic_json,
        output_path=output_path,
    )

    # Try to open as DOCX
    doc = Document(str(output_path))
    assert doc is not None
    assert len(doc.paragraphs) > 0


def test_generate_holistic_report_title_present(tmp_path, mock_holistic_json):
    """Test report contains title heading."""
    output_path = tmp_path / "test_report.docx"

    generate_holistic_report(
        holistic=mock_holistic_json,
        output_path=output_path,
    )

    doc = Document(str(output_path))
    text = "\n".join([p.text for p in doc.paragraphs])

    assert "Raport Matching" in text or "Summary" in text


def test_generate_holistic_report_metadata_present(tmp_path, mock_holistic_json):
    """Test report contains metadata (client, date, etc)."""
    output_path = tmp_path / "test_report.docx"

    generate_holistic_report(
        holistic=mock_holistic_json,
        output_path=output_path,
        client_name="TestClient",
    )

    doc = Document(str(output_path))
    text = "\n".join([p.text for p in doc.paragraphs])

    assert "TestClient" in text
    assert "Extraction Version" in text


def test_generate_holistic_report_summary_table(tmp_path, mock_holistic_json):
    """Test report contains summary table with correct metrics."""
    output_path = tmp_path / "test_report.docx"

    generate_holistic_report(
        holistic=mock_holistic_json,
        output_path=output_path,
    )

    doc = Document(str(output_path))

    # Check for tables
    assert len(doc.tables) >= 1

    # Summary table should have metrics
    summary_table = doc.tables[0]
    table_text = "\n".join(
        [
            cell.text
            for row in summary_table.rows
            for cell in row.cells
        ]
    )

    # Verify key metrics are present
    assert "Matched Articles" in table_text
    assert "Ref-Only" in table_text or "ref" in table_text.lower()
    assert "Oferta-Only" in table_text or "oferta" in table_text.lower()


def test_generate_holistic_report_groups_present(tmp_path, mock_holistic_json):
    """Test report contains group sections (matched, ref-only, oferta-only)."""
    output_path = tmp_path / "test_report.docx"

    generate_holistic_report(
        holistic=mock_holistic_json,
        output_path=output_path,
    )

    doc = Document(str(output_path))
    text = "\n".join([p.text for p in doc.paragraphs])

    # Should have section headings for groups
    assert "Matched Groups" in text
    assert "Reference-Only" in text or "Ref-Only" in text
    assert "Offer-Only" in text or "Oferta-Only" in text


def test_generate_holistic_report_articles_listed(tmp_path, mock_holistic_json):
    """Test report contains article details."""
    output_path = tmp_path / "test_report.docx"

    generate_holistic_report(
        holistic=mock_holistic_json,
        output_path=output_path,
    )

    doc = Document(str(output_path))

    # Extract all text including from tables
    text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    # Should contain article codes from matched group
    assert "CA01A" in text
    assert "CA02B" in text
    # Should contain article from ref-only group
    assert "CC01A" in text


def test_generate_holistic_report_empty_groups(tmp_path):
    """Test report generation with empty groups."""
    holistic = {
        "client": "TestClient",
        "di_file": "oferta_1.json",
        "extraction_version": "2.0",
        "matched_groups": [],
        "ref_only_groups": [],
        "oferta_only_groups": [],
        "stats": {
            "matched_articles_count": 0,
            "ref_only_articles_count": 0,
            "oferta_only_articles_count": 0,
            "matched_groups_count": 0,
            "ref_only_groups_count": 0,
            "oferta_only_groups_count": 0,
        },
    }

    output_path = tmp_path / "empty_report.docx"

    generate_holistic_report(
        holistic=holistic,
        output_path=output_path,
    )

    assert output_path.exists()

    doc = Document(str(output_path))
    assert len(doc.paragraphs) > 0
