import pytest
from docx import Document
from shared.lista_oferta_writer import extract_entity_name, _iter_source_groups, _fmt_nr_crt, _fmt_price, _build_table_header, _write_article_row, _write_group_section, build_docx_for_source
import json
import tempfile
import os


def _make_di(pages_lines):
    """Helper: write temp di_*.json and return path."""
    data = {
        "pages": [
            {"page_number": i + 1, "lines": lines}
            for i, lines in enumerate(pages_lines)
        ],
        "tables": [],
    }
    f = tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False)
    json.dump(data, f)
    f.close()
    return f.name


def test_extract_ofertant_found():
    path = _make_di(
        [
            [
                {"content": "Formularul F3"},
                {"content": "CONTRACTANT (OFERTANT)"},
                {"content": "SC. KATO SERVICE SRL"},
            ]
        ]
    )
    try:
        assert extract_entity_name(path, is_referinta=False) == "SC. KATO SERVICE SRL"
    finally:
        os.unlink(path)


def test_extract_ofertant_skips_bare_srl():
    path = _make_di(
        [
            [
                {"content": "CONTRACTANT (OFERTANT)"},
                {"content": "SRL"},
                {"content": "SC. REAL COMPANY SRL"},
            ]
        ]
    )
    try:
        assert extract_entity_name(path, is_referinta=False) == "SC. REAL COMPANY SRL"
    finally:
        os.unlink(path)


def test_extract_proiectant_found():
    path = _make_di(
        [
            [
                {"content": "PROIECTANT"},
                {"content": "SC. ARHI DESIGN SRL"},
            ]
        ]
    )
    try:
        assert extract_entity_name(path, is_referinta=True) == "SC. ARHI DESIGN SRL"
    finally:
        os.unlink(path)


def test_extract_fallback_when_not_found():
    path = _make_di([[{"content": "Random text"}]])
    try:
        assert extract_entity_name(path, is_referinta=False) == "Necunoscut"
    finally:
        os.unlink(path)


def test_extract_searches_only_first_5_pages():
    # Marker on page 6 — should not be found
    pages = [[{"content": "nothing"}]] * 5 + [
        [
            {"content": "CONTRACTANT (OFERTANT)"},
            {"content": "SC. LATE SRL"},
        ]
    ]
    path = _make_di(pages)
    try:
        assert extract_entity_name(path, is_referinta=False) == "Necunoscut"
    finally:
        os.unlink(path)


# Tests for _iter_source_groups
def _make_article(cod="TST01", nr_ordine=1, is_component=False, parent_code=None, deviz_header=None):
    return {
        "cod": cod, "denumire": "Test article", "um": "mc", "cantitate": 1.0,
        "nr_ordine": nr_ordine, "is_component": is_component, "parent_code": parent_code,
        "pret_material": 0.0, "pret_manopera": 0.0, "pret_utilaj": 0.0, "pret_transport": 0.0,
        "val_material": 0.0, "val_manopera": 0.0, "val_utilaj": 0.0, "val_transport": 0.0,
        "deviz_header": deviz_header or {"obiectivul": "OBJ", "obiectul": "OBL", "categoria": "CAT"},
    }


def test_iter_source_groups_oferta():
    art = _make_article("TST01")
    holistic = {
        "matched_groups": [{"oferta_articles": [art], "ref_articles": [], "deviz_denumire": "A|B|C"}],
        "oferta_only_groups": [],
        "ref_only_groups": [],
    }
    groups = list(_iter_source_groups(holistic, source="oferta"))
    assert len(groups) == 1
    header, articles, _ = groups[0]
    assert header["obiectivul"] == "OBJ"
    assert articles[0]["cod"] == "TST01"


def test_iter_source_groups_referinta():
    art = _make_article("REF01", deviz_header={"obiectivul": "R", "obiectul": "RL", "categoria": "RC"})
    holistic = {
        "matched_groups": [{"ref_articles": [art], "oferta_articles": [], "deviz_denumire": "R|RL|RC"}],
        "ref_only_groups": [],
        "oferta_only_groups": [],
    }
    groups = list(_iter_source_groups(holistic, source="referinta"))
    assert len(groups) == 1
    header, articles, _ = groups[0]
    assert header["obiectivul"] == "R"
    assert articles[0]["cod"] == "REF01"


def test_iter_source_groups_includes_only_groups():
    art = _make_article("ONLY01", deviz_header={"obiectivul": "X", "obiectul": "Y", "categoria": "Z"})
    holistic = {
        "matched_groups": [],
        "oferta_only_groups": [{"articles": [art], "deviz_denumire": "X|Y|Z"}],
        "ref_only_groups": [],
    }
    groups = list(_iter_source_groups(holistic, source="oferta"))
    assert len(groups) == 1
    _, articles, _ = groups[0]
    assert articles[0]["cod"] == "ONLY01"


def test_iter_source_groups_skips_empty_article_groups():
    holistic = {
        "matched_groups": [{"oferta_articles": [], "ref_articles": [], "deviz_denumire": "A|B|C"}],
        "oferta_only_groups": [],
        "ref_only_groups": [],
    }
    groups = list(_iter_source_groups(holistic, source="oferta"))
    assert len(groups) == 0


def test_iter_source_groups_fallback_header_from_deviz_denumire():
    """Extract header from ref_only_groups articles via deviz_header field."""
    art = {"cod": "X", "denumire": "D", "um": "mc", "cantitate": 1.0,
           "nr_ordine": 1, "is_component": False, "parent_code": None,
           "pret_material": 0.0, "pret_manopera": 0.0, "pret_utilaj": 0.0, "pret_transport": 0.0,
           "val_material": 0.0, "val_manopera": 0.0, "val_utilaj": 0.0, "val_transport": 0.0,
           "deviz_header": {"obiectivul": "OBJ2", "obiectul": "OBL2", "categoria": "CAT2"}}
    holistic = {
        "matched_groups": [],
        "ref_only_groups": [{"articles": [art], "deviz_denumire": "OBJ2|OBL2|CAT2"}],
        "oferta_only_groups": [],
    }
    groups = list(_iter_source_groups(holistic, source="referinta"))
    assert len(groups) == 1
    header, _, _ = groups[0]
    assert header["obiectivul"] == "OBJ2"


# Tests for _fmt_nr_crt and _fmt_price
def test_fmt_nr_crt_integer():
    assert _fmt_nr_crt(1) == "1"


def test_fmt_nr_crt_string_subcomp():
    assert _fmt_nr_crt("9.1") == "9.1"


def test_fmt_nr_crt_float_becomes_int():
    # nr_ordine sometimes stored as float 1.0
    assert _fmt_nr_crt(1.0) == "1"


def test_fmt_price_zero_returns_empty():
    assert _fmt_price(0.0) == ""


def test_fmt_price_nonzero_returns_formatted():
    assert _fmt_price(1234.5) == "1.234,50"


def test_fmt_price_rounds_to_2_decimals():
    assert _fmt_price(9.999) == "10,00"


def test_build_table_header_structure():
    doc = Document()
    tbl = doc.add_table(rows=2, cols=11)
    _build_table_header(tbl)
    # Row 0 has content in first 7 cells + merged spans for Pret and Val
    assert tbl.rows[0].cells[0].text == "Nr."
    assert tbl.rows[0].cells[1].text == "Nr.crt"
    assert tbl.rows[0].cells[2].text == "Cod"
    assert tbl.rows[0].cells[3].text == "Cod principal"
    assert tbl.rows[0].cells[4].text == "Denumire"
    assert tbl.rows[0].cells[5].text == "UM"
    assert tbl.rows[0].cells[6].text == "Cantitate"
    # Row 0 merged price group header
    assert tbl.rows[0].cells[7].text == "Pret unitar (lei/UM)"
    # Row 1 price sub-headers (no Valoare section)
    assert tbl.rows[1].cells[7].text == "Material"
    assert tbl.rows[1].cells[8].text == "Manoperă"
    assert tbl.rows[1].cells[9].text == "Utilaje"
    assert tbl.rows[1].cells[10].text == "Transport"


# Tests for _write_article_row
def _make_full_article(cod="TST01", nr_ordine=1, is_component=False, parent_code=None,
                       pret_material=0.0, val_material=0.0, pret_manopera=0.0, val_manopera=0.0,
                       pret_utilaj=0.0, val_utilaj=0.0, pret_transport=0.0, val_transport=0.0):
    return {
        "cod": cod, "denumire": "Test article", "um": "mc", "cantitate": 2.5,
        "nr_ordine": nr_ordine, "is_component": is_component, "parent_code": parent_code,
        "pret_material": pret_material, "val_material": val_material,
        "pret_manopera": pret_manopera, "val_manopera": val_manopera,
        "pret_utilaj": pret_utilaj, "val_utilaj": val_utilaj,
        "pret_transport": pret_transport, "val_transport": val_transport,
        "deviz_header": {"obiectivul": "", "obiectul": "", "categoria": ""},
    }


def test_write_article_row_principal():
    doc = Document()
    tbl = doc.add_table(rows=0, cols=11)
    row = tbl.add_row()
    art = _make_full_article(cod="TSD06XA", nr_ordine=3)
    _write_article_row(row, seq_nr=3, article=art)
    cells = row.cells
    assert cells[0].text == "3"        # Nr sequential
    assert cells[1].text == "3"        # Nr.crt
    assert cells[2].text == "TSD06XA"  # Cod
    assert cells[3].text == ""          # Cod principal (empty for principal)
    assert cells[4].text == "Test article"
    assert cells[5].text == "mc"
    assert cells[6].text == "2.50"     # cantitate 2 decimals
    assert cells[7].text == ""          # pret_material = 0 → empty


def test_write_article_row_subcomponent():
    doc = Document()
    tbl = doc.add_table(rows=0, cols=11)
    row = tbl.add_row()
    art = _make_full_article(cod="IZF16A", nr_ordine="9.1", is_component=True, parent_code="TRA01A10P")
    _write_article_row(row, seq_nr=10, article=art)
    cells = row.cells
    assert cells[0].text == "10"         # Nr sequential
    assert cells[1].text == "9.1"        # Nr.crt from nr_ordine
    assert cells[2].text == "IZF16A"
    assert cells[3].text == "TRA01A10P"  # Cod principal


def test_write_article_row_with_prices():
    doc = Document()
    tbl = doc.add_table(rows=0, cols=11)
    row = tbl.add_row()
    art = _make_full_article(cod="X", pret_material=100.0, val_material=250.0,
                              pret_manopera=50.5, val_manopera=126.25)
    _write_article_row(row, seq_nr=1, article=art)
    cells = row.cells
    assert cells[7].text == "100,00"    # pret_material
    assert cells[8].text == "50,50"     # pret_manopera
    assert cells[9].text == ""           # pret_utilaj = 0
    assert cells[10].text == ""          # pret_transport = 0


# Tests for _write_group_section
def test_write_group_section_adds_paragraph_and_table():
    doc = Document()
    header = {"obiectivul": "PROIECT X", "obiectul": "25.4 CAV", "categoria": "1 Copertina"}
    art1 = _make_full_article(cod="TSD06XA", nr_ordine=1)
    art2 = _make_full_article(cod="IZF16A", nr_ordine="1.1", is_component=True, parent_code="TSD06XA")
    _write_group_section(doc, header, [art1, art2])
    # Should have: 1 paragraph (group title) + 1 table
    tables = doc.tables
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert len(tables) == 1
    assert "Copertina" in paragraphs[0].text or "25.4" in paragraphs[0].text
    # Table: 2 header rows + 2 article rows + 1 total row = 5 rows
    assert len(tables[0].rows) == 5


def test_write_group_section_seq_resets_per_group():
    doc = Document()
    header = {"obiectivul": "X", "obiectul": "Y", "categoria": "Z"}
    arts = [_make_full_article(cod=f"A{i}", nr_ordine=i) for i in range(1, 4)]
    _write_group_section(doc, header, arts)
    # Nr. in col 0 must start at 1 and go to 3
    tbl = doc.tables[0]
    article_rows = tbl.rows[2:5]  # skip 2 header rows, skip total row
    assert article_rows[0].cells[0].text == "1"
    assert article_rows[2].cells[0].text == "3"


# Tests for build_docx_for_source
def _make_holistic_with_one_group():
    """Create a minimal holistic JSON with one matched group."""
    art = {
        "cod": "TSD06XA", "denumire": "Test article", "um": "mc", "cantitate": 2.5,
        "nr_ordine": 1, "is_component": False, "parent_code": None,
        "pret_material": 0.0, "pret_manopera": 0.0, "pret_utilaj": 0.0, "pret_transport": 0.0,
        "val_material": 0.0, "val_manopera": 0.0, "val_utilaj": 0.0, "val_transport": 0.0,
        "deviz_header": {"obiectivul": "OBJ", "obiectul": "OBL", "categoria": "CAT"},
    }
    return {
        "matched_groups": [{"oferta_articles": [art], "ref_articles": [], "deviz_denumire": "OBJ|OBL|CAT"}],
        "oferta_only_groups": [],
        "ref_only_groups": [],
    }


def test_build_docx_for_source_creates_file():
    holistic = _make_holistic_with_one_group()
    out = tempfile.mktemp(suffix=".docx")
    try:
        build_docx_for_source(
            holistic=holistic,
            source="oferta",
            entity_name="SC. TEST SRL",
            client_name="Test Client",
            label="Oferta 1",
            output_path=out,
        )
        assert os.path.exists(out)
        assert os.path.getsize(out) > 5000
    finally:
        if os.path.exists(out):
            os.unlink(out)


def test_build_docx_for_source_header_contains_entity():
    holistic = _make_holistic_with_one_group()
    out = tempfile.mktemp(suffix=".docx")
    try:
        build_docx_for_source(
            holistic=holistic,
            source="oferta",
            entity_name="SC. KATO SERVICE SRL",
            client_name="CAV Maneciu",
            label="Oferta 1",
            output_path=out,
        )
        doc = Document(out)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "KATO" in full_text
        assert "CAV Maneciu" in full_text
    finally:
        if os.path.exists(out):
            os.unlink(out)


# ── Tests for template-exact formatting ──────────────────────────────────────

def test_col_widths_exact_twips():
    """tblGrid trebuie să conțină exact lățimile din template (twips)."""
    from docx.oxml.ns import qn
    from shared.lista_oferta_writer import _COL_WIDTHS_TWIPS, _build_table_header
    doc = Document()
    tbl = doc.add_table(rows=2, cols=11)
    _build_table_header(tbl)
    tblGrid = tbl._tbl.find(qn('w:tblGrid'))
    assert tblGrid is not None, "tblGrid missing"
    widths = [int(gc.get(qn('w:w'))) for gc in tblGrid.findall(qn('w:gridCol'))]
    assert widths == _COL_WIDTHS_TWIPS


def test_suppress_table_borders_xml():
    """_suppress_table_borders adaugă tblBorders cu toate laturile none."""
    from docx.oxml.ns import qn
    from shared.lista_oferta_writer import _suppress_table_borders
    doc = Document()
    tbl = doc.add_table(rows=2, cols=11)
    tbl.style = "Table Grid"
    _suppress_table_borders(tbl)
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    borders = tblPr.find(qn('w:tblBorders'))
    assert borders is not None, "tblBorders element missing"
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{side}'))
        assert el is not None, f"missing border side: {side}"
        assert el.get(qn('w:val')) == 'none', f"{side} val != none"


def test_group_section_header_pipe_format():
    """Paragraf grup: 'OBJ | OBL | CAT' — fara label, fara newlines."""
    from shared.lista_oferta_writer import _write_group_section
    doc = Document()
    header = {"obiectivul": "PROIECT X", "obiectul": "25.4 CAV", "categoria": "1 Copertina"}
    _write_group_section(doc, header, [], deviz_denumire="")
    # Primul paragraf non-gol este header-ul grupului
    para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert len(para_texts) >= 1
    title = para_texts[0]
    assert title == "PROIECT X | 25.4 CAV | 1 Copertina"
    assert "Obiectivul:" not in title
    assert "\n" not in title


def test_group_section_header_no_label():
    """Daca obiectivul e numeric, foloseste deviz_denumire pipe-separat."""
    from shared.lista_oferta_writer import _write_group_section
    doc = Document()
    header = {"obiectivul": "0232 000000232", "obiectul": "", "categoria": ""}
    _write_group_section(doc, header, [], deviz_denumire="DRUMURI | Ob 1 | Str. X")
    para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "DRUMURI" in para_texts[0]
    assert "Obiectivul:" not in para_texts[0]


def test_document_header_all_bold():
    """Toate 4 linii din header document trebuie sa fie bold."""
    holistic = {"matched_groups": [], "ref_only_groups": [], "oferta_only_groups": []}
    out = tempfile.mktemp(suffix=".docx")
    try:
        build_docx_for_source(
            holistic=holistic, source="oferta",
            entity_name="TestOfertant", client_name="TestClient",
            label="Oferta 1", output_path=out,
        )
        doc = Document(out)
        header_paras = [p for p in doc.paragraphs if p.text.strip()][:4]
        assert len(header_paras) == 4
        for p in header_paras:
            for run in p.runs:
                assert run.bold, f"Run not bold in: {p.text!r}"
    finally:
        if os.path.exists(out): os.unlink(out)


def test_document_margins_1_5cm():
    """Margini document: 1.5cm toate laturile."""
    from docx.shared import Cm
    holistic = {"matched_groups": [], "ref_only_groups": [], "oferta_only_groups": []}
    out = tempfile.mktemp(suffix=".docx")
    try:
        build_docx_for_source(
            holistic=holistic, source="oferta",
            entity_name="E", client_name="C", label="Oferta 1", output_path=out,
        )
        doc = Document(out)
        s = doc.sections[0]
        assert abs(s.top_margin - Cm(1.5)) < 500, f"top_margin={s.top_margin}"
        assert abs(s.bottom_margin - Cm(1.5)) < 500, f"bottom_margin={s.bottom_margin}"
    finally:
        if os.path.exists(out): os.unlink(out)
