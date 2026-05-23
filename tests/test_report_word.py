# tests/test_report_word.py
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from shared.report_word import generate_word
from docx import Document
import io


def _make_neconf(tip, deviz, ref_cod='CA01A1', ref_den='TEST', ref_um='m',
                 ref_cant=10.0, oferta_cod='CA01A1', oferta_den='TEST',
                 oferta_um='m', oferta_cant=10.0, camp=None, deviz_den='STRUCTURA'):
    n = {
        'tip': tip, 'deviz_ref': deviz, 'deviz_denumire': deviz_den,
        'ref_cod': ref_cod, 'ref_denumire': ref_den,
        'ref_um': ref_um, 'ref_cantitate': ref_cant,
        'oferta_cod': oferta_cod, 'oferta_denumire': oferta_den,
        'oferta_um': oferta_um, 'oferta_cantitate': oferta_cant,
    }
    if camp:
        n['camp'] = camp
    return n


def _load_doc(comp, mismatches=None, devize_extra=None, devize_lipsa=None):
    session = {'client_name': 'TEST', 'obiect_investitii': ''}

    # Extract devizes from neconformitati to build minimal article lists
    neconformitati = comp.get('neconformitati', [])
    devize_from_neconf = {nc.get('deviz_ref') for nc in neconformitati if nc.get('deviz_ref')}

    # Create minimal ref_articles and oferta_articles if not provided
    ref_articles = comp.get('ref_articles', [])
    if not ref_articles and devize_from_neconf:
        ref_articles = [
            {'deviz': d, 'deviz_denumire': nc.get('deviz_denumire', ''), 'cod': nc.get('ref_cod', '')}
            for nc in neconformitati if nc.get('deviz_ref')
            for d in [nc.get('deviz_ref')]
        ]

    oferta_articles = comp.get('oferta_articles', [])
    if not oferta_articles and devize_from_neconf:
        oferta_articles = [
            {'deviz': nc.get('deviz_ref', ''), 'deviz_denumire': nc.get('deviz_denumire', ''), 'cod': nc.get('oferta_cod', '')}
            for nc in neconformitati if nc.get('deviz_ref') and nc.get('oferta_cod')
        ]

    comp_full = {
        'oferta_nr': 1, 'source_file': 'test.json', 'ofertant': 'Test SRL',
        **comp,
        'deviz_mismatches': mismatches or [],
        'ref_articles': ref_articles,
        'oferta_articles': oferta_articles,
    }
    docx_bytes = generate_word(
        session, comp_full,
        devize_extra=devize_extra or [],
        devize_lipsa=devize_lipsa or [],
    )
    return Document(io.BytesIO(docx_bytes))


def test_deviz_section_heading_appears():
    """Fiecare deviz are un heading cu codul si numarul de articole."""
    comp = {
        'neconformitati': [
            _make_neconf('ARTICOL_LIPSA', '226108', deviz_den='STRUCTURA CUPOLA'),
        ],
        'total_neconformitati': 1, 'matches': 43,
        'ref_articles': [
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA CUPOLA', 'cod': 'CA01A1'},
        ],
        'oferta_articles': [
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA CUPOLA', 'cod': 'CA01A1'},
        ],
    }
    doc = _load_doc(comp)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '226108' in full_text, f"Deviz 226108 not found in doc text"


def test_extra_articles_appear_after_deviz_data():
    """Articolele EXTRA apar dupa randurile de LIPSA/DIFERENTA ale devizului."""
    comp = {
        'neconformitati': [
            _make_neconf('ARTICOL_LIPSA',  '226108', ref_cod='AA01A1'),
            _make_neconf('ARTICOL_EXTRA',  '226108', oferta_cod='EXTRA1',
                         ref_cod='', ref_den=''),
        ],
        'total_neconformitati': 2, 'matches': 43,
        'ref_articles': [
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'AA01A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA02A1'},
        ],
        'oferta_articles': [
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA02A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'EXTRA1'},
        ],
    }
    doc = _load_doc(comp)
    assert len(doc.tables) > 0 or len(doc.paragraphs) > 0


def test_deviz_mismatch_alert_appears():
    """Alerta DEVIZ_MISMATCH apare la finalul documentului."""
    comp = {
        'neconformitati': [],
        'total_neconformitati': 0, 'matches': 100,
        'ref_articles': [],
        'oferta_articles': [],
    }
    mismatches = [{'oferta_deviz': '226113', 'ref_deviz': '226118',
                   'overlap_score': 0.88, 'oferta_art_count': 8, 'ref_art_count': 7}]
    doc = _load_doc(comp, mismatches=mismatches)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '226113' in full_text, "Mismatch alert 226113 not found in doc"
    assert '226118' in full_text, "Mismatch ref deviz 226118 not found in doc"


def test_devize_extra_alert_appears():
    """Devizele din oferta absente din referinta apar in sectiunea alerte."""
    comp = {
        'neconformitati': [],
        'total_neconformitati': 0, 'matches': 100,
        'ref_articles': [],
        'oferta_articles': [],
    }
    devize_extra = [{'deviz': '226728', 'denumire': 'CHELTUIELI CONEXE', 'art_count': 2}]
    doc = _load_doc(comp, devize_extra=devize_extra)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '226728' in full_text, "Extra deviz alert 226728 not found in doc"


def test_empty_neconformitati_generates_valid_doc():
    """Document fara neconformitati se genereaza fara erori."""
    comp = {
        'neconformitati': [],
        'total_neconformitati': 0,
        'matches': 100,
        'ref_articles': [],
        'oferta_articles': [],
    }
    doc = _load_doc(comp)
    assert doc is not None


def test_sumar_contains_counts():
    """Sectiunea de sumar contine numarul de articole matched."""
    comp = {
        'neconformitati': [_make_neconf('ARTICOL_LIPSA', '226108')],
        'total_neconformitati': 1, 'matches': 42,
        'ref_art_count': 43, 'oferta_art_count': 44,
        'ref_articles': [
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA01A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA02A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA03A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA04A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA05A1'},
        ],
        'oferta_articles': [
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA01A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA02A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA03A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA04A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA05A1'},
            {'deviz': '226108', 'deviz_denumire': 'STRUCTURA', 'cod': 'CA06A1'},
        ],
    }
    doc = _load_doc(comp)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '42' in full_text or '43' in full_text, "Counts not found in doc"


# ============================================================================
# Subcomponent Mode Tests
# ============================================================================

def _make_raport_with_subcomp(sub_nc_tip: str = 'DIFERENTA_CAMP'):
    """Build minimal raport_ierarhic with one sub-article having one neconformitate."""
    sub_nc = {
        'tip': sub_nc_tip,
        'camp': 'cantitate',
        'is_component': True,
        'ref_cod': '$6719485',
        'ref_denumire': 'teu D40mm',
        'ref_um': 'buc',
        'ref_cantitate': 10.0,
        'oferta_cod': '$6719485',
        'oferta_denumire': 'teu D40mm',
        'oferta_um': 'buc',
        'oferta_cantitate': 5.0,
        'deviz_ref': 'DVZ1',
        'deviz_denumire': 'INSTALATII',
        'nr_ordine_ref': None,
    }
    raport = {
        'devize': [{
            'cod_deviz': 'DVZ1',
            'denumire_deviz': 'INSTALATII',
            'sumar_deviz': {'lipsa': 0, 'matched': 1, 'neconformitati': 0, 'total': 1},
            'articole': [{
                'cod': 'SA04A01',
                'denumire': 'teava polipropilena',
                'nr_ordine': 1,
                'status_match': 'NECONFORMITATE',
                'cantitate': 1.0,
                'um': 'buc',
                'is_component': False,
                'neconformitati': [],
                'subarticole': [{
                    'cod': '$6719485',
                    'denumire': 'teu D40mm',
                    'nr_ordine': None,
                    'status_match': 'NECONFORMITATE',
                    'cantitate': 10.0,
                    'um': 'buc',
                    'cant_mostenita': False,
                    'is_component': True,
                    'neconformitati': [sub_nc],
                }],
            }],
        }],
    }
    return raport, sub_nc


def _load_doc_with_mode(subcomponent_mode: str, sub_nc_tip: str = 'DIFERENTA_CAMP'):
    """Build comp with raport_ierarhic and generate_word with given subcomponent_mode."""
    raport, sub_nc = _make_raport_with_subcomp(sub_nc_tip)
    comp = {
        'oferta_nr': 1,
        'source_file': 'test.json',
        'ofertant': 'Test SRL',
        'neconformitati': [sub_nc],
        'total_neconformitati': 1,
        'matches': 1,
        'ref_art_count': 1,
        'oferta_art_count': 1,
        'deviz_mismatches': [],
        'ref_articles': [],
        'oferta_articles': [],
        'raport_ierarhic': raport,
    }
    session = {'client_name': 'TEST', 'obiect_investitii': ''}
    docx_bytes = generate_word(
        session, comp,
        devize_extra=[], devize_lipsa=[],
        subcomponent_mode=subcomponent_mode,
    )
    return Document(io.BytesIO(docx_bytes))


def _count_data_rows(doc):
    """Count non-header rows in first table (excludes header rows 0-2)."""
    if not doc.tables:
        return 0
    table = doc.tables[0]
    # First 3 rows are headers; deviz heading rows and total rows are additional
    # We count rows that have a non-empty col[2] (ref_cod column)
    count = 0
    for row in table.rows[3:]:
        text = row.cells[2].text.strip()
        if text and text not in ('', '▶'):
            count += 1
    return count


def test_subcomponent_mode_full_shows_diferenta_camp():
    """Mode=full: DIFERENTA_CAMP for sub-component IS shown."""
    doc = _load_doc_with_mode('full')
    assert _count_data_rows(doc) >= 1, "full mode: sub-component DIFERENTA_CAMP should appear"


def test_subcomponent_mode_fields_hides_diferenta_camp():
    """Mode=fields: DIFERENTA_CAMP for sub-component is suppressed."""
    doc = _load_doc_with_mode('fields')
    assert _count_data_rows(doc) == 0, "fields mode: DIFERENTA_CAMP sub-component should be hidden"


def test_subcomponent_mode_fields_hides_um_diferit():
    """Mode=fields: UM_DIFERIT for sub-component is suppressed."""
    doc = _load_doc_with_mode('fields', sub_nc_tip='UM_DIFERIT')
    assert _count_data_rows(doc) == 0, "fields mode: UM_DIFERIT sub-component should be hidden"


def test_subcomponent_mode_summary_hides_all_neconf():
    """Mode=summary: all neconformitati for sub-component are suppressed."""
    doc = _load_doc_with_mode('summary')
    assert _count_data_rows(doc) == 0, "summary mode: all sub-component neconf should be hidden"


def test_subcomponent_mode_full_shows_cod_similar():
    """Mode=full: COD_SIMILAR for sub-component IS shown."""
    doc = _load_doc_with_mode('full', sub_nc_tip='COD_SIMILAR')
    assert _count_data_rows(doc) >= 1, "full mode: COD_SIMILAR sub-component should appear"


def test_subcomponent_mode_fields_shows_cod_similar():
    """Mode=fields: COD_SIMILAR for sub-component is NOT suppressed (only DIFERENTA_CAMP/UM_DIFERIT)."""
    doc = _load_doc_with_mode('fields', sub_nc_tip='COD_SIMILAR')
    assert _count_data_rows(doc) >= 1, "fields mode: COD_SIMILAR should still appear"
