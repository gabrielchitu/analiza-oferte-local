# shared/diagnostics_word.py
from __future__ import annotations
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from shared.diagnostics_builder import Phase0Result, Phase1Result, Phase2Result

_ALARM_EMOJI = {"red": "🔴", "yellow": "🟡", "green": "✅"}


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_table_row(table, cells: list[str], bold: bool = False) -> None:
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True


def _sumar_oferte_table(doc: Document, oferte: list[dict]) -> None:
    headers = ["Ofertă", "Matched", "LIPSA", "EXTRA", "DEVIZ_MM"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        hdr_row.cells[i].text = h
        for run in hdr_row.cells[i].paragraphs[0].runs:
            run.bold = True
    for o in oferte:
        s = o["sumar"]
        _add_table_row(table, [
            f"Oferta {o['oferta_idx']}",
            str(s["matched"]),
            str(s["lipsa"]),
            str(s["extra"]),
            str(s["deviz_mismatch"]),
        ])


def _phase0_section(doc: Document, p0: Phase0Result) -> None:
    _heading(doc, "Phase 0 — Calitate Referinta", level=2)

    def _item(label: str, items: list[dict], severity: str) -> None:
        em = _ALARM_EMOJI[severity]
        p = doc.add_paragraph()
        run = p.add_run(f"{em} {len(items)} {label}")
        run.bold = len(items) > 0
        if items:
            for a in items[:20]:
                doc.add_paragraph(
                    f"    . {a.get('cod', '?')} — {a.get('denumire', '')[:80]}",
                    style="List Bullet"
                )
            if len(items) > 20:
                doc.add_paragraph(f"    ... si {len(items)-20} mai multe")

    _item("articole fara deviz", p0.fara_deviz, "red" if p0.fara_deviz else "green")
    _item("componente orfane (fara parent identificat)", p0.componente_orfane,
          "red" if len(p0.componente_orfane) > 10 else ("yellow" if p0.componente_orfane else "green"))
    _item("articole incomplete (cant=0 si um lipsa)", p0.incomplete,
          "yellow" if p0.incomplete else "green")


def _phase1_section(doc: Document, oferta_idx: int, p1: Phase1Result) -> None:
    _heading(doc, f"Phase 1 — Articole EXTRA — Oferta {oferta_idx} ({p1.total_extra} total)", level=2)

    if p1.total_extra == 0:
        doc.add_paragraph("✅ Niciun articol extra.")
        return

    if p1.total_extra_dollar > 0:
        doc.add_paragraph(
            f"Atentie: {p1.total_extra_dollar} din {p1.total_extra} sunt $-coduri (resurse eDevize). "
            f"Volum mare poate indica bug extragere referinta — verificati Phase 0 si PDF original."
        )

    for deviz, arts in sorted(p1.by_deviz.items()):
        doc.add_paragraph(f"Deviz {deviz} — {len(arts)} articole extra:", style="List Bullet")
        for a in arts:
            cod = a.get("oferta_cod") or "?"
            den = (a.get("oferta_denumire") or "")[:70]
            cant = a.get("oferta_cantitate")
            um = a.get("oferta_um") or ""
            cant_str = f"  cant: {cant} {um}" if cant else ""
            doc.add_paragraph(f"    . {cod} — {den}{cant_str}", style="List Bullet")


def _phase2_section(doc: Document, oferta_idx: int, p2: Phase2Result) -> None:
    _heading(doc,
             f"Phase 2 — Articole LIPSA — Oferta {oferta_idx} "
             f"({p2.total_lipsa} genuine + {p2.total_deviz_mismatch} deviz mismatch)",
             level=2)

    if p2.total_lipsa == 0 and p2.total_deviz_mismatch == 0:
        doc.add_paragraph("✅ Niciun articol lipsa.")
        return

    if p2.by_deviz:
        for deviz, arts in sorted(p2.by_deviz.items()):
            lipsa_arts = [a for a in arts if a.get("tip") == "ARTICOL_LIPSA"]
            mm_arts = [a for a in arts if a.get("tip") == "DEVIZ_MISMATCH"]
            if lipsa_arts:
                doc.add_paragraph(f"Deviz {deviz} — {len(lipsa_arts)} articole lipsa:", style="List Bullet")
                for a in lipsa_arts:
                    cod = a.get("ref_cod") or "?"
                    den = (a.get("ref_denumire") or "")[:70]
                    cant = a.get("ref_cantitate")
                    um = a.get("ref_um") or ""
                    cant_str = f"  ref: {cant} {um}" if cant else ""
                    doc.add_paragraph(f"    . {cod} — {den}{cant_str}", style="List Bullet")

    if p2.total_deviz_mismatch > 0:
        doc.add_paragraph(
            f"DEVIZ_MISMATCH ({p2.total_deviz_mismatch}): coduri prezente in oferta dar in alt deviz. "
            f"Nu sunt erori reale — verificare manuala daca devizul diferit e acceptabil."
        )


def _global_summary_table(doc: Document, data: dict) -> None:
    _heading(doc, "Sumar Global", level=1)
    sg = data["sumar_global"]
    doc.add_paragraph(f"Clienti analizati: {len(data['clienti'])}")
    doc.add_paragraph(f"Total matched: {sg['total_matched']}")
    doc.add_paragraph(f"Total LIPSA genuine: {sg['total_lipsa']}")
    doc.add_paragraph(f"Total EXTRA: {sg['total_extra']}")
    doc.add_paragraph(f"Total DEVIZ_MISMATCH: {sg['total_deviz_mismatch']}")

    if sg["clienti_cu_alarme_ref"]:
        doc.add_paragraph(
            f"Atentie — Clienti cu alarme Phase 0: {', '.join(sg['clienti_cu_alarme_ref'])}"
        )

    headers = ["Client", "Oferta", "Matched", "LIPSA", "EXTRA", "DEVIZ_MM", "Ref alarm"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        hdr_row.cells[i].text = h
        for run in hdr_row.cells[i].paragraphs[0].runs:
            run.bold = True

    for cr in data["_client_reports"]:
        for o in cr["_oferte_full"]:
            s = o["sumar"]
            alarm = cr["ref_quality"]["alarm_level"]
            _add_table_row(table, [
                cr["client"],
                f"Oferta {o['oferta_idx']}",
                str(s["matched"]),
                str(s["lipsa"]),
                str(s["extra"]),
                str(s["deviz_mismatch"]),
                _ALARM_EMOJI[alarm],
            ])


def generate_diagnostics_docx(data: dict, output_path: Path) -> None:
    doc = Document()
    doc.add_heading("Raport Diagnostic — Analizator Oferte", level=0)
    doc.add_paragraph(f"Generat: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Clienti: {', '.join(data['meta']['clienti_analizati'])}")
    doc.add_page_break()

    for cr in data["_client_reports"]:
        _heading(doc, f"CLIENT: {cr['client']}", level=1)
        _sumar_oferte_table(doc, cr["_oferte_full"])
        doc.add_paragraph()

        _phase0_section(doc, cr["_phase0"])

        for o in cr["_oferte_full"]:
            _phase1_section(doc, o["oferta_idx"], o["_phase1"])
            _phase2_section(doc, o["oferta_idx"], o["_phase2"])

        doc.add_page_break()

    _global_summary_table(doc, data)
    doc.save(output_path)
