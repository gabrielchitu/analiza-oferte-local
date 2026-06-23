"""F3-format DOCX list generator for referinta and offer articles."""

import json
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Generator, Union

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table


# Exact column widths from Template_exact.docx (twips, dxa)
_COL_WIDTHS_TWIPS = [397, 510, 1020, 1020, 3175, 567, 1020, 624, 624, 624, 624]


def _set_cell_width_twips(cell, twips: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)


def _suppress_table_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_tbl_grid(tbl) -> None:
    tbl_el = tbl._tbl
    existing = tbl_el.find(qn('w:tblGrid'))
    if existing is not None:
        tbl_el.remove(existing)
    tblGrid = OxmlElement('w:tblGrid')
    for w in _COL_WIDTHS_TWIPS:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl_el.insert(0, tblGrid)


def extract_entity_name(di_json_path: str, is_referinta: bool) -> str:
    """Extract proiectant/ofertant name from raw DI JSON (first 5 pages)."""
    marker = "PROIECTANT" if is_referinta else "CONTRACTANT (OFERTANT)"
    try:
        with open(di_json_path, encoding="utf-8") as f:
            data = json.load(f)
    except (OSError, json.JSONDecodeError):
        return "Necunoscut"

    pages = data.get("pages", [])
    for page in pages[:5]:
        lines = [ln.get("content", "").strip() for ln in page.get("lines", [])]
        for i, line in enumerate(lines):
            if marker in line:
                for candidate in lines[i + 1 :]:
                    if candidate and candidate.upper() != "SRL":
                        return candidate
    return "Necunoscut"


def _get_header_from_articles(articles: List[Dict]) -> Dict:
    """Extract group header from first article's deviz_header dict."""
    for art in articles:
        dh = art.get("deviz_header")
        if isinstance(dh, dict):
            return dh
    return {"obiectivul": "", "obiectul": "", "categoria": ""}


def _fmt_nr_crt(nr_ordine: Union[int, float, str]) -> str:
    """Format nr_ordine for display. Integers show as '1', subcomponents as '9.1'."""
    if nr_ordine is None:
        return ""
    if isinstance(nr_ordine, float) and nr_ordine == int(nr_ordine):
        return str(int(nr_ordine))
    return str(nr_ordine)


def _fmt_price(value: Optional[float]) -> str:
    """Format price: None/0.0 → empty string; else Romanian locale (1.234,50).

    Uses Python's default rounding (round-half-to-even); adequate for DOCX display.
    """
    if not value:
        return ""
    formatted = f"{value:,.2f}"
    # Convert to Romanian locale: swap . and ,
    return formatted.replace(",", "X").replace(".", ",").replace("X", ".")


def _iter_source_groups(holistic: Dict, source: str) -> Generator[Tuple[Dict, List[Dict]], None, None]:
    """Yield (header_dict, articles_list) for each non-empty group.

    Args:
        holistic: holistic JSON structure with matched_groups, ref_only_groups, oferta_only_groups
        source: "oferta" or "referinta"

    Yields:
        (header: Dict, articles: List[Dict]) tuples
    """
    art_key = "oferta_articles" if source == "oferta" else "ref_articles"
    only_key = "oferta_only_groups" if source == "oferta" else "ref_only_groups"

    def _min_page(arts: List[Dict]) -> int:
        pages = [p for a in arts for p in (a.get("source_pages") or [])]
        return min(pages) if pages else 9999

    matched = sorted(holistic.get("matched_groups", []),
                     key=lambda g: _min_page(g.get(art_key, [])))
    only    = sorted(holistic.get(only_key, []),
                     key=lambda g: _min_page(g.get("articles", [])))

    for group in matched:
        articles = group.get(art_key, [])
        if not articles:
            continue
        header = _get_header_from_articles(articles)
        yield header, articles, group.get("deviz_denumire", "")

    for group in only:
        articles = group.get("articles", [])
        if not articles:
            continue
        header = _get_header_from_articles(articles)
        yield header, articles, group.get("deviz_denumire", "")


HEADER_FILL = "D9D9D9"   # light grey for header rows
# Cols that must never wrap: all numeric/code cols except Denumire (4)
_NO_WRAP_COLS = {0, 1, 2, 3, 5, 6, 7, 8, 9, 10}


def _set_cell_no_wrap(cell) -> None:
    """Prevent cell content from wrapping to next line (w:noWrap)."""
    tcPr = cell._element.get_or_add_tcPr()
    noWrap = OxmlElement("w:noWrap")
    tcPr.append(noWrap)


def _shade_cell(cell, fill_hex: str) -> None:
    """Apply background shading to a cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._element.get_or_add_tcPr().append(shading)


def _cell_text(cell, text: str, bold: bool = False, center: bool = False, font_size: int = 8) -> None:
    """Set text content in a cell with optional formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def _merge_vertical(table: Table, col: int, row_start: int = 0, row_end: int = 1) -> None:
    """Merge cells in same column across rows."""
    table.cell(row_start, col).merge(table.cell(row_end, col))


def _build_table_header(table: Table) -> None:
    """Write 2-row F3 header into an existing 2-row, 11-col table."""
    _set_tbl_grid(table)

    row0 = table.rows[0].cells
    row1 = table.rows[1].cells

    for ci in range(11):
        _set_cell_width_twips(row0[ci], _COL_WIDTHS_TWIPS[ci])
        _set_cell_width_twips(row1[ci], _COL_WIDTHS_TWIPS[ci])
        if ci in _NO_WRAP_COLS:
            _set_cell_no_wrap(row0[ci])
            _set_cell_no_wrap(row1[ci])

    labels_row0 = ["Nr.", "Nr.crt", "Cod", "Cod principal", "Denumire", "UM", "Cantitate"]
    for i, label in enumerate(labels_row0):
        _merge_vertical(table, i)
        _cell_text(row0[i], label, bold=True, center=True)
        _shade_cell(row0[i], HEADER_FILL)

    row0[7].merge(row0[10])
    _cell_text(row0[7], "Pret unitar (lei/UM)", bold=True, center=True)
    _shade_cell(row0[7], HEADER_FILL)

    for i, label in enumerate(["Material", "Manoperă", "Utilaje", "Transport"]):
        _cell_text(row1[7 + i], label, bold=True, center=True)
        _shade_cell(row1[7 + i], HEADER_FILL)


def _write_article_row(row, seq_nr: int, article: Dict) -> None:
    """Write article data into a 15-cell table row.

    Args:
        row: docx table row with 15 cells
        seq_nr: sequential article number (1, 2, 3...)
        article: article dict with cod, denumire, um, cantitate, nr_ordine, is_component,
                 parent_code, pret_*, val_*
    """
    nr_crt = _fmt_nr_crt(article.get("nr_ordine", ""))
    parent_code = article.get("parent_code") or ""
    cantitate = article.get("cantitate", 0)

    cod = article.get("cod", "") or ""
    if not article.get("is_component") and cod.startswith("$"):
        cod = cod[1:]

    values = [
        str(seq_nr),
        nr_crt,
        cod,
        parent_code,
        article.get("denumire", ""),
        article.get("um", ""),
        f"{cantitate:.2f}" if cantitate else "",
        _fmt_price(article.get("pret_material", 0.0)),
        _fmt_price(article.get("pret_manopera", 0.0)),
        _fmt_price(article.get("pret_utilaj", 0.0)),
        _fmt_price(article.get("pret_transport", 0.0)),
    ]

    font_size = 7 if article.get("is_component") else 8
    # 0=Nr center, 1=Nr.crt right, 2=Cod left, 3=CodPrincipal left, 4=Denumire left
    # 5=UM center, 6=Cantitate right, 7-10=Pret cols right
    right_cols = {1, 6, 7, 8, 9, 10}
    center_cols = {0, 5}

    for i, (cell, val) in enumerate(zip(row.cells, values)):
        cell.text = ""
        if i in _NO_WRAP_COLS:
            _set_cell_no_wrap(cell)
        p = cell.paragraphs[0]
        if i in right_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif i in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(font_size)
        if article.get("is_component"):
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _set_table_fixed_layout(tbl: Table) -> None:
    """Lock table column widths — prevent Word autofit from resizing columns."""
    tblPr = tbl._tbl.tblPr
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _is_numeric_only(val: str) -> bool:
    """True when value contains only digits, spaces, dots — no descriptive text."""
    import re
    return bool(val) and bool(re.fullmatch(r'[\d\s.]+', val.strip()))


def _nr_sort_key(art: Dict):
    """Sort key: parse nr_ordine so '1' < '1.1' < '1.2' < '2' < '2.1'."""
    nr = art.get("nr_ordine", "")
    try:
        parts = str(nr).split(".")
        return (int(parts[0]), int(parts[1]) if len(parts) > 1 else 0)
    except (ValueError, TypeError):
        return (9999, 0)


def _write_group_section(doc: Document, header: Dict, articles: List[Dict],
                         deviz_denumire: str = "") -> None:
    """Write group title paragraph + F3 table. Nr. sequence resets to 1 per group."""
    articles = sorted(articles, key=_nr_sort_key)
    obiectivul = header.get("obiectivul", "")
    obiectul   = header.get("obiectul", "")
    categoria  = header.get("categoria", "")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)

    # When obiectivul/obiectul are pure CPV numbers, use deviz_denumire as title
    if _is_numeric_only(obiectivul) and deviz_denumire:
        raw_parts = [s.strip() for s in deviz_denumire.split("|")]
        title = " | ".join(part for part in raw_parts if part)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9)
    else:
        parts = [s.strip() for s in [obiectivul, obiectul, categoria] if s.strip()]
        run = p.add_run(" | ".join(parts))
        run.bold = True
        run.font.size = Pt(9)

    # Count main vs subcomponent articles
    main_count = sum(1 for a in articles if not a.get("is_component"))
    sub_count = sum(1 for a in articles if a.get("is_component"))

    n_rows = 2 + len(articles) + 1
    tbl = doc.add_table(rows=n_rows, cols=11)
    tbl.style = "Table Grid"
    _suppress_table_borders(tbl)
    _set_table_fixed_layout(tbl)

    _build_table_header(tbl)

    for i, art in enumerate(articles):
        row = tbl.rows[2 + i]
        _write_article_row(row, seq_nr=i + 1, article=art)

    total_row = tbl.rows[-1]
    total_row.cells[0].merge(total_row.cells[10])
    total_cell = total_row.cells[0]
    total_cell.text = ""
    run = total_cell.paragraphs[0].add_run(
        f"Total grup: {main_count} articole principale"
        + (f" / {sub_count} subcomponente" if sub_count else "")
    )
    run.bold = True
    run.font.size = Pt(8)
    _shade_cell(total_cell, "F2F2F2")


def build_docx_for_source(
    holistic: Dict,
    source: str,
    entity_name: str,
    client_name: str,
    label: str,
    output_path: str,
) -> None:
    """Generate full F3-format DOCX lista for given source ('oferta' or 'referinta').

    Args:
        holistic: loaded holistic_oferta_N.json dict
        source: "oferta" or "referinta"
        entity_name: ofertant or proiectant name
        client_name: display client name
        label: "Oferta 1" or "Referinta"
        output_path: where to save the .docx
    """
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    entity_label = "Ofertant" if source == "oferta" else "Proiectant"
    for line, bold, size in [
        (f"Lista articole — {label}", True, 14),
        (f"Client: {client_name}", True, 11),
        (f"{entity_label}: {entity_name}", True, 11),
        (f"Generat: {date.today().isoformat()}", True, 9),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)

    for header, articles, deviz_den in _iter_source_groups(holistic, source=source):
        _write_group_section(doc, header, articles, deviz_denumire=deviz_den)

    doc.save(output_path)
