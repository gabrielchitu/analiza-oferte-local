import io
from datetime import date
from itertools import groupby
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT

RED = RGBColor(0xCC, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
YELLOW_FILL = "FFFF99"
ORANGE_FILL = "FFB347"
GRAY_FILL = "D9D9D9"
SUBCOMP_GRAY_FILL = "E8E8E8"  # Light gray for subcomponents
LILA_FILL = "C8A0DC"  # Lila pentru DESCRIERE_DIFERITA

# Tipuri de neconformitate suprimate per mod subcomponente
SUPPRESSED_BY_MODE: dict[str, frozenset] = {
    "full":    frozenset(),
    "fields":  frozenset({"DIFERENTA_CAMP", "UM_DIFERIT"}),
    "summary": frozenset({"DIFERENTA_CAMP", "UM_DIFERIT", "COD_SIMILAR",
                           "DESCRIERE_DIFERITA", "EROARE_ARITMETICA"}),
}

# Column widths in cm (total ~24cm for landscape A4 with 2cm margins)
COL_WIDTHS_CM = [0.7, 3.0, 1.5, 3.5, 0.8, 1.2, 1.5, 3.5, 0.8, 1.2, 6.3]


def _group_articles_by_parent(articles: list) -> list:
    """Group component articles under their parent articles.

    Returns list of dicts:
    [
        {
            "parent": {...article...},
            "components": [{...article...}, ...]
        }
    ]
    """
    result = []
    parents_dict = {}

    # Separate parents and components
    for art in articles:
        if not art.get("is_component"):
            parents_dict[art["cod"]] = {
                "parent": art,
                "components": []
            }
        else:
            parent_cod = art.get("parent_code")
            if parent_cod and parent_cod in parents_dict:
                parents_dict[parent_cod]["components"].append(art)

    # Build result preserving order
    seen_parents = set()
    for art in articles:
        if not art.get("is_component") and art["cod"] not in seen_parents:
            result.append(parents_dict[art["cod"]])
            seen_parents.add(art["cod"])

    return result


def _format_article_with_hierarchy(article: dict, parent_code: str = None) -> str:
    """Format article text with component hierarchy information.

    For components, includes [COMPONENT] prefix and parent reference.
    """
    cod = article.get("cod", "")
    is_component = article.get("is_component", False)

    if is_component:
        parent_ref = f" (parent: {parent_code})" if parent_code else ""
        return f"[COMPONENT] {cod}{parent_ref}"
    return cod


def _observatie_text(neconf: dict) -> str:
    """Return human-readable Romanian observation text for a nonconformity."""
    tip = neconf.get("tip", "")
    camp = neconf.get("camp", "")
    CAMP_LABELS = {
        "cantitate": "Cantitate",
        "pret_material": "Preț material",
        "pret_manopera": "Preț manoperă",
        "pret_utilaj": "Preț utilaj",
        "pret_transport": "Preț transport",
        "val_material": "Valoare material",
        "val_manopera": "Valoare manoperă",
        "val_utilaj": "Valoare utilaj",
        "val_transport": "Valoare transport",
    }
    if tip == "ARTICOL_LIPSA":
        return "Articol lipsă din ofertă"
    if tip == "ARTICOL_EXTRA":
        return "Articol suplimentar în ofertă (nu există în referință)"
    if tip == "UM_DIFERIT":
        return (f"Unitate de măsură diferită: referință {neconf.get('ref_um', '')},"
                f" ofertat {neconf.get('oferta_um', '')}")
    if tip == "DIFERENTA_CAMP":
        label = CAMP_LABELS.get(camp, camp)
        ref_val = neconf.get("ref", neconf.get(f"ref_{camp}", ""))
        oferta_val = neconf.get("oferta", neconf.get(f"oferta_{camp}", ""))
        um = neconf.get("ref_um", "") if camp == "cantitate" else "lei"
        return f"{label} diferită: referință {ref_val} {um}, ofertat {oferta_val} {um}".strip()
    if tip == "EROARE_ARITMETICA":
        return (f"Eroare aritmetică: {camp} declarat {neconf.get('declarat', '')} lei"
                f" ≠ calculat {neconf.get('calculat', '')} lei")
    if tip == "COD_SIMILAR":
        ref_cod = neconf.get("ref_cod", "")
        oferta_cod = neconf.get("oferta_cod", "")
        motiv = neconf.get("motiv_similaritate", "")
        return (f"Cod similar — posibilă eroare OCR sau variație: "
                f"referință '{ref_cod}', ofertat '{oferta_cod}'. "
                f"{motiv}. Necesită verificare manuală.")
    if tip == "ARTICOL_ORPHAN":
        ref_deviz = neconf.get("deviz_ref", "")
        cod = neconf.get("ref_cod", "")
        motiv = neconf.get("motiv", "")
        return f"ORPHAN: Cod '{cod}' identic, dar categorii DIFERITE. {motiv}. VERIFICARE MANUALA NECESARA!"
    if tip == "DESCRIERE_DIFERITA":
        sim = neconf.get("similaritate", "")
        ref_val = neconf.get("ref", "")
        oferta_val = neconf.get("oferta", "")
        return (f"Descriere diferită (similaritate {sim}): "
                f"referință '{ref_val}', ofertat '{oferta_val}'")
    return tip


def _set_cell_shading(cell, fill_hex: str):
    """Set background color of a cell (e.g. 'FFFF99' for yellow)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}" w:val="clear"/>')
    tcPr.append(shd)


def _style_cell(cell, size_pt: float, bold: bool = False,
                color: RGBColor = None, center: bool = False):
    """Apply font size, bold, color, and alignment to all runs in a cell."""
    for para in cell.paragraphs:
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(size_pt)
            run.bold = bold
            if color:
                run.font.color.rgb = color


def _set_landscape(doc):
    """Set page orientation to landscape A4 with 2cm margins."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(2.0))


def _set_col_widths(table):
    """Apply fixed column widths to all cells in each column."""
    for i, w in enumerate(COL_WIDTHS_CM):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)


def _get_subcomponent_style():
    """Return style dict for subcomponent rows: light gray background."""
    return {
        'background_color': SUBCOMP_GRAY_FILL,  # Light gray
        'indent': 0.2,  # inches
    }


def _get_subcomponent_badge():
    """Return text badge for subcomponent marking."""
    return '[Subcomponent]'


def _format_page_range(pages: list) -> str:
    """Formateaza lista de pagini: [5]→'5', [3,4,5]→'3-5', [3,5,7]→'3, 5, 7'."""
    if not pages:
        return ""
    pages = sorted(set(int(p) for p in pages))
    if len(pages) == 1:
        return str(pages[0])
    if pages[-1] - pages[0] == len(pages) - 1:
        return f"{pages[0]}-{pages[-1]}"
    return ", ".join(str(p) for p in pages)


def _add_group_heading(table, group_type: str, ref_header, oferta_header,
                       ref_deviz_cod: str = "", oferta_deviz_cod: str = "", deviz_denumire: str = "") -> None:
    """Rand separator de grup cu informatii 3-layer (OBIECTIVUL/Obiectul/Categoria)."""
    sep_cells = table.add_row().cells
    sep_cells[0].merge(sep_cells[10])

    header = ref_header or oferta_header
    if header:
        obj1 = (getattr(header, 'obiectivul', None) or "").strip()
        obj2 = (getattr(header, 'obiectul', None) or "").strip()
        cat = (getattr(header, 'categoria', None) or "").strip()
        label = f"OBIECTIVUL: {obj1}  │  Obiectul: {obj2}  │  Categoria: {cat}"
    elif deviz_denumire:
        label = f"Deviz: {deviz_denumire}"
    else:
        label = f"Deviz: {ref_deviz_cod or oferta_deviz_cod}"

    if group_type == "ref_only":
        label = f"[GRUP ABSENT DIN OFERTA]  {label}"
        fill = "FFB3B3"
    elif group_type == "oferta_only":
        label = f"[GRUP ABSENT DIN REFERINTA]  {label}"
        fill = "FFE0B3"
    else:
        fill = GRAY_FILL

    run = sep_cells[0].paragraphs[0].add_run(label)
    run.bold = True
    _style_cell(sep_cells[0], 8, bold=True)
    _set_cell_shading(sep_cells[0], fill)


def _add_deviz_heading(table, deviz_cod: str, deviz_den: str,
                       ref_count: int, oferta_count: int,
                       ref_pages: list = None, oferta_pages: list = None) -> None:
    """Adaugă rând separator de deviz cu numărătoare ref vs ofertă + pagini PDF."""
    sep_cells = table.add_row().cells
    sep_cells[0].merge(sep_cells[10])
    delta = oferta_count - ref_count
    delta_str = f"{delta:+d}" if delta != 0 else "0 ✓"
    den_short = deviz_den[:40] + "..." if len(deviz_den) > 40 else deviz_den
    label = (
        f"Capitol de lucrări {deviz_cod}"
        + (f" — {den_short}" if den_short else "")
        + f"  │  LIPSA: {ref_count}"
        + f"  │  EXTRA: {oferta_count}"
        + f"  │  Delta: {delta_str}"
    )
    ref_range = _format_page_range(ref_pages or [])
    oferta_range = _format_page_range(oferta_pages or [])
    if ref_range:
        label += f"  │  Ref: PDF pag. {ref_range}"
    if oferta_range:
        label += f"  │  Oferta: PDF pag. {oferta_range}"
    run = sep_cells[0].paragraphs[0].add_run(label)
    run.bold = True
    _style_cell(sep_cells[0], 9, bold=True)
    _set_cell_shading(sep_cells[0], GRAY_FILL)


def _add_extra_subheader(table) -> None:
    """Adaugă rând sub-separator pentru secțiunea 'Extra în ofertă'."""
    sub = table.add_row().cells
    sub[0].merge(sub[10])
    run = sub[0].paragraphs[0].add_run(
        "▸ Articole extra în ofertă — verificare manuală recomandată"
    )
    run.italic = True
    _style_cell(sub[0], 8)
    _set_cell_shading(sub[0], YELLOW_FILL)


def _add_quality_alerts(doc, deviz_mismatches: list,
                        devize_extra: list, devize_lipsa: list) -> None:
    """Adaugă secțiunea ALERTE DE CALITATE la finalul documentului."""
    if not any([deviz_mismatches, devize_extra, devize_lipsa]):
        return

    doc.add_page_break()
    h = doc.add_heading("ALERTE DE CALITATE — VERIFICARE MANUALĂ", level=1)
    if h.runs:
        h.runs[0].font.color.rgb = RED

    if deviz_mismatches:
        doc.add_heading("Capitol de lucrări cu cod diferit față de referință (Mismatch)", level=2)
        for m in deviz_mismatches:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(
                f"Devizul {m['oferta_deviz']} din ofertă (~{m['overlap_score']:.0%} overlap) "
                f"pare echivalentul devizului {m['ref_deviz']} din proiect "
                f"({m['oferta_art_count']} vs {m['ref_art_count']} articole)."
            )
            run.bold = True
            run.font.color.rgb = RED
            doc.add_paragraph(
                "   → Ofertantul poate fi utilizat o numerotare diferită a categoriilor. "
                "Verificați dacă articolele corespund celor din proiect.",
                style='List Bullet'
            )

    if devize_extra:
        doc.add_heading("Capitole de lucrări în ofertă, absente din referință", level=2)
        doc.add_paragraph(
            "Aceste devize NU au putut fi comparate cu proiectul. "
            "Verificați dacă sunt lucrări suplimentare sau F3 neextras din referință."
        )
        for d in devize_extra:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(
                f"{d['deviz']}"
                + (f" — {d['denumire']}" if d.get('denumire') else "")
                + f" ({d.get('art_count', '?')} articole în ofertă)"
            ).bold = True

    if devize_lipsa:
        doc.add_heading("Capitole de lucrări din referință neacoperite de ofertă", level=2)
        doc.add_paragraph(
            "Aceste categorii de lucrări din proiect nu au nicio ofertă corespunzătoare."
        )
        for d in devize_lipsa:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(
                f"{d['deviz']}"
                + (f" — {d['denumire']}" if d.get('denumire') else "")
                + f" ({d.get('art_count', '?')} articole în referință)"
            ).bold = True


def _build_header(table, ofertant_name: str):
    """Build 3-row merged header for the 11-column table."""
    r0 = table.rows[0].cells
    r1 = table.rows[1].cells
    r2 = table.rows[2].cells

    # Vertical merges (rowspan 3): Nr (col 0), Deviz (col 1), Observatii (col 10)
    r0[0].merge(r2[0])
    r0[1].merge(r2[1])
    r0[10].merge(r2[10])

    # Row 0: horizontal group merges
    r0[2].merge(r0[5])   # CERINTA spans cols 2-5
    r0[6].merge(r0[9])   # CE A OFERTAT spans cols 6-9

    # Row 1: ofertant sub-header under CE A OFERTAT (cols 6-9)
    r1[6].merge(r1[9])

    # Set texts via add_run on fresh paragraphs
    r0[0].paragraphs[0].add_run("Nr.\ncrt.")
    r0[1].paragraphs[0].add_run("Categoria\nde lucrări")
    r0[2].paragraphs[0].add_run("CERINȚĂ")
    r0[6].paragraphs[0].add_run("CE A OFERTAT")
    r0[10].paragraphs[0].add_run("OBSERVAȚII")
    r1[6].paragraphs[0].add_run(ofertant_name)

    for i, txt in enumerate(["Cod", "Denumire", "UM", "Cant."], start=2):
        r2[i].paragraphs[0].add_run(txt)
    for i, txt in enumerate(["Cod", "Denumire", "UM", "Cant."], start=6):
        r2[i].paragraphs[0].add_run(txt)

    # Style header cells — only cells with actual content (not merged continuations)
    for cell in [r0[0], r0[1], r0[2], r0[6], r0[10]]:
        _style_cell(cell, 9, bold=True, center=True)
    _style_cell(r1[6], 9, bold=True, center=True)
    for i in range(2, 10):
        _style_cell(r2[i], 9, bold=True, center=True)


def _add_audit_section(doc, audit_data: dict) -> None:
    """Adaugă secțiunea de audit la finalul documentului Word."""
    from shared.f3_auditor import _is_false_positive
    audit_status = audit_data.get("audit_status", "pending")
    all_disc = audit_data.get("discrepante", [])
    # Filtrăm false pozitive (manoperă $NNNN, totale T2-T9 etc.) — doar discrepanțe reale
    discrepante = [d for d in all_disc if not _is_false_positive(d.get("cod", ""))]

    doc.add_page_break()
    doc.add_heading("Audit extracție articole F3", level=1)

    if audit_status == "pending":
        doc.add_paragraph(
            "Auditul LLM este în curs. Regenerați raportul după finalizare."
        )
        return

    if audit_status == "ok":
        p = doc.add_paragraph("Auditorul AI nu a detectat nicio discrepanta fata de parserul regex.")
        p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        return

    # discrepante_gasite
    doc.add_paragraph(
        f"Auditorul AI a detectat {len(discrepante)} cod(uri) reale prezente in documentul OCR "
        f"dar neextrase de parser (false pozitive filtrate: {len(all_disc) - len(discrepante)}). "
        f"Verificati manual."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Cod'
    hdr[1].text = 'Obiect probabil'
    hdr[2].text = 'Context OCR'
    for d in discrepante:
        row = table.add_row().cells
        row[0].text = d.get('cod', '')
        row[1].text = d.get('deviz_probabil', '')
        ctx = d.get('context', '')
        row[2].text = (ctx[:200] + '...') if len(ctx) > 200 else ctx


def _add_neconf_row(table, row_nr: int, neconf: dict, deviz_map: dict,
                    use_ref_ordine: bool = False) -> None:
    """Adaugă un rând de neconformitate în tabel."""
    row = table.add_row().cells
    tip  = neconf.get("tip", "")
    camp = neconf.get("camp", "")
    is_suspect = bool(neconf.get("suspect"))
    is_subcomp = neconf.get('is_component', False)

    nr_ordine_ref = neconf.get("nr_ordine_ref")
    nr_ordine_of = neconf.get("nr_ordine_oferta")
    ref_pages = neconf.get("ref_source_pages") or []
    of_pages = neconf.get("oferta_source_pages") or []

    ref_pag = str(ref_pages[0]) if ref_pages else "-"
    of_pag = str(of_pages[0]) if of_pages else "-"

    # Linia 1: pagini fizice PDF (ref/oferta)
    pag_text = f"pag.{ref_pag}/{of_pag}"

    # Linia 2: nr ordine in grup (ref/oferta)
    nr_ref_str = str(nr_ordine_ref) if nr_ordine_ref is not None else "-"
    nr_of_str = str(nr_ordine_of) if nr_ordine_of is not None else "-"
    ord_text = f"({nr_ref_str}/{nr_of_str})"

    nr_text = f"{pag_text}\n{ord_text}"
    row[0].paragraphs[0].add_run(nr_text)

    # Col 1: categorie de lucrări (deviz) — identic cu v6.f
    ref_cod = neconf.get('ref_cod', '')
    parent_cod_ref = neconf.get("parent_cod_ref") if is_subcomp else None
    # display_parent_cod afisat pt subcomponente reale (is_component=True)
    # SI pt coduri $ cu parinte (resurse breviar asociate unui articol normativ)
    # Distinctia de principal/secundar e vizuala — UM-ul nu se mai mosteneste
    is_dollar_with_parent = ref_cod.startswith('$') and neconf.get("display_parent_cod")
    display_parent = neconf.get("display_parent_cod") if (is_subcomp or is_dollar_with_parent) else None
    effective_parent = parent_cod_ref or display_parent

    deviz_den_full = neconf.get("deviz_denumire", "") or neconf.get("deviz_ref", "")
    parts = [p.strip() for p in deviz_den_full.split(" | ") if p.strip()]
    # Group heading already shows OBIECTIVUL — column shows Obiectul | Categoria (specific parts)
    deviz_display = " | ".join(parts[-2:]) if len(parts) >= 2 else deviz_den_full
    row[1].paragraphs[0].add_run(deviz_display).bold = True

    # Col 2: cod REF ierarhic — principal sus+stânga, secundar jos+dreapta indentat
    badge = _get_subcomponent_badge() if is_subcomp else ''
    if is_subcomp or display_parent:
        # Subarticol: principal sus (stânga), secundar jos (dreapta indentat)
        p2 = row[2].paragraphs[0]
        if effective_parent:
            par_run = p2.add_run(effective_parent)
            par_run.bold = True
            par_run.font.size = Pt(9)
            par_run.underline = True
            p_sub = row[2].add_paragraph()
        else:
            p_sub = p2
        sub_label = f"{badge} {ref_cod}".strip() if badge else ref_cod
        sub_run = p_sub.add_run(sub_label)
        sub_run.font.size = Pt(8)
        p_sub.paragraph_format.left_indent = Pt(14)
    else:
        cod_run = row[2].paragraphs[0].add_run(ref_cod)
        cod_run.bold = True
        cod_run.font.size = Pt(9)

    denom = str(neconf.get("ref_denumire", ""))
    if len(denom) > 50:
        denom = denom[:47] + "..."
    row[3].paragraphs[0].add_run(denom)

    # Add indentation to denomination for subcomponents
    if is_subcomp:
        paragraph = row[3].paragraphs[0]
        paragraph.paragraph_format.left_indent = Pt(18)  # 18 points indent

    ref_um_run   = row[4].paragraphs[0].add_run(str(neconf.get("ref_um", "")))
    ref_cant_run = row[5].paragraphs[0].add_run(str(neconf.get("ref_cantitate", "")))

    oferta_um_run = oferta_cant_run = None
    if tip != "ARTICOL_LIPSA":
        oferta_cod = str(neconf.get("oferta_cod", ""))
        # principal oferta: din ref (pt articole matched/neconf) sau din oferta (pt EXTRA)
        # oferta_display_parent_cod pt subcomponente reale SI coduri $ cu parinte
        is_oferta_dollar_with_parent = oferta_cod.startswith('$') and neconf.get("oferta_display_parent_cod")
        oferta_parent = effective_parent or (neconf.get("oferta_display_parent_cod") if (is_subcomp or is_oferta_dollar_with_parent) else None)
        # Col 6: oferta cod ierarhic — principal sus+stânga (bold), secundar jos+dreapta indentat
        if oferta_parent and oferta_cod and oferta_cod != oferta_parent:
            p6 = row[6].paragraphs[0]
            par6_run = p6.add_run(oferta_parent)
            par6_run.bold = True
            par6_run.font.size = Pt(9)
            par6_run.underline = True
            p6_sub = row[6].add_paragraph()
            sub6_run = p6_sub.add_run(oferta_cod)
            sub6_run.font.size = Pt(8)
            p6_sub.paragraph_format.left_indent = Pt(14)
        else:
            oferta_cod_run = row[6].paragraphs[0].add_run(oferta_cod)
            oferta_cod_run.bold = True
            oferta_cod_run.font.size = Pt(9)
        oferta_denom = str(neconf.get("oferta_denumire", ""))
        if len(oferta_denom) > 50:
            oferta_denom = oferta_denom[:47] + "..."
        row[7].paragraphs[0].add_run(oferta_denom)

        # Add indentation to offer denomination for subcomponents
        if is_subcomp:
            paragraph = row[7].paragraphs[0]
            paragraph.paragraph_format.left_indent = Pt(18)  # 18 points indent
        oferta_um_run   = row[8].paragraphs[0].add_run(str(neconf.get("oferta_um", "")))
        oferta_cant_run = row[9].paragraphs[0].add_run(str(neconf.get("oferta_cantitate", "")))

    obs_text = _observatie_text(neconf)
    if is_suspect:
        motiv = neconf.get("motiv_suspiciune", "")
        obs_text += f"\n⚠ {motiv}" if motiv else "\n⚠"
    obs_run = row[10].paragraphs[0].add_run(obs_text)
    obs_run.bold = True
    obs_run.font.color.rgb = RED

    for cell in row:
        _style_cell(cell, 8)

    # Apply gray background for subcomponents first
    if is_subcomp:
        for cell in row:
            _set_cell_shading(cell, SUBCOMP_GRAY_FILL)

    if is_suspect:
        for cell in row: _set_cell_shading(cell, YELLOW_FILL)
    if tip == "COD_SIMILAR":
        for cell in row: _set_cell_shading(cell, ORANGE_FILL)
    if tip == "ARTICOL_LIPSA":
        for cell in row: _set_cell_shading(cell, "FFB3B3")  # Light red for missing articles
    if tip == "ARTICOL_EXTRA":
        for cell in row: _set_cell_shading(cell, YELLOW_FILL)
    if tip == "ARTICOL_ORPHAN":
        for cell in row: _set_cell_shading(cell, "FFCC99")
    if tip == "DESCRIERE_DIFERITA":
        for cell in row: _set_cell_shading(cell, LILA_FILL)

    if tip == "DIFERENTA_CAMP" and camp == "cantitate":
        ref_cant_run.font.color.rgb = RED
        if oferta_cant_run: oferta_cant_run.font.color.rgb = RED
    elif tip == "UM_DIFERIT":
        ref_um_run.font.color.rgb = RED
        if oferta_um_run: oferta_um_run.font.color.rgb = RED


def _add_deviz_summary_row(table, row_nr: int, neconf_count: int, ref_total: int, offer_total: int):
    """Add summary footer for a deviz: shows neconformities count and totals."""
    row = table.add_row()
    cells = row.cells

    # Column 1: "SUMAR" label
    cells[1].text = "SUMAR"
    _style_cell(cells[1], 9, bold=True, color=BLACK)
    _set_cell_shading(cells[1], GRAY_FILL)

    # Columns 2-5: summary text
    summary_text = f"Neconformitati: {neconf_count} din {ref_total} articole ref."
    cells[2].text = summary_text
    _style_cell(cells[2], 9, bold=True, color=BLACK)
    _set_cell_shading(cells[2], GRAY_FILL)
    for i in range(3, 6):
        cells[i].text = ""
        _set_cell_shading(cells[i], GRAY_FILL)

    # Column 6: Reference total article count
    cells[6].text = str(ref_total)
    _style_cell(cells[6], 9, bold=True, center=True, color=BLACK)
    _set_cell_shading(cells[6], GRAY_FILL)

    # Columns 7-8: empty (filler)
    cells[7].text = ""
    _set_cell_shading(cells[7], GRAY_FILL)
    cells[8].text = ""
    _set_cell_shading(cells[8], GRAY_FILL)

    # Column 9: Offer total article count
    cells[9].text = str(offer_total)
    _style_cell(cells[9], 9, bold=True, center=True, color=BLACK)
    _set_cell_shading(cells[9], GRAY_FILL)

    cells[10].text = ""
    _set_cell_shading(cells[10], GRAY_FILL)


def _count_main_articles(articles: list) -> int:
    return sum(1 for a in articles if not a.get("is_component", False))


def _add_group_totals_row(table, ref_count: Optional[int], oferta_count: Optional[int]) -> None:
    """Append mirrored totals row after a holistic group.

    Cols 0-1: "TOTAL GRUP"
    Cols 2-5: "Referință: N articole" (empty when ref_count is None)
    Cols 6-9: "Ofertă: M articole"   (empty when oferta_count is None)
    Col 10:   empty
    """
    cells = table.add_row().cells

    # Label: merge cols 0-1
    cells[0].merge(cells[1])
    cells[0].paragraphs[0].add_run("TOTAL GRUP").bold = True
    _style_cell(cells[0], 9, bold=True)
    _set_cell_shading(cells[0], GRAY_FILL)

    # Ref side: merge cols 2-5
    cells[2].merge(cells[5])
    if ref_count is not None:
        cells[2].paragraphs[0].add_run(f"Referință: {ref_count} articole principale").bold = True
        _style_cell(cells[2], 9, bold=True)
    _set_cell_shading(cells[2], GRAY_FILL)

    # Offer side: merge cols 6-9
    cells[6].merge(cells[9])
    if oferta_count is not None:
        cells[6].paragraphs[0].add_run(f"Ofertă: {oferta_count} articole principale").bold = True
        _style_cell(cells[6], 9, bold=True)
    _set_cell_shading(cells[6], GRAY_FILL)

    # Col 10: empty
    _set_cell_shading(cells[10], GRAY_FILL)


# ── Hierarchical DOCX helpers ──────────────────────────────────────────────


def _add_deviz_section_header(doc, dv: dict) -> None:
    """Adaugă heading deviz cu sumar matched/lipsa/neconformitati."""
    cod = dv.get('cod_deviz', '')
    den = dv.get('denumire_deviz', '')
    sumar = dv.get('sumar_deviz', {})
    heading = doc.add_heading(f"DEVIZ: {cod} — {den}", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    p.add_run(
        f"Total: {sumar.get('total', 0)}  │  "
        f"Matched: {sumar.get('matched', 0)}  │  "
        f"Lipsă: {sumar.get('lipsa', 0)}  │  "
        f"Neconformități: {sumar.get('neconformitati', 0)}"
    ).bold = True


def _has_neconformitate(art: dict) -> bool:
    """True dacă articolul principal sau oricare subarticol are neconformitate."""
    if art.get('neconformitati') or art.get('status_match') in ('LIPSA', 'NECONFORMITATE'):
        return True
    return any(
        s.get('neconformitati') or s.get('status_match') in ('LIPSA', 'NECONFORMITATE')
        for s in art.get('subarticole', [])
    )


def _add_neconf_detail(doc, nc: dict, indent: int = 1) -> None:
    """Adaugă rând detaliu neconformitate."""
    tip = nc.get('tip', '')
    camp = nc.get('camp', '')
    val_ref = nc.get('ref_cantitate') or nc.get('ref_um') or ''
    val_oferta = nc.get('oferta_cantitate') or nc.get('oferta_um') or ''
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(indent * 24)
    run = p.add_run(f"→ {tip}")
    if camp:
        p.add_run(f" [{camp}]")
    if val_ref or val_oferta:
        p.add_run(f": ref={val_ref}  ofertă={val_oferta}")


def _add_subarticol_block(doc, sub: dict) -> None:
    """Adaugă rând subarticol indentat cu marker ⚠ dacă are neconformitate."""
    status = sub.get('status_match', '')
    has_nc = bool(sub.get('neconformitati')) or status in ('LIPSA', 'NECONFORMITATE')
    nr = sub.get('nr_ordine', '')

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    label = f"  [{nr}] {sub.get('cod', '')} — {sub.get('denumire', '')}"
    if has_nc:
        run = p.add_run(f"⚠ {label.strip()}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    else:
        p.add_run(label)

    detail = f"    Ref: {sub.get('cantitate', '')} {sub.get('um', '')}  │  {status}"
    if sub.get('cant_mostenita'):
        detail += "  [cantitate moștenită]"
    p.add_run(f"\n{detail}")

    for nc in sub.get('neconformitati', []):
        _add_neconf_detail(doc, nc, indent=2)


def _add_article_block(doc, art: dict) -> None:
    """Adaugă blocul articol principal + subarticole."""
    status = art.get('status_match', '')
    icon = '⚠' if status in ('LIPSA', 'NECONFORMITATE') else '✓'
    nr = art.get('nr_ordine', '')

    p = doc.add_paragraph()
    run_nr = p.add_run(f"[{nr}] {art.get('cod', '')} — {art.get('denumire', '')}")
    run_nr.bold = True
    p.add_run(f"\n    Ref: {art.get('cantitate', '')} {art.get('um', '')}  │  {icon} {status}")

    for nc in art.get('neconformitati', []):
        _add_neconf_detail(doc, nc, indent=1)

    for sub in art.get('subarticole', []):
        _add_subarticol_block(doc, sub)


def _add_principal_context_row(table, art: dict, deviz_denumire: str) -> None:
    """Rând context pentru articolul principal MATCHED când subarticolele lui au neconformitati."""
    row = table.add_row().cells
    nr = art.get('nr_ordine')
    row[0].paragraphs[0].add_run(str(nr) if nr is not None else '')
    parts = [p.strip() for p in deviz_denumire.split(" | ") if p.strip()]
    deviz_display = " | ".join(parts[-2:]) if len(parts) >= 2 else deviz_denumire
    row[1].paragraphs[0].add_run(deviz_display).bold = True
    cod_run = row[2].paragraphs[0].add_run(str(art.get('cod', '')))
    cod_run.bold = True
    cod_run.font.size = Pt(9)
    denom = str(art.get('denumire', ''))
    row[3].paragraphs[0].add_run(denom[:50] + '...' if len(denom) > 50 else denom)
    row[4].paragraphs[0].add_run(str(art.get('um', '')))
    row[5].paragraphs[0].add_run(str(art.get('cantitate', '') or ''))
    obs = row[10].paragraphs[0].add_run('▶ articol principal (matched)')
    obs.italic = True
    for cell in row:
        _style_cell(cell, 8)
        _set_cell_shading(cell, 'D6EAD6')


def _add_deviz_total_row_hierarchical(table, n_main: int, n_sub: int,
                                      n_offer_main: int = 0) -> None:
    """Linie total deviz: ref principale (stânga) vs ofertă principale (dreapta).
    Diferența colorată: roșu=LIPSA, portocaliu=EXTRA, verde=egal.
    """
    row = table.add_row().cells
    _set_cell_shading(row[0], GRAY_FILL)

    row[1].paragraphs[0].add_run("TOTAL DEVIZ").bold = True
    _style_cell(row[1], 9, bold=True, color=BLACK)
    _set_cell_shading(row[1], GRAY_FILL)

    # Cols 2-5: referință
    row[2].merge(row[5])
    ref_txt = f"Ref: {n_main} principale"
    if n_sub:
        ref_txt += f"  +  {n_sub} subarticole"
    row[2].paragraphs[0].add_run(ref_txt).bold = True
    _style_cell(row[2], 9, color=BLACK)
    _set_cell_shading(row[2], GRAY_FILL)

    # Cols 6-9: ofertă cu semnal vizual
    row[6].merge(row[9])
    p_offer = row[6].paragraphs[0]
    offer_run = p_offer.add_run(f"Ofertă: {n_offer_main} principale")
    offer_run.bold = True
    if n_offer_main < n_main:
        offer_run.font.color.rgb = RED
        sig = p_offer.add_run(f"  ▼ {n_main - n_offer_main} lipsă")
        sig.font.color.rgb = RED
    elif n_offer_main > n_main:
        offer_run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        sig = p_offer.add_run(f"  ▲ {n_offer_main - n_main} extra")
        sig.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
    else:
        offer_run.font.color.rgb = RGBColor(0x00, 0x77, 0x00)
        p_offer.add_run("  ✓").font.color.rgb = RGBColor(0x00, 0x77, 0x00)
    _style_cell(row[6], 9, color=BLACK)
    _set_cell_shading(row[6], GRAY_FILL)

    _set_cell_shading(row[10], GRAY_FILL)


def _generate_word_hierarchical(doc, raport: dict, comp: dict,
                                deviz_mismatches_list: list, devize_extra: list,
                                devize_lipsa: list, audit_data: dict,
                                subcomponent_mode: str = "full") -> None:
    """Tabel ierarhic: TOATE devizele in ordinea ref.
    Devize fara neconf: cap + linie total.
    Devize cu neconf: cap + randuri principale/subarticole + linie total.
    Nr ordine = cel din referinta. Principalul apare intotdeauna inaintea subarticolelor sale.
    """
    from collections import defaultdict as _dd

    deviz_map: dict = {}
    for dv in raport.get('devize', []):
        deviz_map[dv.get('cod_deviz', '')] = dv.get('denumire_deviz', '')

    extra_by_deviz: dict = _dd(list)
    for nc in comp.get('neconformitati', []):
        if nc.get('tip') == 'ARTICOL_EXTRA':
            extra_by_deviz[nc.get('deviz_ref', '')].append(nc)

    ofertant_name = comp.get("ofertant") or comp.get("source_file", "")
    table = doc.add_table(rows=3, cols=11)
    table.style = "Table Grid"
    _build_header(table, ofertant_name)

    row_nr = 0

    suppressed = SUPPRESSED_BY_MODE.get(subcomponent_mode, frozenset())

    def _visible(ncs: list) -> list:
        if not suppressed:
            return ncs
        return [nc for nc in ncs if nc.get('tip') not in suppressed]

    for dv in raport.get('devize', []):
        dv_cod = dv.get('cod_deviz', '')
        dv_den = dv.get('denumire_deviz', '')
        sumar = dv.get('sumar_deviz', {})
        n_lipsa = sumar.get('lipsa', 0)
        n_extra = len(extra_by_deviz.get(dv_cod, []))
        n_main = sumar.get('total', 0)
        n_matched = sumar.get('matched', 0)
        n_neconf = sumar.get('neconformitati', 0)  # prezente în ofertă dar cu diferențe
        n_offer_main = n_matched + n_neconf + n_extra  # articole principale în ofertă
        n_sub = sum(len(art.get('subarticole', [])) for art in dv.get('articole', []))

        # Pre-check: will there be any visible rows for this deviz?
        has_visible_rows = False
        for art in dv.get('articole', []):
            is_comp = art.get('is_component', False) or art.get('cod', '').startswith('$')
            raw_ncs = art.get('neconformitati', [])
            principal_ncs = _visible(raw_ncs) if is_comp else raw_ncs
            if principal_ncs:
                has_visible_rows = True
                break
            subs_with_ncs = [
                s for s in art.get('subarticole', [])
                if _visible(s.get('neconformitati', []))
            ]
            if subs_with_ncs:
                has_visible_rows = True
                break
        if extra_by_deviz.get(dv_cod, []):
            has_visible_rows = True

        # Only add rows if there are visible neconformities
        if not has_visible_rows:
            continue

        # Colecteaza source_pages din articolele devizului curent
        _ref_src_pages = []
        _oferta_src_pages = []
        for art in dv.get('articole', []):
            _ref_src_pages.extend(art.get('ref_source_pages', []))
            _oferta_src_pages.extend(art.get('oferta_source_pages', []))

        _add_deviz_heading(table, dv_cod, dv_den, ref_count=n_lipsa, oferta_count=n_extra,
                           ref_pages=_ref_src_pages, oferta_pages=_oferta_src_pages)

        for art in dv.get('articole', []):
            is_comp = art.get('is_component', False) or art.get('cod', '').startswith('$')
            raw_ncs = art.get('neconformitati', [])
            principal_ncs = _visible(raw_ncs) if is_comp else raw_ncs

            subs_with_ncs = [
                s for s in art.get('subarticole', [])
                if _visible(s.get('neconformitati', []))
            ]

            if not principal_ncs and not subs_with_ncs:
                continue

            if not principal_ncs and subs_with_ncs:
                # Principal MATCHED dar subarticole cu neconf: rand context verde
                _add_principal_context_row(table, art, dv_den or dv_cod)
            else:
                for nc in principal_ncs:
                    row_nr += 1
                    _add_neconf_row(table, row_nr, nc, deviz_map, use_ref_ordine=True)

            for sub in subs_with_ncs:
                for nc in _visible(sub.get('neconformitati', [])):
                    row_nr += 1
                    _add_neconf_row(table, row_nr, nc, deviz_map, use_ref_ordine=True)

        extra_items = extra_by_deviz.get(dv_cod, [])
        if extra_items:
            _add_extra_subheader(table)
            for nc in extra_items:
                row_nr += 1
                _add_neconf_row(table, row_nr, nc, deviz_map, use_ref_ordine=True)

        _add_deviz_total_row_hierarchical(table, n_main, n_sub, n_offer_main)

    _set_col_widths(table)
    doc.add_paragraph()
    _add_quality_alerts(doc, deviz_mismatches_list, devize_extra, devize_lipsa)
    if audit_data:
        _add_audit_section(doc, audit_data)

    # Caz B: articole fără deviz — secțiune la final
    nelocalizate = raport.get('articole_nelocalizate', [])
    if nelocalizate:
        doc.add_heading("Articole nelocalizate — verificare manuală", level=2)
        doc.add_paragraph(
            "Articolele de mai jos au fost extrase dar nu au putut fi alocate "
            "unui deviz. Verificați manual în documentele PDF."
        )
        for art in nelocalizate:
            sursa = art.get('sursa', 'ref').upper()
            tip = 'subarticol' if art.get('is_component') else 'principal'
            doc.add_paragraph(
                f"• [{sursa}] {art.get('cod')} — {art.get('denumire')} "
                f"[tip: {tip}]",
                style='List Bullet'
            )


def _generate_word_flat(doc, neconformitati: list, deviz_mismatches_list: list,
                        devize_extra: list, devize_lipsa: list, audit_data: dict,
                        comp: dict) -> None:
    """Generează corpul documentului plat (stil vechi cu tabel)."""
    if not neconformitati:
        doc.add_paragraph("Nicio neconcordanță detectată.")
        doc.add_paragraph()
        _add_quality_alerts(doc, deviz_mismatches_list, devize_extra, devize_lipsa)
        if audit_data:
            _add_audit_section(doc, audit_data)
        return

    from collections import Counter, defaultdict as _dd_total
    from itertools import groupby as _groupby

    # Build per-deviz TOTAL article counts from full article lists (not just non-conforming)
    _ref_deviz_totals = _dd_total(int)
    _oferta_deviz_totals = _dd_total(int)
    for art in comp.get('ref_articles', []):
        d = art.get('deviz', '')
        if d:
            _ref_deviz_totals[d] += 1
    for art in comp.get('oferta_articles', []):
        d = art.get('deviz', '')
        if d:
            _oferta_deviz_totals[d] += 1

    # deviz_map: deviz_cod -> deviz_den
    deviz_map: dict = {}
    for nc in neconformitati:
        d = nc.get("deviz_ref", "")
        n = nc.get("deviz_denumire", "")
        if d and n and not n.startswith("REF:"):
            deviz_map[d] = n

    lipsa_by_deviz = Counter(
        nc.get("deviz_ref", "")
        for nc in neconformitati if nc.get("tip") == "ARTICOL_LIPSA"
    )
    extra_by_deviz = Counter(
        nc.get("deviz_ref", "")
        for nc in neconformitati if nc.get("tip") == "ARTICOL_EXTRA"
    )

    nec_normale = [nc for nc in neconformitati if nc.get("tip") != "ARTICOL_EXTRA"]
    nec_extra   = [nc for nc in neconformitati if nc.get("tip") == "ARTICOL_EXTRA"]

    sorted_nec   = sorted(nec_normale, key=lambda x: x.get("deviz_ref", ""))
    sorted_extra = sorted(nec_extra,   key=lambda x: x.get("deviz_ref", ""))

    extra_per_deviz: dict = {}
    for deviz_key, grp in _groupby(sorted_extra, key=lambda x: x.get("deviz_ref", "")):
        extra_per_deviz[deviz_key] = list(grp)

    # Collect all deviz keys in order (normale + only-extra)
    deviz_keys_normale = list(dict.fromkeys(
        nc.get("deviz_ref", "") for nc in sorted_nec
    ))
    only_extra_devize = sorted(set(extra_per_deviz.keys()) - set(deviz_keys_normale))
    all_deviz_keys = deviz_keys_normale + only_extra_devize

    # Paragraph listing deviz codes (visible in doc.paragraphs for navigation)
    if all_deviz_keys:
        p_devize = doc.add_paragraph()
        p_devize.add_run("Devize: " + ", ".join(str(k) for k in all_deviz_keys)).bold = True

    ofertant_name = comp.get("ofertant") or comp.get("source_file", "")
    table = doc.add_table(rows=3, cols=11)
    table.style = "Table Grid"
    _build_header(table, ofertant_name)

    row_nr = 0
    processed_devize: set = set()

    for deviz_key, group_items in _groupby(sorted_nec, key=lambda x: x.get("deviz_ref", "")):
        processed_devize.add(deviz_key)
        items = list(group_items)
        deviz_cod = str(deviz_key) if deviz_key else ""
        deviz_den = deviz_map.get(deviz_cod, "")
        n_lipsa = lipsa_by_deviz.get(deviz_cod, 0)
        n_extra = extra_by_deviz.get(deviz_cod, 0)
        _add_deviz_heading(table, deviz_cod, deviz_den,
                           ref_count=n_lipsa, oferta_count=n_extra)
        for neconf in items:
            row_nr += 1
            _add_neconf_row(table, row_nr, neconf, deviz_map)

        extra_items = extra_per_deviz.get(deviz_cod, [])
        if extra_items:
            _add_extra_subheader(table)
            for neconf in extra_items:
                row_nr += 1
                _add_neconf_row(table, row_nr, neconf, deviz_map)

        # Single summary row at the end of this deviz block
        all_neconf_items = items + extra_items
        neconf_ref_arts = {nc.get('ref_cod') for nc in all_neconf_items if nc.get('ref_cod')}
        neconf_count = len(neconf_ref_arts)
        ref_total = _ref_deviz_totals.get(deviz_cod, 0) or neconf_count
        offer_total = _oferta_deviz_totals.get(deviz_cod, 0)
        _add_deviz_summary_row(table, row_nr + 1, neconf_count, ref_total, offer_total)
        row_nr += 1

    for deviz_key in only_extra_devize:
        deviz_cod = str(deviz_key)
        deviz_den = deviz_map.get(deviz_cod, "")
        n_extra = extra_by_deviz.get(deviz_cod, 0)
        # Colecteaza pagini sursa din articolele EXTRA ale devizului
        _oferta_src_pages = []
        for art in extra_per_deviz.get(deviz_cod, []):
            _oferta_src_pages.extend(art.get("oferta_source_pages", []))
        _add_deviz_heading(table, deviz_cod, deviz_den, ref_count=0, oferta_count=n_extra,
                           ref_pages=[], oferta_pages=_oferta_src_pages)
        _add_extra_subheader(table)
        extra_items = extra_per_deviz[deviz_key]
        for neconf in extra_items:
            row_nr += 1
            _add_neconf_row(table, row_nr, neconf, deviz_map)

        # Summary row: use actual ref/oferta totals from full article lists
        ref_total = _ref_deviz_totals.get(deviz_cod, 0)
        offer_total = _oferta_deviz_totals.get(deviz_cod, 0)
        neconf_count = len(extra_items)  # count of extra articles only
        _add_deviz_summary_row(table, row_nr + 1, neconf_count, ref_total, offer_total)
        row_nr += 1

    _set_col_widths(table)

    doc.add_paragraph()
    _add_quality_alerts(doc, deviz_mismatches_list, devize_extra, devize_lipsa)

    if audit_data:
        _add_audit_section(doc, audit_data)


def _add_unassigned_section(doc, unassigned_articles: list) -> None:
    if not unassigned_articles:
        return

    doc.add_heading("Articole neasignate — deviz neidentificat", level=2)
    intro = doc.add_paragraph(
        "Articolele de mai jos nu au putut fi atribuite unui grup deviz "
        "(Obiectiv + Obiect + Categorie). Verificați manual pagina sursă din document."
    )
    intro.runs[0].italic = True

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Sursă", "Cod", "Denumire", "Deviz detectat", "Pagini"]):
        hdr[i].paragraphs[0].add_run(text).bold = True
        _set_cell_shading(hdr[i], "D3D3D3")

    for art in unassigned_articles:
        row = tbl.add_row().cells
        row[0].text = "REF" if art.get("source") == "ref" else "OFERTĂ"
        row[1].text = art.get("cod", "")
        row[2].text = (art.get("denumire") or "")[:80]
        row[3].text = art.get("deviz", "")
        pages = art.get("source_pages", [])
        row[4].text = ", ".join(str(p) for p in pages) if pages else "—"
        for cell in row:
            _set_cell_shading(cell, "FFFACD")


def _generate_word_holistic(doc, raport_holistic: dict, comp: dict,
                             subcomponent_mode: str = "full") -> None:
    """
    Genereaza continut Word in structura holistica (group-based).
    Sectiuni: matched, ref-only (LIPSA), oferta-only (EXTRA), ungrouped.
    """
    table = doc.add_table(rows=3, cols=11)
    table.style = "Table Grid"
    _set_col_widths(table)
    ofertant_name = comp.get("ofertant") or comp.get("source_file", "")
    _build_header(table, ofertant_name)

    deviz_map: dict = {}
    suppressed = SUPPRESSED_BY_MODE.get(subcomponent_mode, frozenset())
    row_nr = 0

    def _visible(ncs_list: list) -> list:
        if not suppressed:
            return ncs_list
        return [n for n in ncs_list if n.get("tip") not in suppressed]

    # --- Grupuri matchate ---
    for mg in raport_holistic.get("matched_groups", []):
        _add_group_heading(table, "matched",
                           mg.get("ref_header"), mg.get("oferta_header"),
                           mg.get("ref_deviz_cod", ""), mg.get("oferta_deviz_cod", ""),
                           mg.get("deviz_denumire", ""))
        visible = _visible(mg.get("neconformitati", []))
        if visible:
            for nc in visible:
                row_nr += 1
                _add_neconf_row(table, row_nr, nc, deviz_map)
        else:
            info_row = table.add_row().cells
            info_row[0].merge(info_row[10])
            info_row[0].paragraphs[0].add_run("  ✓ Grup conform — fara neconformitati").italic = True
            _style_cell(info_row[0], 8)

    # --- Grupuri doar in referinta (LIPSA) ---
    for rg in raport_holistic.get("ref_only_groups", []):
        _add_group_heading(table, "ref_only",
                           rg.get("ref_header"), None,
                           ref_deviz_cod=rg.get("ref_deviz_cod", ""),
                           deviz_denumire=rg.get("deviz_denumire", ""))
        for nc in rg.get("neconformitati", []):
            row_nr += 1
            _add_neconf_row(table, row_nr, nc, deviz_map)

    # --- Grupuri doar in oferta (EXTRA) ---
    for og in raport_holistic.get("oferta_only_groups", []):
        _add_group_heading(table, "oferta_only",
                           None, og.get("oferta_header"),
                           oferta_deviz_cod=og.get("oferta_deviz_cod", ""),
                           deviz_denumire=og.get("deviz_denumire", ""))
        for nc in og.get("neconformitati", []):
            row_nr += 1
            _add_neconf_row(table, row_nr, nc, deviz_map)

    # --- Articole fara grup ---
    ungrouped = raport_holistic.get("ungrouped", [])
    if ungrouped:
        ug_cells = table.add_row().cells
        ug_cells[0].merge(ug_cells[10])
        run = ug_cells[0].paragraphs[0].add_run(
            f"[ARTICOLE FARA GRUP — {len(ungrouped)} articole fara deviz detectat]"
        )
        run.bold = True
        _style_cell(ug_cells[0], 8, bold=True)
        _set_cell_shading(ug_cells[0], "FFCCFF")

    # --- Articole neasignate (deviz_key INCOMPLETE) ---
    unassigned = raport_holistic.get("unassigned_articles", [])
    _add_unassigned_section(doc, unassigned)


def generate_word(
    session: dict,
    comp: dict,
    comparison_mode: str = "cu_pret",
    audit_data: dict = None,
    devize_extra: list = None,
    devize_lipsa: list = None,
    subcomponent_mode: str = "full",
) -> bytes:
    """Generate a Word document for a single comparatie (oferta).

    Args:
        session: session dict with client_name, obiect_investitii
        comp: comparatie dict with neconformitati, oferta_nr, source_file, deviz_mismatches, etc.
        comparison_mode: "cu_pret" or "fara_pret"
        devize_extra: list[dict] cu {'deviz', 'denumire', 'art_count'} —
                      devize prezente in oferta dar absente din referinta.
        devize_lipsa: list[dict] cu {'deviz', 'denumire', 'art_count'} —
                      devize din referinta fara nicio oferta.

    Returns:
        bytes: DOCX document content
    """
    doc = Document()
    _set_landscape(doc)

    # Document header
    title = doc.add_heading("TABEL NECONCORDANȚE", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Client: {session.get('client_name', '')}")
    obiect = session.get("obiect_investitii", "")
    if obiect:
        doc.add_paragraph(f"Obiectul investiției: {obiect}")
    doc.add_paragraph(f"Data: {date.today()}")

    nr_oferta = comp.get("oferta_nr", "?")
    source = comp.get("source_file", "")
    ofertant_name = comp.get("ofertant") or source
    doc.add_heading(f"OFERTA {nr_oferta} — {ofertant_name}", level=1)

    neconformitati = comp.get("neconformitati", [])
    deviz_mismatches_list = comp.get("deviz_mismatches", [])
    devize_extra = devize_extra or []
    devize_lipsa = devize_lipsa or []

    # ── SUMAR ────────────────────────────────────────────────────────
    ref_art_count    = comp.get("ref_art_count", "?")
    oferta_art_count = comp.get("oferta_art_count", "?")

    # In holistic mode, get correct matched count from holistic sumar
    raport_holistic_comp = comp.get("raport_holistic") or {}
    holistic_sumar = raport_holistic_comp.get("sumar", {})
    if holistic_sumar:
        total_matched = holistic_sumar.get("total_matched_articles", 0)
        total_neconf = sum(holistic_sumar.get("neconformitati_by_tip", {}).values())
    else:
        total_matched = comp.get("matches", 0)
        total_neconf = comp.get("total_neconformitati", 0)

    p_sumar = doc.add_paragraph()
    p_sumar.add_run(
        f"Articole referință: {ref_art_count}  │  "
        f"Articole ofertă: {oferta_art_count}  │  "
        f"Matched: {total_matched}  │  "
        f"Neconformități: {total_neconf}"
    ).bold = True

    # ── Dispatch: holistic, hierarchical, or flat ────────────────────
    raport_holistic = comp.get('raport_holistic')
    raport_ierarhic = comp.get('raport_ierarhic')
    if raport_holistic:
        _generate_word_holistic(doc, raport_holistic, comp,
                                subcomponent_mode=subcomponent_mode)
    elif raport_ierarhic:
        _generate_word_hierarchical(doc, raport_ierarhic, comp,
                                    deviz_mismatches_list, devize_extra,
                                    devize_lipsa, audit_data,
                                    subcomponent_mode=subcomponent_mode)
    else:
        _generate_word_flat(doc, neconformitati, deviz_mismatches_list,
                            devize_extra, devize_lipsa, audit_data, comp)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
