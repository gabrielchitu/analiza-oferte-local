"""Side-by-side comparison DOCX: all ref articles on left, matched oferta on right."""
from __future__ import annotations

from datetime import date
from typing import Dict, List, Optional, Tuple

from rapidfuzz import fuzz as _fuzz

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table

# ── Colors ────────────────────────────────────────────────────────────────────
FILL_EXTRA       = "CCFFCC"   # light green  — articol extra în ofertă (NC relevant)
FILL_LIPSA       = "FFCCCC"   # light red    — articol lipsă în ofertă (NC relevant)
FILL_COD_SIMILAR = "F0F0F0"   # light grey   — caracteristici similare, cod diferit
FILL_NC          = "FFFACD"   # pale yellow  — alte neconformități (pereche matched)
FILL_COD_NORM    = "FFD966"   # orange-yellow — COD_NORMATIV_DIFERIT
FILL_SPEC_DIFF   = "FFC000"   # amber        — SPECIFICATIE_DIFERITA
FILL_PARAM_DIFF  = "F0F0F0"   # pale gray    — DIFERENTA_PARAMETRU (de-emphasized; obs text red)
FILL_REAL_NC     = "FFD966"   # amber-yellow — DIFERENTA_CAMP + UM_DIFERIT (real actionable NCs)
FILL_HDR   = "D9D9D9"
FILL_TOTAL = "F2F2F2"

# NC types not printed in raport (COD_SIMILAR = match logic artifact, not actionable)
_SUPPRESSED_TIPS: frozenset = frozenset({"COD_SIMILAR"})
TEXT_RED   = RGBColor(0xCC, 0x00, 0x00)

# ── Table layout ──────────────────────────────────────────────────────────────
# Landscape A4: 29.7 - 2×1.5cm margins = 26.7cm available
# 0=NrRef 1=CodRef 2=DenRef 3=UMRef 4=CantRef | 5=NrOf 6=CodOf 7=DenOf 8=UMOf 9=CantOf | 10=NC
N_COLS = 11
COL_W   = [1.0, 1.6, 5.2, 1.0, 1.4,  1.0, 1.6, 5.2, 1.0, 1.4,  4.3]
NO_WRAP = {0, 3, 4, 5, 8, 9}   # Nr, UM, Cant — Cod (1,6) allowed to wrap for long codes


# ── XML / layout helpers ──────────────────────────────────────────────────────

def _shade(cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    cell._element.get_or_add_tcPr().append(shd)


def _no_wrap(cell) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(OxmlElement("w:noWrap"))


def _fixed_layout(tbl: Table) -> None:
    tblPr = tbl._tbl.tblPr
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _set_tbl_grid(tbl: Table) -> None:
    """Define column widths at table level (twips) so fixed layout is stable."""
    tblGrid = OxmlElement("w:tblGrid")
    for w_cm in COL_W:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(round(w_cm * 567))))  # 1cm = 567 twips
        tblGrid.append(gc)
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is not None:
        tbl_pr.addnext(tblGrid)
    else:
        tbl_el.insert(0, tblGrid)


def _set_cell_margins(tbl: Table) -> None:
    """Reduce default cell padding (1.9mm → top/bottom 0.5mm, left/right 1mm)."""
    tblPr = tbl._tbl.tblPr
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, twips in [("top", 28), ("left", 57), ("bottom", 28), ("right", 57)]:
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:w"), str(twips))
        elem.set(qn("w:type"), "dxa")
        tblCellMar.append(elem)
    tblPr.append(tblCellMar)


def _repeat_header_rows(tbl: Table) -> None:
    """Mark first 2 rows as header rows so they repeat on each page."""
    for row in tbl.rows[:2]:
        tr = row._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)


def _set_landscape(doc: Document) -> None:
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin   = Cm(1.5)
    sec.right_margin  = Cm(1.5)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)


# ── Formatting helpers ────────────────────────────────────────────────────────

def _fmt_nr(nr_ordine) -> str:
    if nr_ordine is None:
        return ""
    if isinstance(nr_ordine, float) and nr_ordine == int(nr_ordine):
        return str(int(nr_ordine))
    return str(nr_ordine)


def _fmt_nr_with_page(art: dict) -> tuple[str, str]:
    """Return (nr_text, page_text) for a two-line nr cell."""
    nr = _fmt_nr(art.get("nr_ordine"))
    pages = art.get("source_pages") or []
    page = f"p.{pages[0]}" if pages else ""
    return nr, page


def _fmt_cant(v) -> str:
    try:
        f = float(v)
        return f"{f:.2f}" if f else ""
    except (TypeError, ValueError):
        return str(v) if v else ""


# ── NC text building ──────────────────────────────────────────────────────────

def _nc_text(nc: dict) -> str:
    tip = nc.get("tip", "")
    if tip == "ARTICOL_EXTRA":
        den = nc.get("oferta_denumire", "")
        return f"EXTRA: {den}" if den else "EXTRA: absent din referință"
    if tip == "ARTICOL_LIPSA":
        den = nc.get("ref_denumire", "")
        return f"LIPSĂ: {den}" if den else "LIPSĂ: absent din ofertă"
    if tip == "COD_SIMILAR":
        return nc.get("motiv_similaritate") or "COD_SIMILAR"
    if tip == "DIFERENTA_PARAMETRU":
        ref_p = ", ".join(nc.get("ref_parametri", [])) or "—"
        off_p = ", ".join(nc.get("oferta_parametri", [])) or "—"
        return f"PARAM DIFERIȚI: REF [{ref_p}] ↔ OF [{off_p}]"
    if tip == "DIFERENTA_CAMP":
        camp = nc.get("camp", "")
        ref_v = nc.get("ref", "")
        of_v  = nc.get("oferta", "")
        return f"Câmp {camp}: REF={ref_v} → OF={of_v}"
    if tip == "DIFERENTA_CANTITATE":
        ref_c = nc.get("ref_cantitate", "")
        of_c  = nc.get("oferta_cantitate", "")
        return f"Cant: REF={ref_c} → OF={of_c}"
    if tip == "COD_NORMATIV_DIFERIT":
        diffs = "; ".join(
            f"{d.get('camp','')}: {d.get('ref','')}→{d.get('oferta','')}"
            if "ref" in d else d.get("detaliu", "")
            for d in nc.get("diferente", [])
        )
        motiv = nc.get("motiv_llm", "")
        parts = [p for p in [diffs, motiv] if p]
        ref_cod = nc.get("ref_cod", "")
        off_cod = nc.get("oferta_cod", "")
        prefix = f"COD DIFERIT ({ref_cod}↔{off_cod})"
        return f"{prefix}: {' | '.join(parts)}" if parts else prefix
    if tip == "SPECIFICATIE_DIFERITA":
        nota = nc.get("nota_specialist", "")
        return f"SPEC DIFERITA: {nota}" if nota else "SPEC DIFERITA: specificație tehnic diferită"
    return tip


def _ncs_for(ncs: list, ref_art: Optional[dict], of_art: Optional[dict]) -> list:
    ref_cod = (ref_art or {}).get("cod") or ""
    of_cod  = (of_art or {}).get("cod") or ""
    ref_nr  = (ref_art or {}).get("nr_ordine")
    of_nr   = (of_art or {}).get("nr_ordine")
    result  = []
    for nc in ncs:
        nc_ref = nc.get("ref_cod") or ""
        nc_of  = nc.get("oferta_cod") or ""
        match_ref = bool(ref_cod and nc_ref == ref_cod)
        match_of  = bool(of_cod  and nc_of  == of_cod)
        if not (match_ref or match_of):
            continue
        # Narrow by nr_ordine when NC carries it — prevents same-cod cross-contamination
        nc_ref_nr = nc.get("nr_ordine_ref")
        nc_of_nr  = nc.get("nr_ordine_oferta")
        if match_ref and ref_nr is not None and nc_ref_nr is not None:
            if str(nc_ref_nr) != str(ref_nr):
                match_ref = False
        if match_of and of_nr is not None and nc_of_nr is not None:
            if str(nc_of_nr) != str(of_nr):
                match_of = False
        if not (match_ref or match_of):
            continue
        if ref_art and of_art:
            if (not nc_ref or match_ref) and (not nc_of or match_of):
                result.append(nc)
        else:
            result.append(nc)
    return result


# ── Fuzzy suggestion ──────────────────────────────────────────────────────────

def _fuzzy_suggest(art: Optional[dict], candidates: List[dict], min_score: int = 85) -> Optional[str]:
    """Find best denomination fuzzy-match for art in candidates.

    Returns annotation string "POSIBIL ACELAȘI MATERIAL (N%): COD — den" or None.
    Only considers principal articles (not sub-components) with denomination ≥ 5 chars.
    """
    if not art or not candidates:
        return None
    den = (art.get("denumire") or "").lower().strip()
    if len(den) < 5:
        return None
    best_score = 0
    best_cand = None
    for cand in candidates:
        if cand.get("is_component"):
            continue
        cand_den = (cand.get("denumire") or "").lower().strip()
        if len(cand_den) < 5:
            continue
        score = _fuzz.token_set_ratio(den, cand_den)
        if score > best_score:
            best_score = score
            best_cand = cand
    if best_score >= min_score and best_cand:
        cand_cod = best_cand.get("cod") or ""
        if not best_cand.get("is_component") and cand_cod.startswith("$"):
            cand_cod = cand_cod[1:]
        cand_den_short = (best_cand.get("denumire") or "")[:45]
        return f"POSIBIL ACELAȘI MATERIAL ({int(best_score)}%): {cand_cod} — {cand_den_short}"
    return None


# ── Nr helpers ────────────────────────────────────────────────────────────────

def _major_nr(art: dict) -> int:
    nr = art.get("nr_ordine")
    if nr is None:
        return 0
    try:
        return int(str(nr).split(".")[0])
    except (ValueError, TypeError):
        return 0


def _sort_key(art: dict):
    nr = str(art.get("nr_ordine", "0") or "0")
    parts = nr.split(".")
    try:
        return (int(parts[0]), int(parts[1]) if len(parts) > 1 else 0)
    except (ValueError, IndexError):
        return (0, 0)


# ── Row building ──────────────────────────────────────────────────────────────
# RowTuple = (ref_art | None, of_art | None, fill_hex | None, [nc_text, ...])
RowTuple = Tuple[Optional[dict], Optional[dict], Optional[str], List[str]]


def _sub_rows(ref_arts, of_arts, ncs, parent_nr, parent_fill) -> List[RowTuple]:
    """Pair sub-components by nr_ordine. Unpaired → SUBCOMP LIPSĂ / SUBCOMP EXTRA."""
    ref_subs = sorted(
        [a for a in ref_arts if a.get("is_component") and _major_nr(a) == parent_nr],
        key=_sort_key,
    )
    of_subs = sorted(
        [a for a in of_arts if a.get("is_component") and _major_nr(a) == parent_nr],
        key=_sort_key,
    )
    ref_by_nr = {a.get("nr_ordine"): a for a in ref_subs}
    of_by_nr  = {a.get("nr_ordine"): a for a in of_subs}
    all_nrs   = sorted(
        set(ref_by_nr) | set(of_by_nr),
        key=lambda x: _sort_key({"nr_ordine": x}),
    )
    rows: List[RowTuple] = []
    for nr in all_nrs:
        rsub = ref_by_nr.get(nr)
        osub = of_by_nr.get(nr)
        if rsub and osub:
            rows.append((rsub, osub, parent_fill, []))
        elif rsub:
            den = rsub.get("denumire", "") or ""
            rows.append((rsub, None, FILL_LIPSA, [f"SUBCOMP LIPSĂ: {den}" if den else "SUBCOMP LIPSĂ"]))
        else:
            den = osub.get("denumire", "") or ""
            rows.append((None, osub, FILL_EXTRA, [f"SUBCOMP EXTRA: {den}" if den else "SUBCOMP EXTRA"]))
    return rows


def _pair_fill(ref_art, of_art, pair_ncs, of_main, ref_main):
    """Shared fill + NC-text logic for one row pair."""
    nc_texts = [_nc_text(nc) for nc in pair_ncs]
    if ref_art and not of_art:
        fill = FILL_LIPSA
        if not nc_texts:
            nc_texts = ["LIPSĂ: absent din ofertă"]
        suggestion = _fuzzy_suggest(ref_art, of_main)
        if suggestion:
            nc_texts.append(suggestion)
    elif not ref_art and of_art:
        fill = FILL_EXTRA
        if not nc_texts:
            nc_texts = ["EXTRA: absent din referință"]
        suggestion = _fuzzy_suggest(of_art, ref_main)
        if suggestion:
            nc_texts.append(suggestion)
    elif pair_ncs:
        tips = {nc.get("tip") for nc in pair_ncs}
        if "COD_NORMATIV_DIFERIT" in tips:
            fill = FILL_COD_NORM
        elif "SPECIFICATIE_DIFERITA" in tips:
            fill = FILL_SPEC_DIFF
        elif "DIFERENTA_PARAMETRU" in tips:
            fill = FILL_PARAM_DIFF
        elif "DIFERENTA_CAMP" in tips or "UM_DIFERIT" in tips:
            fill = FILL_REAL_NC
        elif tips <= {"COD_SIMILAR", "DESCRIERE_DIFERITA"}:
            fill = FILL_COD_SIMILAR
        else:
            fill = FILL_NC
    else:
        fill = None
    return fill, nc_texts


def _build_rows_from_matches(ref_main, of_main, ref_arts, of_arts, ncs, matches) -> List[RowTuple]:
    """Pair rows from the comparator's verdicts, not by nr_ordine.

    nr_ordine pairing breaks on Naipu: eDevize offers have nr=None (all collapse
    to key 0) and N:1-merged groups have overlapping nr sequences.

    Ground truth: ARTICOL_LIPSA/ARTICOL_EXTRA NCs say WHICH articles are
    unmatched. Everything else was matched by the comparator (possibly with
    OCR-normalized codes) — after removing the LIPSA refs and EXTRA offers,
    the two sides align POSITIONALLY (both follow document order).
    """
    from collections import Counter

    lipsa_q: Counter = Counter(
        str(nc.get("ref_cod") or "") for nc in ncs
        if nc.get("tip") == "ARTICOL_LIPSA" and not nc.get("is_component"))
    extra_q: Counter = Counter(
        str(nc.get("oferta_cod") or "") for nc in ncs
        if nc.get("tip") == "ARTICOL_EXTRA" and not nc.get("is_component"))

    from collections import defaultdict, deque

    def _k(c) -> str:
        return str(c or "").lstrip("$").upper()

    ref_flags = []  # (ref_art, is_lipsa)
    for a in ref_main:
        c = str(a.get("cod") or "")
        if lipsa_q.get(c, 0) > 0:
            lipsa_q[c] -= 1
            ref_flags.append((a, True))
        else:
            ref_flags.append((a, False))

    of_extra_list, of_rest = [], []
    for a in of_main:
        c = str(a.get("cod") or "")
        if extra_q.get(c, 0) > 0:
            extra_q[c] -= 1
            of_extra_list.append(a)
        else:
            of_rest.append(a)

    # Pas 1: pereche pe cod (FIFO); dublurile ref refolosesc ultimul offer cu
    # acelasi cod (comparatorul face N:M — nu genereaza LIPSA pentru ele).
    of_pool: dict = defaultdict(deque)
    for a in of_rest:
        of_pool[_k(a.get("cod"))].append(id(a))
    of_by_id = {id(a): a for a in of_rest}
    last_by_cod: dict = {}
    paired: dict = {}
    used_of = set()
    for ref_art, is_lipsa in ref_flags:
        if is_lipsa:
            continue
        rk = _k(ref_art.get("cod"))
        if of_pool.get(rk):
            oid = of_pool[rk].popleft()
            paired[id(ref_art)] = of_by_id[oid]
            used_of.add(oid)
            last_by_cod[rk] = of_by_id[oid]
        elif rk in last_by_cod:
            paired[id(ref_art)] = last_by_cod[rk]  # N:M reuse pt dubluri
    # Pas 2: pozitional pentru resturile de pe ambele parti (ordine document)
    rest_ref = [a for a, fl in ref_flags if not fl and id(a) not in paired]
    rest_of = [a for a in of_rest if id(a) not in used_of]
    for ref_art, of_art in zip(rest_ref, rest_of):
        paired[id(ref_art)] = of_art
        used_of.add(id(of_art))

    rows: List[RowTuple] = []
    for ref_art, is_lipsa in ref_flags:
        of_art = None if is_lipsa else paired.get(id(ref_art))
        pair_ncs = [nc for nc in _ncs_for(ncs, ref_art, of_art)
                    if nc.get("tip") not in _SUPPRESSED_TIPS]
        if of_art is None and not is_lipsa:
            # comparatorul NU a marcat LIPSA — nu inventa eticheta (rest N:M)
            fill, nc_texts = None, [_nc_text(nc) for nc in pair_ncs]
        else:
            fill, nc_texts = _pair_fill(ref_art, of_art, pair_ncs, of_main, ref_main)
        rows.append((ref_art, of_art, fill, nc_texts))
        parent_nr = _major_nr(ref_art)
        if parent_nr:
            rows.extend(_sub_rows(ref_arts, of_arts, ncs, parent_nr, fill))

    leftover_of = [a for a in of_rest if id(a) not in used_of]
    for of_art in leftover_of:
        # nematchuit dar fara NC EXTRA — afiseaza fara eticheta
        rows.append((None, of_art, None, []))
    for of_art in of_extra_list:
        pair_ncs = [nc for nc in _ncs_for(ncs, None, of_art)
                    if nc.get("tip") not in _SUPPRESSED_TIPS]
        fill, nc_texts = _pair_fill(None, of_art, pair_ncs, of_main, ref_main)
        rows.append((None, of_art, fill, nc_texts))
    return rows


def _build_matched_rows(group: dict) -> List[RowTuple]:
    ncs      = group.get("neconformitati", [])
    ref_arts = group.get("ref_articles", [])
    of_arts  = group.get("oferta_articles", [])

    ref_main = sorted([a for a in ref_arts if not a.get("is_component")], key=_sort_key)
    of_main  = sorted([a for a in of_arts  if not a.get("is_component")], key=_sort_key)

    matches = [m for m in (group.get("matches") or []) if isinstance(m, dict)]
    if matches:
        return _build_rows_from_matches(ref_main, of_main, ref_arts, of_arts, ncs, matches)

    # EXTRA oferta articles identified by (cod, nr_ordine) — not cod alone
    # (multiple articles can share a cod; only the specific nr_ordine is EXTRA)
    extra_of_keys = {
        (nc["oferta_cod"], nc.get("nr_ordine_oferta"))
        for nc in ncs
        if nc.get("tip") == "ARTICOL_EXTRA" and nc.get("oferta_cod")
    }

    def _art_key(a: dict):
        return (a.get("cod"), a.get("nr_ordine"))

    ref_by_nr = {_major_nr(a): a for a in ref_main}
    of_normal = [a for a in of_main if _art_key(a) not in extra_of_keys]
    of_extra  = [a for a in of_main if _art_key(a) in extra_of_keys]
    of_by_nr  = {_major_nr(a): a for a in of_normal}

    # nrs from ref + nrs in both (normal oferta); extras deferred to end
    paired_nrs = sorted(set(ref_by_nr) | set(of_by_nr))

    rows: List[RowTuple] = []

    for nr in paired_nrs:
        ref_art = ref_by_nr.get(nr)
        of_art  = of_by_nr.get(nr)

        pair_ncs  = [nc for nc in _ncs_for(ncs, ref_art, of_art)
                     if nc.get("tip") not in _SUPPRESSED_TIPS]
        nc_texts  = [_nc_text(nc) for nc in pair_ncs]

        if ref_art and not of_art:
            fill = FILL_LIPSA
            if not nc_texts:
                nc_texts = ["LIPSĂ: absent din ofertă"]
            suggestion = _fuzzy_suggest(ref_art, of_main)
            if suggestion:
                nc_texts.append(suggestion)
        elif not ref_art and of_art:
            fill = FILL_EXTRA
            if not nc_texts:
                nc_texts = ["EXTRA: absent din referință"]
            suggestion = _fuzzy_suggest(of_art, ref_main)
            if suggestion:
                nc_texts.append(suggestion)
        elif pair_ncs:
            tips = {nc.get("tip") for nc in pair_ncs}
            if "COD_NORMATIV_DIFERIT" in tips:
                fill = FILL_COD_NORM
            elif "SPECIFICATIE_DIFERITA" in tips:
                fill = FILL_SPEC_DIFF
            elif "DIFERENTA_PARAMETRU" in tips:
                fill = FILL_PARAM_DIFF
            elif "DIFERENTA_CAMP" in tips or "UM_DIFERIT" in tips:
                fill = FILL_REAL_NC
            elif tips <= {"COD_SIMILAR", "DESCRIERE_DIFERITA"}:
                fill = FILL_COD_SIMILAR
            else:
                fill = FILL_NC
        else:
            fill = None

        rows.append((ref_art, of_art, fill, nc_texts))
        rows.extend(_sub_rows(ref_arts, of_arts, ncs, nr, fill))

    # EXTRA oferta articles appended at end
    for of_art in of_extra:
        pair_ncs = [nc for nc in _ncs_for(ncs, None, of_art)
                    if nc.get("tip") not in _SUPPRESSED_TIPS]
        nc_texts = [_nc_text(nc) for nc in pair_ncs] or ["EXTRA: absent din referință"]
        suggestion = _fuzzy_suggest(of_art, ref_main)
        if suggestion:
            nc_texts.append(suggestion)
        rows.append((None, of_art, FILL_EXTRA, nc_texts))
        # Only include subs whose parent_code matches this extra article (avoid double-counting
        # subs that were already paired in the LIPSA section above)
        _extra_cod = of_art.get("cod")
        _of_arts_for_extra = [a for a in of_arts if not a.get("is_component") or a.get("parent_code") == _extra_cod]
        rows.extend(_sub_rows([], _of_arts_for_extra, ncs, _major_nr(of_art), FILL_EXTRA))

    return rows


def _build_only_rows(group: dict, side: str) -> List[RowTuple]:
    arts = sorted(group.get("articles", []), key=_sort_key)
    fill = FILL_LIPSA if side == "ref" else FILL_EXTRA
    msg  = "Grup absent din ofertă" if side == "ref" else "Grup absent din referință"
    rows: List[RowTuple] = []
    for art in arts:
        if side == "ref":
            rows.append((art, None, fill, [msg] if not art.get("is_component") else []))
        else:
            rows.append((None, art, fill, [msg] if not art.get("is_component") else []))
    return rows


# ── Table / cell writing ──────────────────────────────────────────────────────

def _apply_widths(tbl: Table) -> None:
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(COL_W[i])
            if i in NO_WRAP:
                _no_wrap(cell)


def _build_header(tbl: Table) -> None:
    _apply_widths(tbl)
    r0 = tbl.rows[0].cells
    r1 = tbl.rows[1].cells

    def _hcell(cell, text):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        _shade(cell, FILL_HDR)

    # Row 0: group spans
    r0[0].merge(r0[4]);  _hcell(r0[0], "REFERINȚĂ")
    r0[5].merge(r0[9]);  _hcell(r0[5], "OFERTĂ")
    r0[10].merge(r1[10]); _hcell(r0[10], "Neconformități")

    # Row 1: column sub-labels
    sub = ["Nr", "Cod", "Denumire", "UM", "Cant"]
    for base in (0, 5):
        for i, lbl in enumerate(sub):
            cell = r1[base + i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(lbl)
            run.bold = True
            run.font.size = Pt(7)
            _shade(cell, FILL_HDR)


def _write_5_cols(cells, art: Optional[dict], offset: int, fill: Optional[str]) -> None:
    is_comp = (art or {}).get("is_component", False)
    fs      = 6 if is_comp else 7
    grey    = RGBColor(0x55, 0x55, 0x55) if is_comp else None

    def _wcell(rel, val, center=False, right=False, two_lines: Optional[str] = None):
        cell = cells[offset + rel]
        cell.text = ""
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(val)
        run.font.size = Pt(fs)
        if grey:
            run.font.color.rgb = grey
        if two_lines is not None:
            # Second line in smaller, greyed text
            p.add_run().add_break()
            r2 = p.add_run(two_lines)
            r2.font.size = Pt(fs - 1)
            r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    if art is None:
        for rel in range(5):
            cells[offset + rel].text = ""
        return

    nr_text, page_text = _fmt_nr_with_page(art)
    cod = art.get("cod", "") or ""
    if not is_comp and cod.startswith("$"):
        cod = cod[1:]
    # Hide synthetic SPEC_ prefix — these are embedded material specs, not real codes
    if cod.startswith("SPEC_"):
        cod = ""
    _wcell(0, nr_text, center=True, two_lines=page_text if page_text else None)
    _wcell(1, cod)
    _wcell(2, art.get("denumire", "") or "")
    _wcell(3, art.get("um", "") or "", center=True)
    _wcell(4, _fmt_cant(art.get("cantitate")), right=True)


def _write_data_row(row, ref_art, of_art, fill, nc_texts) -> None:
    cells = row.cells

    if fill:
        for cell in cells:
            _shade(cell, fill)
    for i in NO_WRAP:
        _no_wrap(cells[i])

    _write_5_cols(cells, ref_art, offset=0, fill=fill)
    _write_5_cols(cells, of_art,  offset=5, fill=fill)

    nc_cell = cells[10]
    nc_cell.text = ""
    if nc_texts:
        p = nc_cell.paragraphs[0]
        for i, t in enumerate(nc_texts):
            if i > 0:
                p.add_run().add_break()
            run = p.add_run(t)
            run.font.size = Pt(6)
            run.font.color.rgb = TEXT_RED


# ── Group section ─────────────────────────────────────────────────────────────

def _header_label(arts: list) -> str:
    """Extract 'obiectivul | obiectul | categoria' from first article with deviz_header."""
    for a in arts:
        dh = a.get("deviz_header") or {}
        parts = [dh.get("obiectivul", ""), dh.get("obiectul", ""), dh.get("categoria", "")]
        label = " | ".join(p for p in parts if p)
        if label:
            return label
    return "?"


def _min_ref_page(group: dict, group_type: str) -> int:
    """Return min source_page from ref side — used to sort groups by DI order."""
    if group_type == "oferta_only":
        arts = group.get("articles", [])
    elif group_type == "ref_only":
        arts = group.get("articles", [])
    else:
        arts = group.get("ref_articles", [])
    pages = [p for a in arts for p in (a.get("source_pages") or [])]
    return min(pages) if pages else 9999


def _write_group_section(doc: Document, group: dict, group_type: str, nc_only: bool = False) -> None:
    ref_arts = group.get("ref_articles") or (group.get("articles", []) if group_type == "ref_only" else [])
    of_arts  = group.get("oferta_articles") or (group.get("articles", []) if group_type == "oferta_only" else [])

    if group_type == "matched":
        rows = _build_matched_rows(group)
    elif group_type == "ref_only":
        rows = _build_only_rows(group, "ref")
    else:
        rows = _build_only_rows(group, "oferta")

    if nc_only:
        rows = [(r, o, f, n) for r, o, f, n in rows if f is not None]

    if not rows:
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True

    if group_type == "matched":
        ref_label = _header_label(ref_arts)
        of_label  = _header_label(of_arts)
        r1 = p.add_run(f"REF:    {ref_label}")
        r1.bold = True; r1.font.size = Pt(8)
        p.add_run().add_break()
        r2 = p.add_run(f"OFERTĂ: {of_label}")
        r2.bold = True; r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x00, 0x4C, 0x99)
    elif group_type == "ref_only":
        r1 = p.add_run(f"REF:    {_header_label(ref_arts)}")
        r1.bold = True; r1.font.size = Pt(8)
    else:
        r2 = p.add_run(f"OFERTĂ: {_header_label(of_arts)}")
        r2.bold = True; r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x00, 0x4C, 0x99)

    tbl = doc.add_table(rows=2 + len(rows) + 1, cols=N_COLS)
    tbl.style = "Table Grid"
    _fixed_layout(tbl)
    _set_tbl_grid(tbl)
    _set_cell_margins(tbl)
    _build_header(tbl)
    _repeat_header_rows(tbl)

    for i, (ref_art, of_art, fill, nc_texts) in enumerate(rows):
        _write_data_row(tbl.rows[2 + i], ref_art, of_art, fill, nc_texts)

    # Total row
    total_row = tbl.rows[-1]
    total_row.cells[0].merge(total_row.cells[10])
    tc = total_row.cells[0]
    tc.text = ""
    n_ref = sum(1 for r, o, f, n in rows if r and not r.get("is_component"))
    n_of  = sum(1 for r, o, f, n in rows if o and not o.get("is_component"))
    run = tc.paragraphs[0].add_run(
        f"Total grup: REFERINȚĂ {n_ref} articole principale | OFERTĂ {n_of} articole principale"
    )
    run.bold = True
    run.font.size = Pt(7)
    _shade(tc, FILL_TOTAL)


# ── Entry point ───────────────────────────────────────────────────────────────

def _strip_components(holistic: dict) -> dict:
    """Copy of holistic without subcomponent articles and their NCs.

    Reference documents (listing format) carry almost no subcomponents while
    eDevize offers list every resource — side-by-side component rows are
    one-sided noise.
    """
    result = dict(holistic)
    for key in ("matched_groups", "ref_only_groups", "oferta_only_groups"):
        groups = holistic.get(key)
        if not isinstance(groups, list):
            continue
        new_groups = []
        for g in groups:
            g2 = dict(g)
            for akey in ("ref_articles", "oferta_articles", "articles"):
                if isinstance(g2.get(akey), list):
                    g2[akey] = [a for a in g2[akey] if not a.get("is_component")]
            if isinstance(g2.get("neconformitati"), list):
                g2["neconformitati"] = [nc for nc in g2["neconformitati"]
                                        if not nc.get("is_component")]
            new_groups.append(g2)
        result[key] = new_groups
    return result


def build_comparatie_docx(holistic: dict, client_name: str, oferta_nr: int, output_path: str,
                          nc_only: bool = False, include_components: bool = True) -> None:
    """Generate landscape A4 side-by-side comparison DOCX.

    nc_only=True → Raport extract: only rows with neconformitati (LIPSA/EXTRA/NC).
    nc_only=False → Full comparatie: all rows including clean matches.
    include_components=False → hide subcomponent rows (resources) entirely.
    """
    from shared.holistic_filters import strip_codeless
    holistic = strip_codeless(holistic)
    if not include_components:
        holistic = _strip_components(holistic)
    doc = Document()
    _set_landscape(doc)

    if nc_only:
        title = f"Raport Neconformități Ofertă {oferta_nr} — {client_name}"
    else:
        title = f"Comparație Referință vs Ofertă {oferta_nr} — {client_name}"

    for text, bold, size in [
        (title, True, 14),
        (f"Generat: {date.today().isoformat()}", False, 9),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    # Sort each group list by minimum ref source_page → preserves DI referință order
    matched   = sorted(holistic.get("matched_groups", []),      key=lambda g: _min_ref_page(g, "matched"))
    ref_only  = sorted(holistic.get("ref_only_groups", []),     key=lambda g: _min_ref_page(g, "ref_only"))
    of_only   = sorted(holistic.get("oferta_only_groups", []),  key=lambda g: _min_ref_page(g, "oferta_only"))

    for group in matched:
        _write_group_section(doc, group, "matched", nc_only=nc_only)

    for group in ref_only:
        _write_group_section(doc, group, "ref_only", nc_only=nc_only)

    for group in of_only:
        _write_group_section(doc, group, "oferta_only", nc_only=nc_only)

    doc.save(output_path)
