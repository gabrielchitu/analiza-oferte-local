"""Generate F3 landscape DOCX from verified extracted sursa-incarcare data."""

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment

# Column widths in cm: Nr | Denumire | UM | Cant | Pret | Total
_COL_W = [1.2, 9.0, 1.5, 2.5, 2.8, 3.0]

_STOPWORDS = {'DE', 'LA', 'PE', 'SI', 'IN', 'CU', 'DIN', 'A', 'AL', 'SA', 'O', 'UN'}

_GRAY_HEX = 'EEEEEE'
_YELLOW_HEX = 'FFF2CC'
_RED_HEX = 'FF0000'
_HEADER_GRAY = 'D9D9D9'

_XLS_GRAY   = PatternFill('solid', fgColor='EEEEEE')
_XLS_YELLOW = PatternFill('solid', fgColor='FFF2CC')
_XLS_RED    = PatternFill('solid', fgColor='FF0000')
_XLS_HEADER = PatternFill('solid', fgColor='D9D9D9')
_XLS_PU_HDR = PatternFill('solid', fgColor='DCE6F1')  # blue tint for PU group
_XLS_VAL_HDR= PatternFill('solid', fgColor='EBF1DE')  # green tint for Val group
_XLS_BOLD        = Font(bold=True)
_XLS_BOLD_WHITE  = Font(bold=True, color='FFFFFF')
_XLS_SMALL       = Font(size=8, italic=True)
_XLS_SUBITEM     = Font(size=8)
_XLS_CENTER = Alignment(horizontal='center', vertical='center', wrap_text=True)
_XLS_RIGHT  = Alignment(horizontal='right',  vertical='center')
_XLS_LEFT   = Alignment(horizontal='left',   vertical='center', wrap_text=True)

# 17-col layout: 7 desc + 5 PU + 5 Val (matches DOCX v2)
_XLSX_COLS = [
    ('Nr.',          5),   # 0
    ('Nr.crt',       7),   # 1
    ('Cod',         12),   # 2
    ('Cod principal',14),  # 3
    ('Denumire',    40),   # 4
    ('UM',           7),   # 5
    ('Cantitate',   11),   # 6
    ('PU Mat',      12),   # 7
    ('PU Man',      12),   # 8
    ('PU Util',     12),   # 9
    ('PU Trans',    12),   # 10
    ('PU Total',    12),   # 11
    ('Val Mat',     14),   # 12
    ('Val Man',     14),   # 13
    ('Val Util',    14),   # 14
    ('Val Trans',   14),   # 15
    ('Val Total',   14),   # 16
]


def make_acronym(obiectivul: str) -> str:
    """Generate max-6-char acronym from OBIECTIVUL string.

    Leading numeric tokens (e.g. '0232 000000232') are stripped before processing.

    Examples:
        '0232 000000232 DRUMURI TATARANI'               -> 'DT'
        'CONSTRUIRE UNITATE DE CAZARE - TARGOVISTE'     -> 'CUCT'
    """
    # Strip leading numeric-only tokens (codes like '0232 000000232')
    text = re.sub(r'^\d[\d\s]+', '', obiectivul).strip()
    # Keep only Romanian-alphabet letters and spaces (drop punctuation/hyphens)
    text = re.sub(r'[^A-ZĂÂÎȘȚ ]', '', text.upper())
    words = text.split()
    letters = [w[0] for w in words if w and w not in _STOPWORDS and len(w) >= 2]
    return ''.join(letters[:6])


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _set_landscape(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)


def _set_tbl_grid(tbl) -> None:
    tblGrid = OxmlElement('w:tblGrid')
    for w_cm in _COL_W:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(round(w_cm * 567))))
        tblGrid.append(gc)
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn('w:tblPr'))
    if tbl_pr is not None:
        tbl_pr.addnext(tblGrid)
    else:
        tbl_el.insert(0, tblGrid)


def _set_cell_margins(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    for side, twips in [('top', 28), ('left', 57), ('bottom', 28), ('right', 57)]:
        elem = OxmlElement(f'w:{side}')
        elem.set(qn('w:w'), str(twips))
        elem.set(qn('w:type'), 'dxa')
        tblCellMar.append(elem)
    tblPr.append(tblCellMar)


def _repeat_header_rows(tbl, n_rows: int = 2) -> None:
    for row in tbl.rows[:n_rows]:
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        trPr.append(OxmlElement('w:tblHeader'))


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _fmt_num(value: float, decimals: int = 2) -> str:
    if value == 0.0:
        return '0.00'
    return f'{value:,.{decimals}f}'


def _cell_write(cell, text: str, bold: bool = False, size: float = 8,
                center: bool = False, right: bool = False,
                color: str | None = None) -> None:
    cell.text = ''
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_header_rows(tbl, obiectivul: str, obiectul: str, categoria: str) -> None:
    row0 = tbl.rows[0]
    for i in range(1, 6):
        row0.cells[0].merge(row0.cells[i])
    cell = row0.cells[0]
    cell.text = ''
    for label, value in [('Obiectivul', obiectivul), ('Obiectul', obiectul), ('Categoria', categoria)]:
        p = cell.add_paragraph()
        run = p.add_run(f'{label}: {value}')
        run.font.size = Pt(9)
        run.font.bold = True
    # remove the empty first paragraph that python-docx adds by default
    first_p = cell.paragraphs[0]._p
    first_p.getparent().remove(first_p)

    row1 = tbl.rows[1]
    headers = [
        'Nr.',
        'Capitol de lucrări',
        'U.M.',
        'Cantitatea',
        'Preț unitar (fără TVA) — Lei',
        'TOTALUL (fără TVA) — Lei',
    ]
    for i, h in enumerate(headers):
        _shade_cell(row1.cells[i], _HEADER_GRAY)
        _cell_write(row1.cells[i], h, bold=True, size=7.5, center=True)


def _add_capitol_row(tbl, titlu: str) -> None:
    row = tbl.add_row()
    for i in range(2, 6):
        row.cells[1].merge(row.cells[i])
    _shade_cell(row.cells[0], _GRAY_HEX)
    _shade_cell(row.cells[1], _GRAY_HEX)
    _cell_write(row.cells[0], '', bold=True, size=8)
    _cell_write(row.cells[1], titlu, bold=True, size=8.5)


def _add_article_row(tbl, art: dict) -> None:
    row = tbl.add_row()
    cod_den = (f"{art['cod']} - {art['denumire']}"
               if art.get('cod') else art.get('denumire', ''))
    vals = [
        art.get('nr_crt', ''),
        cod_den,
        art.get('um', ''),
        _fmt_num(art.get('cantitate', 0), 3),
        _fmt_num(art.get('pret_unitar', 0)),
        _fmt_num(art.get('total', 0)),
    ]
    for i, v in enumerate(vals):
        _cell_write(row.cells[i], v, size=8,
                    center=(i in {0, 2}), right=(i in {3, 4, 5}))


def _add_breakdown_rows(tbl, breakdown: dict) -> None:
    for key in ('material', 'manopera', 'utilaj', 'transport'):
        bd = breakdown.get(key, {})
        row = tbl.add_row()
        _cell_write(row.cells[0], '', size=7)
        cell1 = row.cells[1]
        cell1.text = ''
        p = cell1.paragraphs[0]
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{key}:")
        run.font.size = Pt(7)
        run.font.italic = True
        _cell_write(row.cells[2], '', size=7)
        _cell_write(row.cells[3], '', size=7)
        _cell_write(row.cells[4], _fmt_num(bd.get('pret', 0)), size=7, right=True)
        _cell_write(row.cells[5], _fmt_num(bd.get('total', 0)), size=7, right=True)


def _add_sub_item_row(tbl, sub: dict) -> None:
    row = tbl.add_row()
    cod_den = (f"{sub['cod']} - {sub['denumire']}"
               if sub.get('cod') else sub.get('denumire', ''))
    vals = [
        sub.get('nr_crt', ''),
        cod_den,
        sub.get('um', ''),
        _fmt_num(sub.get('cantitate', 0), 3),
        _fmt_num(sub.get('pret_unitar', 0)),
        _fmt_num(sub.get('total', 0)),
    ]
    for i, v in enumerate(vals):
        _cell_write(row.cells[i], v, size=7.5,
                    center=(i in {0, 2}), right=(i in {3, 4, 5}))


def _add_total_capitol_row(tbl, titlu: str, total: float) -> None:
    row = tbl.add_row()
    for i in range(1, 5):
        row.cells[0].merge(row.cells[i])
    _cell_write(row.cells[0], f'TOTAL {titlu}', bold=True, size=8)
    _cell_write(row.cells[5], _fmt_num(total), bold=True, size=8, right=True)


def _add_total_deviz_row(tbl, total: float, is_red: bool = False) -> None:
    row = tbl.add_row()
    for i in range(1, 5):
        row.cells[0].merge(row.cells[i])
    if is_red:
        for cell in row.cells:
            _shade_cell(cell, _RED_HEX)
        _cell_write(row.cells[0],
                    'TOTAL NECONFIRMAT — verificare manuală necesară',
                    bold=True, size=8, color='FFFFFF')
        _cell_write(row.cells[5], _fmt_num(total), bold=True, size=8,
                    right=True, color='FFFFFF')
    else:
        for cell in row.cells:
            _shade_cell(cell, _YELLOW_HEX)
        _cell_write(row.cells[0], 'TOTAL 1 (Cheltuieli directe)', bold=True, size=9)
        _cell_write(row.cells[5], _fmt_num(total), bold=True, size=9, right=True)


def _build_table(doc: Document, deviz: dict) -> None:
    tbl = doc.add_table(rows=2, cols=6)
    tbl.style = 'Table Grid'
    _set_tbl_grid(tbl)
    _set_cell_margins(tbl)
    _repeat_header_rows(tbl, n_rows=2)
    _add_header_rows(tbl,
                     deviz.get('obiectivul', ''),
                     deviz.get('obiectul', ''),
                     deviz.get('categoria', ''))

    for cap in deviz.get('capitole', []):
        _add_capitol_row(tbl, cap['titlu'])
        for art in cap.get('articole', []):
            _add_article_row(tbl, art)
            if art.get('breakdown'):
                _add_breakdown_rows(tbl, art['breakdown'])
            for sub in art.get('sub_items', []):
                _add_sub_item_row(tbl, sub)
        if cap.get('total_capitol') is not None:
            _add_total_capitol_row(tbl, cap['titlu'], cap['total_capitol'])

    is_red = deviz.get('status') == 'RED'
    _add_total_deviz_row(tbl, deviz.get('total_deviz', 0.0), is_red)


def write_docx(devize: list[dict], output_path: Path) -> None:
    """Write F3 landscape DOCX with one table per deviz."""
    doc = Document()
    _set_landscape(doc)
    for idx, deviz in enumerate(devize):
        if idx > 0:
            doc.add_page_break()
        _build_table(doc, deviz)
    doc.save(str(output_path))


def write_xlsx(devize: list[dict], output_path: Path, footer: dict | None = None) -> None:
    """17-col XLS matching DOCX v2 landscape layout: 7 desc + 5 PU + 5 Val."""
    wb = Workbook()
    wb.remove(wb.active)

    N = len(_XLSX_COLS)  # 17

    def _num(v):
        try:
            return float(v) if v else 0.0
        except (TypeError, ValueError):
            return 0.0

    for deviz in devize:
        sheet_name = (deviz.get('categoria') or 'Deviz')[:31]
        ws = wb.create_sheet(title=sheet_name)

        # Col widths
        for ci, (_, w) in enumerate(_XLSX_COLS, 1):
            ws.column_dimensions[ws.cell(1, ci).column_letter].width = w

        r = 1
        # Row 0: deviz header merged across all 17 cols
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=N)
        header_val = (
            f"Obiectivul: {deviz.get('obiectivul', '')}  |  "
            f"Obiectul: {deviz.get('obiectul', '')}  |  "
            f"Categoria: {deviz.get('categoria', '')}"
        )
        cell = ws.cell(r, 1, value=header_val)
        cell.font = _XLS_BOLD
        cell.alignment = Alignment(horizontal='left', wrap_text=True, vertical='top')
        cell.fill = _XLS_HEADER
        ws.row_dimensions[r].height = 36
        r += 1

        # Row 1: group labels (cols 1-7 blank, 8-12 "Pret unitar", 13-17 "Total lei")
        ws.merge_cells(start_row=r, start_column=8, end_row=r, end_column=12)
        c = ws.cell(r, 8, value='Pret unitar (lei/UM)')
        c.font = _XLS_BOLD; c.fill = _XLS_PU_HDR; c.alignment = _XLS_CENTER
        for ci in range(8, 13): ws.cell(r, ci).fill = _XLS_PU_HDR

        ws.merge_cells(start_row=r, start_column=13, end_row=r, end_column=17)
        c = ws.cell(r, 13, value='Total (lei)')
        c.font = _XLS_BOLD; c.fill = _XLS_VAL_HDR; c.alignment = _XLS_CENTER
        for ci in range(13, 18): ws.cell(r, ci).fill = _XLS_VAL_HDR
        r += 1

        # Row 2: column labels
        labels_desc = ['Nr.', 'Nr.crt', 'Cod', 'Cod principal', 'Denumire', 'UM', 'Cantitate']
        labels_pu   = ['Material', 'Manoperă', 'Utilaj', 'Transport', 'Total']
        labels_val  = ['Material', 'Manoperă', 'Utilaj', 'Transport', 'Total']
        all_labels  = labels_desc + labels_pu + labels_val
        for ci, lbl in enumerate(all_labels, 1):
            c = ws.cell(r, ci, value=lbl)
            c.font = _XLS_BOLD
            c.alignment = _XLS_CENTER
            if ci <= 7:
                c.fill = _XLS_HEADER
            elif ci <= 12:
                c.fill = _XLS_PU_HDR
            else:
                c.fill = _XLS_VAL_HDR
        ws.row_dimensions[r].height = 30
        ws.freeze_panes = ws.cell(r + 1, 1)
        r += 1

        seq = 0
        for cap in deviz.get('capitole', []):
            # Capitol header
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=N)
            c = ws.cell(r, 1, value=cap.get('titlu', ''))
            c.font = _XLS_BOLD; c.fill = _XLS_GRAY
            c.alignment = Alignment(horizontal='left', vertical='center')
            for ci in range(2, N + 1): ws.cell(r, ci).fill = _XLS_GRAY
            r += 1

            for art in cap.get('articole', []):
                if not art.get('cod') and not art.get('denumire') and not (art.get('total') or art.get('cantitate')):
                    continue
                seq += 1
                bd = art.get('breakdown') or {}

                def _pu(comp): return _num((bd.get(comp) or {}).get('pret'))
                def _val(comp):
                    v = _num((bd.get(comp) or {}).get('total'))
                    if v: return v
                    return _num(art.get('cantitate')) * _pu(comp)

                pu_mat, pu_man, pu_util, pu_tran = _pu('material'), _pu('manopera'), _pu('utilaj'), _pu('transport')
                pu_tot = pu_mat + pu_man + pu_util + pu_tran
                val_mat  = _val('material')
                val_man  = _val('manopera')
                val_util = _val('utilaj')
                val_tran = _val('transport')
                val_tot  = _num(art.get('total')) or (val_mat + val_man + val_util + val_tran)
                cant     = _num(art.get('cantitate'))

                is_sub = bool(art.get('parent_cod'))
                fn = Font(size=8) if is_sub else None

                def _w(col, val, align=_XLS_RIGHT):
                    c = ws.cell(r, col, value=val if val != 0.0 else None)
                    c.alignment = align
                    if fn: c.font = fn
                    return c

                _w(1, seq, _XLS_CENTER)
                _w(2, art.get('nr_crt', ''), _XLS_CENTER)
                _w(3, art.get('cod', '') or '', _XLS_LEFT)
                _w(4, art.get('parent_cod', '') or '', _XLS_LEFT)
                _w(5, art.get('denumire', ''), _XLS_LEFT)
                _w(6, art.get('um', ''), _XLS_CENTER)
                _w(7, cant or None)
                _w(8, pu_mat or None)
                _w(9, pu_man or None)
                _w(10, pu_util or None)
                _w(11, pu_tran or None)
                _w(12, pu_tot or None)
                _w(13, val_mat or None)
                _w(14, val_man or None)
                _w(15, val_util or None)
                _w(16, val_tran or None)
                _w(17, val_tot or None)

                # Format numeric cols as number
                for ci in range(7, 18):
                    ws.cell(r, ci).number_format = '#,##0.00'
                r += 1

        # Total deviz row
        is_red = deviz.get('status') == 'RED'
        total_label = 'TOTAL NECONFIRMAT — verificare manuală necesară' if is_red else 'TOTAL (Cheltuieli directe)'
        fill = _XLS_RED if is_red else _XLS_YELLOW
        font = _XLS_BOLD_WHITE if is_red else _XLS_BOLD
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=16)
        for ci in range(1, N + 1):
            ws.cell(r, ci).fill = fill
        c = ws.cell(r, 1, value=total_label)
        c.font = font
        c.alignment = Alignment(horizontal='left', vertical='center')
        c = ws.cell(r, 17, value=_num(deviz.get('total_deviz')) or None)
        c.font = font
        c.alignment = _XLS_RIGHT
        c.number_format = '#,##0.00'
        r += 1

    # Recapitulatie sheet — 4 grand-total rows
    ws_r = wb.create_sheet(title='Recapitulatie')
    for ci, (_, w) in enumerate(_XLSX_COLS, 1):
        ws_r.column_dimensions[ws_r.cell(1, ci).column_letter].width = w
    t1 = sum(d.get('total_deviz', 0.0) or 0.0 for d in devize)
    ft = footer or {}
    recap_rows = [
        ('TOTAL 1 (Cheltuieli directe)',   t1),
        ('Total general (fără TVA)',        ft.get('total_general_fara_tva')),
        (f"TVA ({ft.get('tva_pct', 0.0):.2f}%)", ft.get('tva_val')),
        ('Total general (inclusiv TVA)',    ft.get('total_cu_tva')),
    ]
    for ri, (label, val) in enumerate(recap_rows, 1):
        ws_r.merge_cells(start_row=ri, start_column=1, end_row=ri, end_column=16)
        c = ws_r.cell(ri, 1, value=label)
        c.font = _XLS_BOLD
        c.alignment = Alignment(horizontal='left', vertical='center')
        ws_r.row_dimensions[ri].height = 22
        c = ws_r.cell(ri, 17, value=val)
        c.font = _XLS_BOLD
        c.alignment = _XLS_RIGHT
        c.number_format = '#,##0.00'
        c.fill = _XLS_YELLOW

    wb.save(str(output_path))


def write_pdf(docx_path: Path, output_dir: Path) -> bool:
    """Convert DOCX to PDF via LibreOffice CLI. Returns True on success, False if unavailable."""
    try:
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf',
             '--outdir', str(output_dir), str(docx_path)],
            capture_output=True, timeout=60,
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def write_pdf_native(devize: list[dict], output_path: Path) -> bool:
    """Generate searchable PDF with reportlab (no LibreOffice needed). Returns True on success."""
    try:
        import os
        from reportlab.lib.pagesizes import A4, landscape as rl_landscape
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, PageBreak,
        )
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return False

    # -- Font (Arial with Romanian character support, fallback Helvetica) --
    _FONT = 'Helvetica'
    _FONT_BOLD = 'Helvetica-Bold'
    _arial_r = '/System/Library/Fonts/Supplemental/Arial.ttf'
    _arial_b = '/System/Library/Fonts/Supplemental/Arial Bold.ttf'
    if os.path.exists(_arial_r):
        try:
            pdfmetrics.registerFont(TTFont('_PdfArial', _arial_r))
            _FONT = '_PdfArial'
            if os.path.exists(_arial_b):
                pdfmetrics.registerFont(TTFont('_PdfArialBold', _arial_b))
                _FONT_BOLD = '_PdfArialBold'
            else:
                _FONT_BOLD = '_PdfArial'
        except Exception:
            pass

    # -- Paragraph styles --
    def _ps(name, font, size, leading_mult=1.25):
        return ParagraphStyle(name, fontName=font, fontSize=size,
                              leading=size * leading_mult, spaceAfter=0, spaceBefore=0)

    ps_normal  = _ps('N',  _FONT,      7)
    ps_bold    = _ps('B',  _FONT_BOLD, 7)
    ps_header  = _ps('H',  _FONT_BOLD, 8.5, 1.3)
    ps_colhdr  = _ps('CH', _FONT_BOLD, 7,   1.2)
    ps_small   = _ps('S',  _FONT,      6.5)
    ps_total   = _ps('T',  _FONT_BOLD, 8)
    ps_total_w = _ps('TW', _FONT_BOLD, 8)
    ps_total_w.textColor = colors.white

    # -- Column widths: scaled proportionally to fill A4 landscape (26.7 cm usable) --
    COL_W = [w * cm for w in (1.6, 12.0, 2.0, 3.2, 3.6, 4.3)]

    # -- Colors --
    C_GRAY    = colors.HexColor('#EEEEEE')
    C_DGRAY   = colors.HexColor('#D9D9D9')
    C_YELLOW  = colors.HexColor('#FFF2CC')
    C_RED     = colors.HexColor('#FF0000')

    def _esc(s: str) -> str:
        return str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def _p(text, style):
        return Paragraph(_esc(text), style)

    def _build_table(deviz: dict) -> Table:
        rows = []
        cmds = [
            ('GRID',          (0, 0), (-1, -1), 0.4,  colors.black),
            ('FONTNAME',      (0, 0), (-1, -1), _FONT),
            ('FONTSIZE',      (0, 0), (-1, -1), 7),
            ('TOPPADDING',    (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ('LEFTPADDING',   (0, 0), (-1, -1), 3),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 3),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ]

        # Row 0: deviz header (spans all 6 cols)
        hdr = (
            f"<b>Obiectivul:</b> {_esc(deviz.get('obiectivul', ''))}<br/>"
            f"<b>Obiectul:</b> {_esc(deviz.get('obiectul', ''))}<br/>"
            f"<b>Categoria:</b> {_esc(deviz.get('categoria', ''))}"
        )
        rows.append([Paragraph(hdr, ps_header), '', '', '', '', ''])
        cmds += [
            ('SPAN',          (0, 0), (5, 0)),
            ('BACKGROUND',    (0, 0), (5, 0), C_DGRAY),
            ('TOPPADDING',    (0, 0), (5, 0), 5),
            ('BOTTOMPADDING', (0, 0), (5, 0), 5),
        ]

        # Row 1: column headers
        col_labels = [
            'Nr.', 'Capitol de lucrări', 'U.M.',
            'Cantitatea', 'Preț unitar\n(fără TVA) Lei', 'TOTALUL\n(fără TVA) Lei',
        ]
        rows.append([Paragraph(h, ps_colhdr) for h in col_labels])
        cmds += [
            ('BACKGROUND', (0, 1), (5, 1), C_DGRAY),
            ('ALIGN',      (0, 1), (5, 1), 'CENTER'),
            ('TOPPADDING', (0, 1), (5, 1), 3),
            ('BOTTOMPADDING', (0, 1), (5, 1), 3),
        ]

        ri = 2

        for cap in deviz.get('capitole', []):
            # Chapter header row: col 0 empty, cols 1-5 merged
            rows.append(['', _p(cap['titlu'], ps_bold), '', '', '', ''])
            cmds += [
                ('SPAN',       (1, ri), (5, ri)),
                ('BACKGROUND', (0, ri), (5, ri), C_GRAY),
                ('FONTNAME',   (0, ri), (5, ri), _FONT_BOLD),
            ]
            ri += 1

            for art in cap.get('articole', []):
                cod_den = (f"{_esc(art['cod'])} - {_esc(art['denumire'])}"
                           if art.get('cod') else _esc(art.get('denumire', '')))
                rows.append([
                    art.get('nr_crt', ''),
                    Paragraph(cod_den, ps_normal),
                    art.get('um', ''),
                    _fmt_num(art.get('cantitate', 0), 3),
                    _fmt_num(art.get('pret_unitar', 0)),
                    _fmt_num(art.get('total', 0)),
                ])
                cmds += [
                    ('ALIGN', (0, ri), (0, ri), 'CENTER'),
                    ('ALIGN', (2, ri), (2, ri), 'CENTER'),
                    ('ALIGN', (3, ri), (5, ri), 'RIGHT'),
                ]
                ri += 1

                if art.get('breakdown'):
                    for key in ('material', 'manopera', 'utilaj', 'transport'):
                        bd = art['breakdown'].get(key, {})
                        rows.append([
                            '', Paragraph(f'<i>  {key}:</i>', ps_small), '', '',
                            _fmt_num(bd.get('pret', 0)),
                            _fmt_num(bd.get('total', 0)),
                        ])
                        cmds += [
                            ('FONTSIZE', (0, ri), (5, ri), 6.5),
                            ('ALIGN',    (4, ri), (5, ri), 'RIGHT'),
                        ]
                        ri += 1

                for sub in art.get('sub_items', []):
                    cod_den_s = (f"{_esc(sub['cod'])} - {_esc(sub['denumire'])}"
                                 if sub.get('cod') else _esc(sub.get('denumire', '')))
                    rows.append([
                        sub.get('nr_crt', ''),
                        Paragraph(cod_den_s, ps_small),
                        sub.get('um', ''),
                        _fmt_num(sub.get('cantitate', 0), 3),
                        _fmt_num(sub.get('pret_unitar', 0)),
                        _fmt_num(sub.get('total', 0)),
                    ])
                    cmds += [
                        ('FONTSIZE', (0, ri), (5, ri), 6.5),
                        ('ALIGN', (0, ri), (0, ri), 'CENTER'),
                        ('ALIGN', (2, ri), (2, ri), 'CENTER'),
                        ('ALIGN', (3, ri), (5, ri), 'RIGHT'),
                    ]
                    ri += 1

            if cap.get('total_capitol') is not None:
                rows.append([
                    _p(f'TOTAL {cap["titlu"]}', ps_bold),
                    '', '', '', '',
                    _fmt_num(cap['total_capitol']),
                ])
                cmds += [
                    ('SPAN',     (0, ri), (4, ri)),
                    ('FONTNAME', (0, ri), (5, ri), _FONT_BOLD),
                    ('ALIGN',    (5, ri), (5, ri), 'RIGHT'),
                ]
                ri += 1

        # Deviz total row
        is_red = deviz.get('status') == 'RED'
        total_label = ('TOTAL NECONFIRMAT — verificare manuală necesară'
                       if is_red else 'TOTAL 1 (Cheltuieli directe)')
        bg = C_RED if is_red else C_YELLOW
        lbl_style = ps_total_w if is_red else ps_total

        rows.append([
            Paragraph(_esc(total_label), lbl_style),
            '', '', '', '',
            _fmt_num(deviz.get('total_deviz', 0.0)),
        ])
        cmds += [
            ('SPAN',       (0, ri), (4, ri)),
            ('BACKGROUND', (0, ri), (5, ri), bg),
            ('FONTNAME',   (0, ri), (5, ri), _FONT_BOLD),
            ('FONTSIZE',   (0, ri), (5, ri), 8),
            ('ALIGN',      (5, ri), (5, ri), 'RIGHT'),
            ('TOPPADDING', (0, ri), (5, ri), 3),
            ('BOTTOMPADDING', (0, ri), (5, ri), 3),
        ]
        if is_red:
            cmds.append(('TEXTCOLOR', (5, ri), (5, ri), colors.white))

        tbl = Table(rows, colWidths=COL_W, repeatRows=2)
        tbl.setStyle(TableStyle(cmds))
        return tbl

    # -- Build document --
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=rl_landscape(A4),
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )
    story = []
    for idx, deviz in enumerate(devize):
        if idx > 0:
            story.append(PageBreak())
        story.append(_build_table(deviz))

    doc.build(story)
    return True


# ── V2: Template-exact format ──────────────────────────────────────────────

# Landscape A4 (26.7cm usable): 7 desc + 5 PU + 5 Val = 17 cols, sum=15120 twips
_COL_WIDTHS_TWIPS_V2 = [
    250, 300, 720, 720, 2270, 380, 550,   # 0-6: desc
    993, 993, 993, 993, 993,               # 7-11: Pret unitar (Material/Man/Util/Trans/Total)
    993, 993, 993, 993, 993,               # 12-16: Total valoare (Material/Man/Util/Trans/Total)
]
_NO_WRAP_COLS_V2 = {0, 1, 2, 3, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16}
_HEADER_FILL_V2 = 'D9D9D9'
_TOTAL_FILL_V2 = 'F2F2F2'


def _set_cell_width_twips_v2(cell, twips: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)


def _set_cell_no_wrap_v2(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(OxmlElement('w:noWrap'))


def _suppress_borders_v2(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_tbl_grid_v2(tbl) -> None:
    tbl_el = tbl._tbl
    existing = tbl_el.find(qn('w:tblGrid'))
    if existing is not None:
        tbl_el.remove(existing)
    tblGrid = OxmlElement('w:tblGrid')
    for w in _COL_WIDTHS_TWIPS_V2:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl_el.insert(0, tblGrid)


def _shade_cell_v2(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell_write_v2(cell, text: str, bold: bool = False, size: float = 8,
                   center: bool = False, right: bool = False) -> None:
    cell.text = ''
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fmt_price_ro(value) -> str:
    """None/0 → ''; else Romanian locale: 1.234,50."""
    if not value:
        return ''
    formatted = f'{value:,.2f}'
    return formatted.replace(',', 'X').replace('.', ',').replace('X', '.')


def _build_sursa_table_header(tbl) -> None:
    """Write 2-row header: 7 desc cols + 5 PU cols + 5 Val cols = 17 cols."""
    row0 = tbl.rows[0].cells
    row1 = tbl.rows[1].cells

    for ci in range(17):
        _set_cell_width_twips_v2(row0[ci], _COL_WIDTHS_TWIPS_V2[ci])
        _set_cell_width_twips_v2(row1[ci], _COL_WIDTHS_TWIPS_V2[ci])
        if ci in _NO_WRAP_COLS_V2:
            _set_cell_no_wrap_v2(row0[ci])
            _set_cell_no_wrap_v2(row1[ci])

    # Cols 0-6: vertical merge
    for i, label in enumerate(["Nr.", "Nr.crt", "Cod", "Cod principal", "Denumire", "UM", "Cantitate"]):
        tbl.cell(0, i).merge(tbl.cell(1, i))
        _cell_write_v2(row0[i], label, bold=True, center=True)
        _shade_cell_v2(row0[i], _HEADER_FILL_V2)

    # Cols 7-11: horizontal merge "Pret unitar (lei/UM)"
    row0[7].merge(row0[11])
    _cell_write_v2(row0[7], "Pret unitar (lei/UM)", bold=True, center=True)
    _shade_cell_v2(row0[7], _HEADER_FILL_V2)
    for i, label in enumerate(["Material", "Manoperă", "Utilaj", "Transport", "Total"]):
        _cell_write_v2(row1[7 + i], label, bold=True, center=True)
        _shade_cell_v2(row1[7 + i], _HEADER_FILL_V2)

    # Cols 12-16: horizontal merge "Total (lei)"
    row0[12].merge(row0[16])
    _cell_write_v2(row0[12], "Total (lei)", bold=True, center=True)
    _shade_cell_v2(row0[12], _HEADER_FILL_V2)
    for i, label in enumerate(["Material", "Manoperă", "Utilaj", "Transport", "Total"]):
        _cell_write_v2(row1[12 + i], label, bold=True, center=True)
        _shade_cell_v2(row1[12 + i], _HEADER_FILL_V2)


def _write_sursa_row_v2(row, seq_nr: int, art: dict, is_sub: bool, parent_cod: str) -> None:
    bd = art.get('breakdown') or {}
    font_size = 7 if is_sub else 8

    cantitate = art.get('cantitate') or 0.0

    def _pu(comp): return (bd.get(comp) or {}).get('pret') or 0.0
    def _val(comp):
        v = (bd.get(comp) or {}).get('total')
        if v:
            return v
        # fallback: cantitate × pret dacă total lipsește din JSON
        return cantitate * _pu(comp)

    pu_mat, pu_man, pu_util, pu_tran = _pu('material'), _pu('manopera'), _pu('utilaj'), _pu('transport')
    pu_total = pu_mat + pu_man + pu_util + pu_tran

    val_mat  = _val('material')
    val_man  = _val('manopera')
    val_util = _val('utilaj')
    val_tran = _val('transport')
    val_total = art.get('total') or (val_mat + val_man + val_util + val_tran)

    values = [
        str(seq_nr),
        str(art.get('nr_crt', '')),
        art.get('cod', '') or '',
        parent_cod or '',
        art.get('denumire', ''),
        art.get('um', ''),
        f"{cantitate:.2f}" if cantitate else '',
        _fmt_price_ro(pu_mat),
        _fmt_price_ro(pu_man),
        _fmt_price_ro(pu_util),
        _fmt_price_ro(pu_tran),
        _fmt_price_ro(pu_total),
        _fmt_price_ro(val_mat),
        _fmt_price_ro(val_man),
        _fmt_price_ro(val_util),
        _fmt_price_ro(val_tran),
        _fmt_price_ro(val_total),
    ]
    right_cols = {1, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16}
    center_cols = {0, 5}

    for ci, (cell, val) in enumerate(zip(row.cells, values)):
        cell.text = ''
        p = cell.paragraphs[0]
        if ci in right_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif ci in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(font_size)
        if is_sub:
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _build_sursa_group(doc, deviz: dict) -> None:
    """Add group header paragraph + 17-col table for one deviz (landscape)."""
    obiectivul = deviz.get('obiectivul', '')
    obiectul   = deviz.get('obiectul', '')
    categoria  = deviz.get('categoria', '')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    parts = [s.strip() for s in [obiectivul, obiectul, categoria] if s.strip()]
    run = p.add_run(' | '.join(parts))
    run.bold = True
    run.font.size = Pt(9)

    # Flatten capitole → articole + sub_items
    flat: list[tuple[bool, dict, str]] = []  # (is_sub, art, parent_cod)
    for cap in deviz.get('capitole', []):
        for art in cap.get('articole', []):
            if not art.get('cod') and not art.get('denumire'):
                continue
            flat.append((False, art, ''))
            for sub in art.get('sub_items', []):
                if not sub.get('cod') and not sub.get('denumire'):
                    continue
                flat.append((True, sub, art.get('cod', '') or ''))

    main_count = sum(1 for is_sub, _, _ in flat if not is_sub)
    sub_count  = sum(1 for is_sub, _, _ in flat if is_sub)

    n_rows = 2 + len(flat) + 1
    tbl = doc.add_table(rows=n_rows, cols=17)
    tbl.style = 'Table Grid'
    _set_tbl_grid_v2(tbl)
    tblPr_el = tbl._tbl.find(qn('w:tblPr'))
    if tblPr_el is not None:
        tbl_total = sum(_COL_WIDTHS_TWIPS_V2)
        existing_tblW = tblPr_el.find(qn('w:tblW'))
        if existing_tblW is not None:
            tblPr_el.remove(existing_tblW)
        tblW_el = OxmlElement('w:tblW')
        tblW_el.set(qn('w:w'), str(tbl_total))
        tblW_el.set(qn('w:type'), 'dxa')
        tblPr_el.append(tblW_el)

    _build_sursa_table_header(tbl)

    for i, (is_sub, art, parent_cod) in enumerate(flat):
        _write_sursa_row_v2(tbl.rows[2 + i], seq_nr=i + 1,
                            art=art, is_sub=is_sub, parent_cod=parent_cod)

    total_row = tbl.rows[-1]
    total_row.cells[0].merge(total_row.cells[16])
    cell = total_row.cells[0]
    cell.text = ''
    suffix = f' / {sub_count} subcomponente' if sub_count else ''
    run = cell.paragraphs[0].add_run(
        f'Total grup: {main_count} articole principale{suffix}'
    )
    run.bold = True
    run.font.size = Pt(8)
    _shade_cell_v2(cell, _TOTAL_FILL_V2)


def write_docx_v2(devize: list, output_path, metadata: dict | None = None, footer: dict | None = None) -> None:
    """Write F3-format DOCX identical with Template_exact.docx."""
    from datetime import date as _date
    meta = metadata or {}

    doc = Document()
    section = doc.sections[0]
    # Landscape A4
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    beneficiar  = meta.get('beneficiar', '') or meta.get('client', '')
    executant   = meta.get('executant', '')  or meta.get('ofertant', '')
    proiectant  = meta.get('proiectant', '')
    gen_date    = meta.get('date', _date.today().isoformat())

    header_lines = [('Lista articole', True, 14)]
    if beneficiar:
        header_lines.append((f'Beneficiar: {beneficiar}', False, 11))
    if executant:
        header_lines.append((f'Executant: {executant}', False, 11))
    if proiectant:
        header_lines.append((f'Proiectant: {proiectant}', False, 11))
    header_lines.append((f'Generat: {gen_date}', False, 9))

    for line, bold, size in header_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)

    for deviz in devize:
        _build_sursa_group(doc, deviz)

    # Recapitulatie generala
    t1 = sum(d.get('total_deviz', 0.0) or 0.0 for d in devize)
    ft = footer or {}
    recap_data = [
        ('TOTAL 1 (Cheltuieli directe)',   t1),
        ('Total general (fără TVA)',        ft.get('total_general_fara_tva')),
        (f"TVA ({ft.get('tva_pct', 0.0):.2f}%)", ft.get('tva_val')),
        ('Total general (inclusiv TVA)',    ft.get('total_cu_tva')),
    ]
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Recapitulație generală')
    run.bold = True
    run.font.size = Pt(10)
    tbl_r = doc.add_table(rows=len(recap_data), cols=2)
    tbl_r.style = 'Table Grid'
    for i, (label, val) in enumerate(recap_data):
        # label cell — left aligned
        c0 = tbl_r.rows[i].cells[0]
        c0.text = ''
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9)
        # value cell — right aligned
        c1 = tbl_r.rows[i].cells[1]
        c1.text = ''
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(_fmt_price_ro(val or 0.0))
        r1.bold = True
        r1.font.size = Pt(9)

    doc.save(str(output_path))
