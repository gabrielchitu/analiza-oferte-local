"""Word report generation from v2 holistic JSON.

Generates Word documents from v1-compatible holistic JSON structure
produced by MatchingOrchestratorV2. Output format matches v1 reports.
"""

from pathlib import Path
from datetime import date
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _shade_cell(cell, fill_hex: str):
    """Shade a cell with a background color (hex, no # prefix)."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def _format_article(article: Dict) -> str:
    """Format single article for display."""
    cod = article.get("cod", "N/A")
    um = article.get("um", "")
    cant = article.get("cant", "")
    desc = article.get("denumire", "")

    parts = [f"[{cod}]"]
    if um:
        parts.append(f"({um})")
    if cant:
        parts.append(f"x{cant}")
    if desc:
        parts.append(f"- {desc}")

    return " ".join(parts)


def generate_holistic_report(
    holistic: Dict,
    output_path: Path,
    client_name: Optional[str] = None,
) -> None:
    """Generate Word report from v2 holistic matching JSON.

    Args:
        holistic: v1-compatible holistic JSON dict with structure:
                  {
                      "client": str,
                      "di_file": str,
                      "extraction_version": str,
                      "matched_groups": [group_dict, ...],
                      "ref_only_groups": [group_dict, ...],
                      "oferta_only_groups": [group_dict, ...],
                      "stats": {
                          "matched_articles_count": int,
                          "ref_only_articles_count": int,
                          "oferta_only_articles_count": int,
                          ...
                      }
                  }
        output_path: Path where to save .docx file
        client_name: Optional client name override (for header)

    Returns:
        None (writes file to output_path)
    """
    doc = Document()

    # Set landscape
    section = doc.sections[0]
    section.page_height = Inches(8.5)
    section.page_width = Inches(11)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Title
    title = doc.add_heading("Raport Matching - Set-Based v2", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    client = client_name or holistic.get("client", "Unknown")
    doc.add_paragraph(f"Client: {client}")
    doc.add_paragraph(f"DI File: {holistic.get('di_file', 'N/A')}")
    doc.add_paragraph(f"Extraction Version: {holistic.get('extraction_version', 'N/A')}")
    doc.add_paragraph(f"Report Date: {date.today()}")

    # Summary Table
    doc.add_heading("Summary", level=1)
    stats = holistic.get("stats", {})

    table = doc.add_table(rows=8, cols=2)
    table.style = "Light Grid Accent 1"

    rows = table.rows
    _shade_cell(rows[0].cells[0], "D3D3D3")
    _shade_cell(rows[0].cells[1], "D3D3D3")

    rows[0].cells[0].text = "Metric"
    rows[0].cells[1].text = "Count"

    metrics = [
        ("Matched Groups", stats.get("matched_groups_count", 0)),
        ("Ref-Only Groups", stats.get("ref_only_groups_count", 0)),
        ("Oferta-Only Groups", stats.get("oferta_only_groups_count", 0)),
        ("Matched Articles", stats.get("matched_articles_count", 0)),
        ("Ref-Only Articles", stats.get("ref_only_articles_count", 0)),
        ("Oferta-Only Articles", stats.get("oferta_only_articles_count", 0)),
        ("Total Groups", (
            stats.get("matched_groups_count", 0)
            + stats.get("ref_only_groups_count", 0)
            + stats.get("oferta_only_groups_count", 0)
        )),
    ]

    for i, (metric, value) in enumerate(metrics, 1):
        rows[i].cells[0].text = metric
        rows[i].cells[1].text = str(value)

    # Matched Groups
    matched_groups = holistic.get("matched_groups", [])
    if matched_groups:
        doc.add_heading("Matched Groups", level=1)
        for group in matched_groups:
            _add_group_section(doc, group, "matched")

    # Ref-Only Groups
    ref_only_groups = holistic.get("ref_only_groups", [])
    if ref_only_groups:
        doc.add_heading("Reference-Only Groups (Not in Offer)", level=1)
        for group in ref_only_groups:
            _add_group_section(doc, group, "ref_only")

    # Oferta-Only Groups
    oferta_only_groups = holistic.get("oferta_only_groups", [])
    if oferta_only_groups:
        doc.add_heading("Offer-Only Groups (Not in Reference)", level=1)
        for group in oferta_only_groups:
            _add_group_section(doc, group, "oferta_only")

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _add_group_section(doc: Document, group: Dict, group_type: str) -> None:
    """Add a group section to the document.

    Args:
        doc: Document to add to
        group: Group dict with deviz_den and articole
        group_type: "matched", "ref_only", or "oferta_only"
    """
    # Group heading
    deviz_den = group.get("deviz_den", "Unknown Group")
    heading_text = deviz_den
    if group_type == "ref_only":
        heading_text += " [REF ONLY]"
    elif group_type == "oferta_only":
        heading_text += " [OFERTA ONLY]"

    doc.add_heading(heading_text, level=2)

    # Articles table
    articole = group.get("articole", [])
    if not articole:
        doc.add_paragraph("(No articles)")
        return

    table = doc.add_table(rows=len(articole) + 1, cols=5)
    table.style = "Light Grid Accent 1"

    # Header
    header = table.rows[0]
    for cell in header.cells:
        _shade_cell(cell, "D3D3D3")

    header.cells[0].text = "Cod"
    header.cells[1].text = "UM"
    header.cells[2].text = "Cantitate"
    header.cells[3].text = "Descriere"
    header.cells[4].text = "Status"

    # Rows
    for i, article in enumerate(articole, 1):
        row = table.rows[i]
        row.cells[0].text = article.get("cod", "")
        row.cells[1].text = article.get("um", "")
        row.cells[2].text = str(article.get("cant", ""))
        row.cells[3].text = article.get("denumire", "")

        # Status based on match
        if group_type == "matched":
            status = "Matched"
        elif group_type == "ref_only":
            status = "Not in offer"
        else:
            status = "Not in ref"
        row.cells[4].text = status

    # Group totals
    doc.add_paragraph(f"Total: {len(articole)} articles in this group")
