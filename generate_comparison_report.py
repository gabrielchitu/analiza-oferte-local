#!/usr/bin/env python3
"""
Generate side-by-side Word report comparing v1 vs v2 extraction on same client.
Usage: python3 generate_comparison_report.py "Drum Tatarani"
"""
import json
import sys
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
OUTPUT_BASE = ROOT / "output_AO"

def load_json(path):
    with open(path) as f:
        return json.load(f)

def count_articles(data):
    if isinstance(data, dict) and "grupos" in data:
        return sum(len(g.get("articole", [])) for g in data.get("grupos", []))
    return 0

def extract_article_summary(data):
    """Extract article details for comparison."""
    articles = []
    if isinstance(data, dict) and "grupos" in data:
        for grupo in data.get("grupos", []):
            for art in grupo.get("articole", []):
                articles.append({
                    "deviz": grupo.get("deviz_cod"),
                    "nr": art.get("nr") or art.get("cod"),
                    "descriere": art.get("descriere")[:50] if art.get("descriere") else "",
                    "um": art.get("um"),
                    "cant": art.get("cantitate"),
                    "source": art.get("extraction_source", "UNKNOWN"),
                })
    return articles

def generate_comparison_report(client_name, oferta_num=1):
    client_dir = OUTPUT_BASE / client_name

    # Load both versions
    v1_file = client_dir / f"holistic_oferta_{oferta_num}.json"
    v2_file = client_dir / f"oferta_{oferta_num}_v2.json"

    if not v1_file.exists() or not v2_file.exists():
        print(f"Missing files for {client_name} oferta_{oferta_num}")
        return

    v1_data = load_json(v1_file)
    v2_data = load_json(v2_file)

    v1_count = count_articles(v1_data)
    v2_count = count_articles(v2_data)

    v1_articles = extract_article_summary(v1_data)
    v2_articles = extract_article_summary(v2_data)

    # Create Word doc
    doc = Document()
    doc.add_heading(f"Comparison Report: v1 vs v2", 0)
    doc.add_paragraph(f"Client: {client_name}")
    doc.add_paragraph(f"Oferta: {oferta_num}")
    doc.add_paragraph(f"Date: {date.today()}")

    # Summary table
    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "v1 (LLM-Powered)"
    hdr[2].text = "v2 (Table-Native)"

    row1 = table.rows[1].cells
    row1[0].text = "Total Articles"
    row1[1].text = str(v1_count)
    row1[2].text = str(v2_count)

    row2 = table.rows[2].cells
    row2[0].text = "Difference"
    row2[1].text = ""
    diff = v2_count - v1_count
    pct = (diff / v1_count * 100) if v1_count > 0 else 0
    row2[2].text = f"+{diff} ({pct:+.1f}%)"

    # Article details (first 20)
    doc.add_heading("Article Sample (first 20)", level=1)

    comp_table = doc.add_table(rows=min(21, max(len(v1_articles), len(v2_articles)) + 1), cols=6)
    comp_table.style = "Light Grid Accent 1"

    hdr = comp_table.rows[0].cells
    hdr[0].text = "Idx"
    hdr[1].text = "v1 Descriere"
    hdr[2].text = "v1 Um"
    hdr[3].text = "v2 Descriere"
    hdr[4].text = "v2 Um"
    hdr[5].text = "Match"

    for i in range(min(20, max(len(v1_articles), len(v2_articles)))):
        row = comp_table.rows[i + 1].cells
        row[0].text = str(i + 1)

        if i < len(v1_articles):
            art_v1 = v1_articles[i]
            row[1].text = art_v1["descriere"]
            row[2].text = art_v1["um"] or ""

        if i < len(v2_articles):
            art_v2 = v2_articles[i]
            row[3].text = art_v2["descriere"]
            row[4].text = art_v2["um"] or ""

        if i < len(v1_articles) and i < len(v2_articles):
            match = "✓" if v1_articles[i]["descriere"] == v2_articles[i]["descriere"] else "✗"
            row[5].text = match

    # Findings
    doc.add_heading("Findings", level=1)
    findings = doc.add_paragraph()
    findings.add_run(f"v1 Articles: {v1_count}\n")
    findings.add_run(f"v2 Articles: {v2_count}\n")
    findings.add_run(f"Difference: +{v2_count - v1_count} (+{(v2_count - v1_count) / v1_count * 100 if v1_count > 0 else 0:.1f}%)\n")
    findings.add_run(f"\nExtraction Sources in v2:\n")

    sources = {}
    for art in v2_articles:
        src = art["source"]
        sources[src] = sources.get(src, 0) + 1

    for src, count in sorted(sources.items()):
        findings.add_run(f"  - {src}: {count}\n")

    # Save
    output_path = client_dir / f"Raport_Comparatie_Oferta_{oferta_num}_v1_vs_v2.docx"
    doc.save(output_path)
    print(f"✓ Saved: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 generate_comparison_report.py 'ClientName' [oferta_num]")
        sys.exit(1)

    client = sys.argv[1]
    for oferta in [1, 2]:
        generate_comparison_report(client, oferta)
